from fastapi import FastAPI, Request, Form, UploadFile, File, HTTPException, BackgroundTasks, Header
from fastapi.responses import HTMLResponse, RedirectResponse, StreamingResponse, JSONResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from email_campaign_module import init_email_campaign_module
from sqlalchemy import create_engine, text
from sqlalchemy.exc import IntegrityError
from pathlib import Path
from pydantic import BaseModel
from io import BytesIO
import openpyxl
import subprocess
import shutil
import zipfile
from docx import Document
import os
import time
import urllib.parse
import urllib.request
import urllib.error
import socket
import hashlib
import secrets
import json
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from typing import Optional, Any
from datetime import datetime, timedelta

BASE_DIR = Path(__file__).resolve().parent
APP_TITLE = "TradeFlow CRM"
SETTINGS_FILE = BASE_DIR / "crm_settings.json"
DEFAULT_SQLITE_URL = f"sqlite:///{(BASE_DIR / 'altahhan_crm.db').as_posix()}"
TRACKING_BASE_URL = os.getenv("TRACKING_BASE_URL", "").strip()
COOKIE_SECURE = os.getenv("COOKIE_SECURE", "true").strip().lower() in {"1", "true", "yes", "on"}
INITIAL_ADMIN_USERNAME = os.getenv("INITIAL_ADMIN_USERNAME", "admin").strip() or "admin"
INITIAL_ADMIN_PASSWORD = os.getenv("INITIAL_ADMIN_PASSWORD", "").strip()
UPLOADS_DIR = BASE_DIR / "uploads"
UPLOADS_DIR.mkdir(exist_ok=True)
CHAT_UPLOADS_DIR = UPLOADS_DIR / "chat"
CHAT_UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
AVATAR_UPLOADS_DIR = UPLOADS_DIR / "avatars"
AVATAR_UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
LEAD_UPLOADS_DIR = UPLOADS_DIR / "leads"
LEAD_UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
CLIENT_UPLOADS_DIR = UPLOADS_DIR / "clients"
CLIENT_UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
DATA_DIR = BASE_DIR / "data"
DATA_DIR.mkdir(exist_ok=True)
TRADE_INTELLIGENCE_FILE = DATA_DIR / "trade_intelligence.json"
EXPORT_CATALOG_FILE = DATA_DIR / "export_catalog.json"

PRODUCT_PROFILE_FILE = BASE_DIR / "product_profile.json"
BOOTSTRAP_ADMIN_NOTE_FILE = BASE_DIR / "bootstrap_admin.txt"


def resolve_bootstrap_admin_password() -> str:
    if INITIAL_ADMIN_PASSWORD:
        return INITIAL_ADMIN_PASSWORD
    generated = f"TF-{secrets.token_hex(5)}"
    note = (
        "TradeFlow CRM bootstrap admin account\n"
        f"username: {INITIAL_ADMIN_USERNAME}\n"
        f"password: {generated}\n\n"
        "Rotate this password immediately after first login or set INITIAL_ADMIN_PASSWORD in the environment.\n"
    )
    try:
        BOOTSTRAP_ADMIN_NOTE_FILE.write_text(note, encoding="utf-8")
    except Exception:
        pass
    return generated


def load_product_profile() -> dict:
    default = {
        "product_name": "TradeFlow CRM",
        "edition_name": "Altahhan Edition",
        "company_name": "Al Tahhan Golden Dates",
        "app_title": "TradeFlow CRM — Altahhan Edition",
        "hero_tagline": "Sales, outreach, shipments and trade operations in one bilingual workspace.",
        "login_title": "Export & import operations platform",
        "login_subtitle": "A polished internal platform for leads, outreach, shipment follow-up and trade documentation.",
        "landing_tagline": "The first deployment for Al Tahhan, designed as a reusable product for future companies.",
        "primary_market": "Export & import departments",
        "support_label": "Bilingual EN / AR experience",
        "outreach_name": "Outreach Hub",
    }
    if PRODUCT_PROFILE_FILE.exists():
        try:
            data = json.loads(PRODUCT_PROFILE_FILE.read_text(encoding="utf-8"))
            if isinstance(data, dict):
                default.update({k: v for k, v in data.items() if v not in (None, "")})
        except Exception:
            pass
    return default


PRODUCT_PROFILE = load_product_profile()
APP_TITLE = PRODUCT_PROFILE.get("app_title") or APP_TITLE


def load_trade_intelligence_seed() -> dict:
    if not TRADE_INTELLIGENCE_FILE.exists():
        return {"agreements": [], "laws": []}
    try:
        return json.loads(TRADE_INTELLIGENCE_FILE.read_text(encoding="utf-8"))
    except Exception:
        return {"agreements": [], "laws": []}


def maybe_seed_trade_intelligence():
    payload = load_trade_intelligence_seed()
    agreements = payload.get('agreements') or []
    laws = payload.get('laws') or []
    if not agreements and not laws:
        return
    with engine.begin() as conn:
        total = conn.execute(text("SELECT COUNT(*) FROM trade_reference_items")).scalar() or 0
        if total:
            return
        for item in agreements + laws:
            conn.execute(text("""
                INSERT INTO trade_reference_items
                (title, category, item_type, region, partner_countries, effective_year, status, summary, benefits, rules_of_origin, source_org, source_url, tags, sort_order, created_at, updated_at)
                VALUES
                (:title, :category, :item_type, :region, :partner_countries, :effective_year, :status, :summary, :benefits, :rules_of_origin, :source_org, :source_url, :tags, :sort_order, :created_at, :updated_at)
            """), {
                'title': item.get('title','').strip(),
                'category': item.get('category','agreement').strip() or 'agreement',
                'item_type': item.get('type','').strip(),
                'region': item.get('region','').strip(),
                'partner_countries': item.get('partner_countries','').strip(),
                'effective_year': item.get('effective_year','').strip(),
                'status': item.get('status','').strip(),
                'summary': item.get('summary','').strip(),
                'benefits': item.get('benefits','').strip(),
                'rules_of_origin': item.get('rules_of_origin','').strip(),
                'source_org': item.get('source_org','').strip(),
                'source_url': item.get('source_url','').strip(),
                'tags': item.get('tags','').strip(),
                'sort_order': 0 if item.get('category') == 'agreement' else 100,
                'created_at': time.time(),
                'updated_at': time.time(),
            })


def agreement_match_score(country: str, item) -> int:
    country = (country or '').strip().lower()
    hay = f"{getattr(item, 'partner_countries', '')} {getattr(item, 'tags', '')} {getattr(item, 'title', '')}".lower()
    if not country:
        return 0
    score = 0
    if country in hay:
        score += 100
    # lightweight country-family boosts
    region_map = {
        'morocco': ['agadir', 'gafta', 'arab'],
        'jordan': ['agadir', 'gafta', 'arab'],
        'tunisia': ['agadir', 'gafta', 'arab'],
        'germany': ['eu', 'european union', 'europe'],
        'france': ['eu', 'european union', 'europe'],
        'italy': ['eu', 'european union', 'europe'],
        'spain': ['eu', 'european union', 'europe'],
        'united kingdom': ['uk', 'britain'],
        'uk': ['uk', 'britain'],
        'turkey': ['turkey', 'türkiye'],
        'türkiye': ['turkey', 'türkiye'],
        'serbia': ['serbia'],
        'switzerland': ['efta'],
        'norway': ['efta'],
        'iceland': ['efta'],
        'liechtenstein': ['efta'],
        'brazil': ['mercosur'],
        'argentina': ['mercosur'],
        'uruguay': ['mercosur'],
        'paraguay': ['mercosur'],
    }
    for token in region_map.get(country, []):
        if token in hay:
            score += 25
    if country in {'egypt'} and getattr(item, 'category', '') == 'law':
        score += 15
    return score



def load_runtime_settings() -> dict:
    default = {
        "db_driver": "sqlite",
        "db_server": "",
        "db_name": "",
        "db_user": "",
        "db_password": "",
        "domain": "",
        "image_prefix": "/uploads",
    }
    if SETTINGS_FILE.exists():
        try:
            data = json.loads(SETTINGS_FILE.read_text(encoding="utf-8"))
            if isinstance(data, dict):
                default.update(data)
        except Exception:
            pass
    return default


def save_runtime_settings(data: dict):
    current = load_runtime_settings()
    current.update(data or {})
    SETTINGS_FILE.write_text(json.dumps(current, ensure_ascii=False, indent=2), encoding="utf-8")


def build_db_url_from_settings(settings: dict) -> str:
    env_url = os.getenv("CRM_DB_URL", "").strip()
    if env_url:
        return env_url

    require_remote_db = os.getenv("RENDER") or os.getenv("REQUIRE_REMOTE_DB", "").strip().lower() in {"1", "true", "yes", "on"}
    if require_remote_db:
        raise RuntimeError("CRM_DB_URL is required in Render/production environment.")

    driver = (settings.get("db_driver") or "sqlite").strip().lower()
    if driver == "mysql":
        server = (settings.get("db_server") or "").strip()
        db_name = (settings.get("db_name") or "").strip()
        db_user = urllib.parse.quote_plus((settings.get("db_user") or "").strip())
        db_password = urllib.parse.quote_plus(settings.get("db_password") or "")
        if server and db_name and db_user:
            return f"mysql+pymysql://{db_user}:{db_password}@{server}/{db_name}?charset=utf8mb4"
    if driver in {"postgres", "postgresql"}:
        server = (settings.get("db_server") or "").strip()
        db_name = (settings.get("db_name") or "").strip()
        db_user = urllib.parse.quote_plus((settings.get("db_user") or "").strip())
        db_password = urllib.parse.quote_plus(settings.get("db_password") or "")
        if server and db_name and db_user:
            return f"postgresql+psycopg://{db_user}:{db_password}@{server}/{db_name}"
    return DEFAULT_SQLITE_URL


def make_engine(db_url: str):
    engine_kwargs = {"future": True, "pool_pre_ping": True}
    if db_url.startswith("sqlite"):
        engine_kwargs["connect_args"] = {"check_same_thread": False, "timeout": 30}
    return create_engine(db_url, **engine_kwargs)


RUNTIME_SETTINGS = load_runtime_settings()
ACTIVE_DB_URL = build_db_url_from_settings(RUNTIME_SETTINGS)
engine = make_engine(ACTIVE_DB_URL)


def reload_engine_from_settings():
    global RUNTIME_SETTINGS, ACTIVE_DB_URL, engine
    RUNTIME_SETTINGS = load_runtime_settings()
    ACTIVE_DB_URL = build_db_url_from_settings(RUNTIME_SETTINGS)
    try:
        engine.dispose()
    except Exception:
        pass
    engine = make_engine(ACTIVE_DB_URL)
    init_db()


app = FastAPI(title=APP_TITLE)
app.mount("/static", StaticFiles(directory=str(BASE_DIR / "static")), name="static")
app.mount("/uploads", StaticFiles(directory=str(UPLOADS_DIR)), name="uploads")
templates = Jinja2Templates(directory=str(BASE_DIR / "templates"))
templates.env.globals["branding"] = PRODUCT_PROFILE
templates.env.globals["product_version"] = "V17 Shipment-Centric"


def datetimeformat_filter(value):
    try:
        return time.strftime('%Y-%m-%d', time.localtime(int(float(value))))
    except Exception:
        return ''


templates.env.filters["datetimeformat"] = datetimeformat_filter


# ----------------------------
# Security / utility helpers
# ----------------------------
def hash_password(password: str) -> str:
    salt = secrets.token_hex(16)
    digest = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt.encode("utf-8"), 200000).hex()
    return f"pbkdf2_sha256$200000${salt}${digest}"


def verify_password(password: str, stored_value: str) -> bool:
    if not stored_value:
        return False
    if stored_value.startswith("pbkdf2_sha256$"):
        try:
            _, rounds, salt, digest = stored_value.split("$", 3)
            check = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt.encode("utf-8"), int(rounds)).hex()
            return secrets.compare_digest(check, digest)
        except Exception:
            return False
    return secrets.compare_digest(password, stored_value)


def safe_int(value, default=0):
    try:
        return int(value)
    except Exception:
        return default


def normalize_stage(value: str) -> str:
    allowed = {'Lead', 'Contacted', 'Negotiation', 'Quotation', 'Won', 'Lost'}
    value = (value or 'Lead').strip().title()
    return value if value in allowed else 'Lead'


def calculate_lead_score(data: dict) -> int:
    score = 0
    if data.get('email'):
        score += 10
    if data.get('phone'):
        score += 10
    if data.get('contact_person'):
        score += 10
    if data.get('country'):
        score += 5
    if data.get('source'):
        score += 8
    if data.get('assigned_to'):
        score += 7
    if data.get('tags'):
        score += 5
    stage = normalize_stage(data.get('stage') or 'Lead')
    score += {
        'Lead': 5,
        'Contacted': 15,
        'Negotiation': 30,
        'Quotation': 40,
        'Won': 60,
        'Lost': 0,
    }.get(stage, 5)
    try:
        value = float(data.get('estimated_value') or 0)
    except Exception:
        value = 0
    if value >= 100000:
        score += 25
    elif value >= 25000:
        score += 15
    elif value > 0:
        score += 8
    return max(0, min(int(score), 100))


def get_preferred_lan_ip() -> str:
    candidates = []
    try:
        for item in socket.getaddrinfo(socket.gethostname(), None, family=socket.AF_INET):
            ip = item[4][0]
            if not ip.startswith("127."):
                candidates.append(ip)
    except Exception:
        pass

    # UDP trick usually returns the active NIC IP.
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        ip = s.getsockname()[0]
        s.close()
        if ip and not ip.startswith("127."):
            candidates.insert(0, ip)
    except Exception:
        pass

    seen = []
    unique = []
    for ip in candidates:
        if ip not in seen:
            seen.append(ip)
            unique.append(ip)

    for prefix in ("192.168.", "10.", "172."):
        for ip in unique:
            if ip.startswith(prefix):
                return ip
    return unique[0] if unique else "127.0.0.1"


def build_tracking_base_url(request: Optional[Request] = None) -> str:
    if TRACKING_BASE_URL:
        return TRACKING_BASE_URL.rstrip("/")
    domain = (RUNTIME_SETTINGS.get("domain") or "").strip()
    if domain:
        if domain.startswith("http://") or domain.startswith("https://"):
            return domain.rstrip("/")
        return f"https://{domain}".rstrip("/")
    if request is not None:
        host = request.headers.get("host") or f"{get_preferred_lan_ip()}:8000"
        proto = request.headers.get("x-forwarded-proto") or request.url.scheme or "http"
        return f"{proto}://{host}".rstrip("/")
    return f"http://{get_preferred_lan_ip()}:8000"


def normalize_prefix(prefix: str) -> str:
    prefix = (prefix or "/uploads").strip()
    if not prefix.startswith("/"):
        prefix = "/" + prefix
    return prefix.rstrip("/") or "/uploads"


def build_image_base_url(request: Optional[Request] = None) -> str:
    return build_tracking_base_url(request) + normalize_prefix(RUNTIME_SETTINGS.get("image_prefix") or "/uploads")


def public_upload_url(filename: str, request: Optional[Request] = None) -> str:
    return build_image_base_url(request) + "/" + urllib.parse.quote(filename)


def redirect(url: str) -> RedirectResponse:
    return RedirectResponse(url, status_code=302)


def insert_and_get_id(conn, sql: str, params: dict):
    sql_clean = sql.strip()
    dialect = engine.dialect.name
    if dialect in {"sqlite", "postgresql"}:
        try:
            return conn.execute(text(sql_clean + "\nRETURNING id"), params).scalar_one()
        except Exception:
            pass
    result = conn.execute(text(sql_clean), params)
    inserted_id = getattr(result, "lastrowid", None)
    if inserted_id:
        return int(inserted_id)
    if dialect == "mysql":
        return int(conn.execute(text("SELECT LAST_INSERT_ID()")).scalar_one())
    raise RuntimeError("Could not determine inserted id")



def safe_filename(filename: str) -> str:
    keep = []
    for ch in (filename or ''):
        if ch.isalnum() or ch in {'-', '_', '.', ' '}:
            keep.append(ch)
    cleaned = ''.join(keep).strip().replace(' ', '_')
    return cleaned[:180] or f"file_{int(time.time())}"


def chat_file_kind(filename: str) -> str:
    ext = (Path(filename).suffix or '').lower()
    if ext in {'.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp'}:
        return 'image'
    if ext == '.pdf':
        return 'pdf'
    return ''

def display_username(row) -> str:
    return (getattr(row, 'display_name', '') or getattr(row, 'username', '') or '').strip()


def avatar_url(avatar_path: str) -> str:
    avatar_path = (avatar_path or '').strip()
    if not avatar_path:
        return '/static/assets/logo.png'
    return f"/uploads/avatars/{avatar_path}"


def render_chat_text(value: str) -> str:
    value = (value or '').strip()
    if value.startswith('[[sticker:') and value.endswith(']]'):
        name = value[10:-2].strip().lower()
        stickers = {
            'party': '🥳', 'fire': '🔥', 'rocket': '🚀', 'love': '❤️', 'ok': '👌',
            'wow': '😮', 'boss': '😎', 'coffee': '☕', 'gift': '🎁', 'check': '✅'
        }
        return stickers.get(name, '✨')
    return value


def channel_cookie_key(channel_id: int) -> str:
    return f"chat_room_ok_{safe_int(channel_id, 0)}"


def is_room_password_verified(request: Request, channel_id: int) -> bool:
    return request.cookies.get(channel_cookie_key(channel_id), '') == '1'


def user_can_access_channel(conn, username: str, role: str, channel_id: int) -> bool:
    channel = conn.execute(text("SELECT * FROM chat_channels WHERE id=:id AND is_active=1"), {'id': channel_id}).fetchone()
    if not channel:
        return False
    if role == 'admin' or not safe_int(getattr(channel, 'only_members', 0), 0):
        return True
    member = conn.execute(text("SELECT id FROM chat_channel_members WHERE channel_id=:cid AND username=:u"), {'cid': channel_id, 'u': username}).fetchone()
    return bool(member)


def user_can_write_channel(conn, username: str, role: str, channel_id: int) -> bool:
    if role == 'admin':
        return True
    channel = conn.execute(text("SELECT * FROM chat_channels WHERE id=:id AND is_active=1"), {'id': channel_id}).fetchone()
    if not channel:
        return False
    if not safe_int(getattr(channel, 'only_members', 0), 0):
        return True
    member = conn.execute(text("SELECT * FROM chat_channel_members WHERE channel_id=:cid AND username=:u"), {'cid': channel_id, 'u': username}).fetchone()
    return bool(member) and safe_int(getattr(member, 'can_write', 1), 1) == 1


def list_accessible_channels(conn, user) -> list[Any]:
    if user.role == 'admin':
        return conn.execute(text("SELECT * FROM chat_channels WHERE is_active=1 ORDER BY id ASC")).fetchall()
    return conn.execute(text("""
        SELECT c.*
        FROM chat_channels c
        LEFT JOIN chat_channel_members m ON m.channel_id=c.id AND m.username=:u
        WHERE c.is_active=1 AND (COALESCE(c.only_members,0)=0 OR m.id IS NOT NULL)
        GROUP BY c.id
        ORDER BY c.id ASC
    """), {'u': user.username}).fetchall()


def verify_channel_password(conn, channel_id: int, password: str) -> bool:
    channel = conn.execute(text("SELECT password_hash FROM chat_channels WHERE id=:id"), {'id': channel_id}).fetchone()
    stored = getattr(channel, 'password_hash', '') if channel else ''
    if not stored:
        return True
    return verify_password(password or '', stored)


def room_member_rows(conn, channel_id: int):
    return conn.execute(text("SELECT * FROM chat_channel_members WHERE channel_id=:cid ORDER BY username ASC"), {'cid': channel_id}).fetchall()



def current_clients_seed_file() -> Path:
    return DATA_DIR / 'Morocco clients.xlsx'


def sync_current_clients_from_excel(filepath: Path, actor_username: str = 'system') -> int:
    if not filepath.exists():
        return 0
    wb = openpyxl.load_workbook(filepath, data_only=True)
    ws = wb[wb.sheetnames[0]]
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return 0
    headers = [str(x).replace("\n", " ").strip() if x is not None else '' for x in rows[0]]
    added = 0
    with engine.begin() as conn:
        for row in rows[1:]:
            if not any(v not in (None, '') for v in row):
                continue
            rec = {headers[i]: ('' if i >= len(row) or row[i] is None else str(row[i]).strip()) for i in range(len(headers))}
            code = rec.get('CUSTOMER  CODE', '') or rec.get('CUSTOMER CODE', '')
            company_ar = rec.get('COMPANY  NAME (ARABIC)', '')
            company_en = rec.get('COMPANY  NAM (ENG)', '') or company_ar or code
            country = rec.get('COUNTRY', '')
            address = rec.get('ADDRESS', '')
            bank_name = rec.get('NAME OF BANK', '')
            bank_address = rec.get('ADDRESS OF  BANK', '')
            swift_code = rec.get('SWIFT CODE', '')
            postal_code = rec.get('POSTAL CODE', '')
            iban = rec.get('IBAN & ACC  NO', '')
            ice = rec.get('ICE', '')
            if not (company_en or company_ar or code):
                continue
            exists = conn.execute(text("SELECT id FROM current_clients WHERE customer_code=:c OR company_en=:e LIMIT 1"), {'c': code, 'e': company_en}).fetchone()
            if exists:
                conn.execute(text("UPDATE current_clients SET country=:country, company_ar=:company_ar, address=:address, ice=:ice, bank_name=:bank_name, bank_address=:bank_address, swift_code=:swift_code, postal_code=:postal_code, iban_account=:iban_account, source_file=:source_file WHERE id=:id"), {
                    'id': exists.id, 'country': country, 'company_ar': company_ar, 'address': address, 'ice': ice,
                    'bank_name': bank_name, 'bank_address': bank_address, 'swift_code': swift_code,
                    'postal_code': postal_code, 'iban_account': iban, 'source_file': filepath.name
                })
                continue
            insert_and_get_id(conn, """
                INSERT INTO current_clients (customer_code,country,company_ar,company_en,address,ice,bank_name,bank_address,swift_code,postal_code,iban_account,source_file,created_at,created_by)
                VALUES (:customer_code,:country,:company_ar,:company_en,:address,:ice,:bank_name,:bank_address,:swift_code,:postal_code,:iban_account,:source_file,:created_at,:created_by)
            """, {
                'customer_code': code, 'country': country, 'company_ar': company_ar, 'company_en': company_en,
                'address': address, 'ice': ice, 'bank_name': bank_name, 'bank_address': bank_address,
                'swift_code': swift_code, 'postal_code': postal_code, 'iban_account': iban, 'source_file': filepath.name,
                'created_at': time.time(), 'created_by': actor_username
            })
            added += 1
    return added


def maybe_seed_current_clients():
    try:
        with engine.begin() as conn:
            count = conn.execute(text("SELECT COUNT(*) FROM current_clients")).scalar() or 0
        if count == 0 and current_clients_seed_file().exists():
            sync_current_clients_from_excel(current_clients_seed_file())
    except Exception:
        pass




def get_table_columns(conn, table_name: str, dialect: str) -> set[str]:
    if dialect == "sqlite":
        return {row[1] for row in conn.exec_driver_sql(f"PRAGMA table_info({table_name})").fetchall()}
    if dialect == "postgresql":
        return {
            row[0]
            for row in conn.execute(
                text("""
                    SELECT column_name
                    FROM information_schema.columns
                    WHERE table_schema = current_schema()
                      AND table_name = :table_name
                """),
                {"table_name": table_name},
            ).fetchall()
        }
    current_db = conn.execute(text("SELECT DATABASE()")).scalar()
    return {
        row[0]
        for row in conn.execute(
            text("SELECT COLUMN_NAME FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_SCHEMA=:db AND TABLE_NAME=:table_name"),
            {"db": current_db, "table_name": table_name},
        ).fetchall()
    }


def ensure_column(conn, table_name: str, column_name: str, sqlite_sql: str, mysql_sql: str, dialect: str):
    cols = get_table_columns(conn, table_name, dialect)
    if column_name not in cols:
        conn.execute(text(sqlite_sql) if dialect in {"sqlite", "postgresql"} else text(mysql_sql))


def is_active_user(conn, username: str) -> bool:
    username = (username or '').strip()
    if not username:
        return False
    row = conn.execute(text("SELECT id FROM users WHERE username=:u AND is_active=1"), {"u": username}).fetchone()
    return bool(row)


def resolve_campaign_sender_username(conn, preferred_username: str, fallback_username: str) -> str:
    preferred_username = (preferred_username or '').strip()
    if preferred_username and is_active_user(conn, preferred_username):
        return preferred_username
    return (fallback_username or '').strip()

# ----------------------------
# Database bootstrap
# ----------------------------
def init_db():
    dialect = engine.dialect.name
    sqlite_mode = dialect == "sqlite"
    postgres_mode = dialect == "postgresql"

    def pg(stmt: str) -> str:
        return (
            stmt.replace("INTEGER PRIMARY KEY AUTOINCREMENT", "SERIAL PRIMARY KEY")
                .replace(" REAL", " DOUBLE PRECISION")
        )

    sqlite_ddl = [
        """
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE,
            password TEXT,
            display_name TEXT DEFAULT '',
            role TEXT DEFAULT 'user',
            is_active INTEGER DEFAULT 1
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS leads (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            company TEXT,
            contact_person TEXT DEFAULT '',
            phone TEXT DEFAULT '',
            email TEXT,
            country TEXT,
            city TEXT DEFAULT '',
            source TEXT DEFAULT '',
            assigned_to TEXT DEFAULT '',
            tags TEXT DEFAULT '',
            type TEXT,
            status TEXT,
            notes TEXT DEFAULT '',
            stage TEXT DEFAULT 'Lead',
            estimated_value REAL DEFAULT 0,
            score INTEGER DEFAULT 0,
            next_followup_at REAL DEFAULT 0,
            last_activity REAL DEFAULT 0,
            created_by TEXT DEFAULT ''
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS notifications (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            message TEXT,
            kind TEXT DEFAULT 'general',
            related_type TEXT DEFAULT '',
            related_id INTEGER DEFAULT 0,
            created_at REAL,
            is_read INTEGER DEFAULT 0,
            target_username TEXT DEFAULT '',
            actor_username TEXT DEFAULT ''
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS chat_messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT,
            message TEXT,
            created_at REAL
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS templates (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT UNIQUE,
            subject TEXT,
            body TEXT,
            created_at REAL
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS campaigns (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT,
            sender_mode TEXT DEFAULT 'outlook',
            template_id INTEGER DEFAULT NULL,
            subject TEXT,
            body TEXT,
            status TEXT DEFAULT 'draft',
            created_by TEXT DEFAULT '',
            created_at REAL
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS campaign_jobs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            campaign_id INTEGER,
            lead_id INTEGER,
            email TEXT,
            status TEXT DEFAULT 'pending',
            error_message TEXT DEFAULT '',
            created_at REAL,
            sent_at REAL DEFAULT 0
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS tracking_events (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            lead_id INTEGER DEFAULT 0,
            email TEXT DEFAULT '',
            campaign_id INTEGER DEFAULT 0,
            event_type TEXT,
            details TEXT DEFAULT '',
            created_at REAL
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS announcements (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT,
            body TEXT,
            author TEXT,
            created_at REAL
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS announcement_replies (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            announcement_id INTEGER,
            author TEXT,
            body TEXT,
            created_at REAL
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS lead_notes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            lead_id INTEGER,
            author TEXT,
            body TEXT,
            created_at REAL
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS activity_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT,
            action TEXT,
            entity_type TEXT DEFAULT '',
            entity_id INTEGER DEFAULT 0,
            details TEXT DEFAULT '',
            created_at REAL
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS tasks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT,
            description TEXT DEFAULT '',
            assigned_to TEXT DEFAULT '',
            status TEXT DEFAULT 'Open',
            priority TEXT DEFAULT 'Medium',
            due_at REAL DEFAULT 0,
            created_by TEXT DEFAULT '',
            created_at REAL,
            completed_at REAL DEFAULT 0
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS task_comments (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            task_id INTEGER,
            author TEXT,
            body TEXT,
            created_at REAL
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS lead_attachments (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            lead_id INTEGER,
            filename TEXT,
            original_name TEXT,
            created_at REAL
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS chat_channels (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT UNIQUE,
            description TEXT DEFAULT '',
            created_by TEXT DEFAULT '',
            created_at REAL,
            is_active INTEGER DEFAULT 1
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS chat_channel_members (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            channel_id INTEGER NOT NULL,
            username TEXT NOT NULL,
            role TEXT DEFAULT 'member',
            can_write INTEGER DEFAULT 1,
            joined_at REAL DEFAULT 0
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS current_clients (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            customer_code TEXT DEFAULT '',
            country TEXT DEFAULT '',
            company_ar TEXT DEFAULT '',
            company_en TEXT DEFAULT '',
            address TEXT DEFAULT '',
            ice TEXT DEFAULT '',
            bank_name TEXT DEFAULT '',
            bank_address TEXT DEFAULT '',
            swift_code TEXT DEFAULT '',
            postal_code TEXT DEFAULT '',
            iban_account TEXT DEFAULT '',
            source_file TEXT DEFAULT '',
            created_at REAL,
            created_by TEXT DEFAULT ''
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS client_invoices (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            client_id INTEGER,
            invoice_no TEXT DEFAULT '',
            invoice_date TEXT DEFAULT '',
            amount REAL DEFAULT 0,
            currency TEXT DEFAULT 'USD',
            notes TEXT DEFAULT '',
            attachment_filename TEXT DEFAULT '',
            attachment_original_name TEXT DEFAULT '',
            created_at REAL,
            created_by TEXT DEFAULT ''
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS trade_reference_items (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            category TEXT DEFAULT 'agreement',
            item_type TEXT DEFAULT '',
            region TEXT DEFAULT '',
            partner_countries TEXT DEFAULT '',
            effective_year TEXT DEFAULT '',
            status TEXT DEFAULT '',
            summary TEXT DEFAULT '',
            benefits TEXT DEFAULT '',
            rules_of_origin TEXT DEFAULT '',
            source_org TEXT DEFAULT '',
            source_url TEXT DEFAULT '',
            tags TEXT DEFAULT '',
            sort_order INTEGER DEFAULT 0,
            created_at REAL DEFAULT 0,
            updated_at REAL DEFAULT 0
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS bridge_agents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT NOT NULL,
            device_name TEXT NOT NULL,
            device_token TEXT NOT NULL UNIQUE,
            outlook_account_email TEXT DEFAULT '',
            is_active INTEGER DEFAULT 1,
            last_seen_at REAL DEFAULT 0,
            created_at REAL DEFAULT 0
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS campaign_job_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            job_id INTEGER NOT NULL,
            status TEXT NOT NULL,
            message TEXT DEFAULT '',
            created_at REAL DEFAULT 0
        )
        """,
    ]

    mysql_ddl = [
        """
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTO_INCREMENT,
            username VARCHAR(255) UNIQUE,
            password TEXT,
            display_name VARCHAR(255) DEFAULT '',
            role VARCHAR(100) DEFAULT 'user',
            is_active INTEGER DEFAULT 1
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS leads (
            id INTEGER PRIMARY KEY AUTO_INCREMENT,
            company TEXT,
            contact_person VARCHAR(255) DEFAULT '',
            phone VARCHAR(255) DEFAULT '',
            email VARCHAR(255),
            country VARCHAR(255),
            city VARCHAR(255) DEFAULT '',
            source VARCHAR(255) DEFAULT '',
            assigned_to VARCHAR(255) DEFAULT '',
            tags TEXT,
            type VARCHAR(255),
            status VARCHAR(255),
            notes LONGTEXT,
            stage VARCHAR(255) DEFAULT 'Lead',
            estimated_value DOUBLE DEFAULT 0,
            score INTEGER DEFAULT 0,
            next_followup_at DOUBLE DEFAULT 0,
            last_activity DOUBLE DEFAULT 0,
            created_by VARCHAR(255) DEFAULT ''
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS notifications (
            id INTEGER PRIMARY KEY AUTO_INCREMENT,
            message TEXT,
            kind VARCHAR(100) DEFAULT 'general',
            related_type VARCHAR(100) DEFAULT '',
            related_id INTEGER DEFAULT 0,
            created_at DOUBLE,
            is_read INTEGER DEFAULT 0,
            target_username VARCHAR(255) DEFAULT '',
            actor_username VARCHAR(255) DEFAULT ''
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS chat_messages (
            id INTEGER PRIMARY KEY AUTO_INCREMENT,
            username VARCHAR(255),
            message LONGTEXT,
            created_at DOUBLE
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS templates (
            id INTEGER PRIMARY KEY AUTO_INCREMENT,
            name VARCHAR(255) UNIQUE,
            subject TEXT,
            body LONGTEXT,
            created_at DOUBLE
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS campaigns (
            id INTEGER PRIMARY KEY AUTO_INCREMENT,
            name VARCHAR(255),
            sender_mode VARCHAR(50) DEFAULT 'outlook',
            template_id INTEGER NULL,
            subject TEXT,
            body LONGTEXT,
            status VARCHAR(50) DEFAULT 'draft',
            created_by VARCHAR(255) DEFAULT '',
            created_at DOUBLE
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS campaign_jobs (
            id INTEGER PRIMARY KEY AUTO_INCREMENT,
            campaign_id INTEGER,
            lead_id INTEGER,
            email VARCHAR(255),
            status VARCHAR(50) DEFAULT 'pending',
            error_message TEXT,
            created_at DOUBLE,
            sent_at DOUBLE DEFAULT 0
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS tracking_events (
            id INTEGER PRIMARY KEY AUTO_INCREMENT,
            lead_id INTEGER DEFAULT 0,
            email VARCHAR(255) DEFAULT '',
            campaign_id INTEGER DEFAULT 0,
            event_type VARCHAR(100),
            details TEXT,
            created_at DOUBLE
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS announcements (
            id INTEGER PRIMARY KEY AUTO_INCREMENT,
            title VARCHAR(255),
            body LONGTEXT,
            author VARCHAR(255),
            created_at DOUBLE
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS announcement_replies (
            id INTEGER PRIMARY KEY AUTO_INCREMENT,
            announcement_id INTEGER,
            author VARCHAR(255),
            body LONGTEXT,
            created_at DOUBLE
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS lead_notes (
            id INTEGER PRIMARY KEY AUTO_INCREMENT,
            lead_id INTEGER,
            author VARCHAR(255),
            body LONGTEXT,
            created_at DOUBLE
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS activity_logs (
            id INTEGER PRIMARY KEY AUTO_INCREMENT,
            username VARCHAR(255),
            action VARCHAR(255),
            entity_type VARCHAR(100) DEFAULT '',
            entity_id INTEGER DEFAULT 0,
            details TEXT,
            created_at DOUBLE
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS tasks (
            id INTEGER PRIMARY KEY AUTO_INCREMENT,
            title VARCHAR(255),
            description LONGTEXT,
            assigned_to VARCHAR(255) DEFAULT '',
            status VARCHAR(50) DEFAULT 'Open',
            priority VARCHAR(50) DEFAULT 'Medium',
            due_at DOUBLE DEFAULT 0,
            created_by VARCHAR(255) DEFAULT '',
            created_at DOUBLE,
            completed_at DOUBLE DEFAULT 0
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS task_comments (
            id INTEGER PRIMARY KEY AUTO_INCREMENT,
            task_id INTEGER,
            author VARCHAR(255),
            body LONGTEXT,
            created_at DOUBLE
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS lead_attachments (
            id INTEGER PRIMARY KEY AUTO_INCREMENT,
            lead_id INTEGER,
            filename VARCHAR(255),
            original_name VARCHAR(255),
            created_at DOUBLE
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS chat_channels (
            id INTEGER PRIMARY KEY AUTO_INCREMENT,
            name VARCHAR(255) UNIQUE,
            description TEXT,
            created_by VARCHAR(255) DEFAULT '',
            created_at DOUBLE,
            is_active INTEGER DEFAULT 1
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS chat_channel_members (
            id INTEGER PRIMARY KEY AUTO_INCREMENT,
            channel_id INTEGER NOT NULL,
            username VARCHAR(255) NOT NULL,
            role VARCHAR(50) DEFAULT 'member',
            can_write INTEGER DEFAULT 1,
            joined_at DOUBLE DEFAULT 0,
            UNIQUE KEY uniq_room_member (channel_id, username)
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS current_clients (
            id INTEGER PRIMARY KEY AUTO_INCREMENT,
            customer_code VARCHAR(255) DEFAULT '',
            country VARCHAR(255) DEFAULT '',
            company_ar TEXT,
            company_en TEXT,
            address TEXT,
            ice VARCHAR(255) DEFAULT '',
            bank_name TEXT,
            bank_address TEXT,
            swift_code VARCHAR(255) DEFAULT '',
            postal_code VARCHAR(255) DEFAULT '',
            iban_account TEXT,
            source_file VARCHAR(255) DEFAULT '',
            created_at DOUBLE,
            created_by VARCHAR(255) DEFAULT ''
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS client_invoices (
            id INTEGER PRIMARY KEY AUTO_INCREMENT,
            client_id INTEGER,
            invoice_no VARCHAR(255) DEFAULT '',
            invoice_date VARCHAR(50) DEFAULT '',
            amount DOUBLE DEFAULT 0,
            currency VARCHAR(20) DEFAULT 'USD',
            notes TEXT,
            attachment_filename VARCHAR(255) DEFAULT '',
            attachment_original_name VARCHAR(255) DEFAULT '',
            created_at DOUBLE,
            created_by VARCHAR(255) DEFAULT ''
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS trade_reference_items (
            id INTEGER PRIMARY KEY AUTO_INCREMENT,
            title VARCHAR(500) NOT NULL,
            category VARCHAR(50) DEFAULT 'agreement',
            item_type VARCHAR(100) DEFAULT '',
            region VARCHAR(255) DEFAULT '',
            partner_countries TEXT,
            effective_year VARCHAR(50) DEFAULT '',
            status VARCHAR(100) DEFAULT '',
            summary LONGTEXT,
            benefits LONGTEXT,
            rules_of_origin LONGTEXT,
            source_org VARCHAR(255) DEFAULT '',
            source_url TEXT,
            tags TEXT,
            sort_order INTEGER DEFAULT 0,
            created_at DOUBLE DEFAULT 0,
            updated_at DOUBLE DEFAULT 0
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS bridge_agents (
            id INTEGER PRIMARY KEY AUTO_INCREMENT,
            username VARCHAR(255) NOT NULL,
            device_name VARCHAR(255) NOT NULL,
            device_token VARCHAR(255) NOT NULL UNIQUE,
            outlook_account_email VARCHAR(255) DEFAULT '',
            is_active INTEGER DEFAULT 1,
            last_seen_at DOUBLE DEFAULT 0,
            created_at DOUBLE DEFAULT 0
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
        """
        CREATE TABLE IF NOT EXISTS campaign_job_logs (
            id INTEGER PRIMARY KEY AUTO_INCREMENT,
            job_id INTEGER NOT NULL,
            status VARCHAR(50) NOT NULL,
            message TEXT,
            created_at DOUBLE DEFAULT 0
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """,
    ]

    ddl_statements = sqlite_ddl if sqlite_mode else [pg(stmt) for stmt in sqlite_ddl] if postgres_mode else mysql_ddl

    with engine.begin() as conn:
        for stmt in ddl_statements:
            conn.execute(text(stmt))

        if sqlite_mode:
            conn.execute(text("CREATE UNIQUE INDEX IF NOT EXISTS idx_chat_channel_members_unique ON chat_channel_members(channel_id, username)"))
        elif postgres_mode:
            conn.execute(text("CREATE UNIQUE INDEX IF NOT EXISTS idx_chat_channel_members_unique ON chat_channel_members(channel_id, username)"))

        lead_cols = get_table_columns(conn, "leads", dialect)
        notif_cols = get_table_columns(conn, "notifications", dialect)
        if 'created_by' not in lead_cols:
            conn.execute(text("ALTER TABLE leads ADD COLUMN created_by TEXT DEFAULT ''") if dialect in {"sqlite", "postgresql"} else text("ALTER TABLE leads ADD COLUMN created_by VARCHAR(255) DEFAULT ''"))
        for col, sqlite_sql, mysql_sql in [
            ('contact_person', "ALTER TABLE leads ADD COLUMN contact_person TEXT DEFAULT ''", "ALTER TABLE leads ADD COLUMN contact_person VARCHAR(255) DEFAULT ''"),
            ('phone', "ALTER TABLE leads ADD COLUMN phone TEXT DEFAULT ''", "ALTER TABLE leads ADD COLUMN phone VARCHAR(255) DEFAULT ''"),
            ('city', "ALTER TABLE leads ADD COLUMN city TEXT DEFAULT ''", "ALTER TABLE leads ADD COLUMN city VARCHAR(255) DEFAULT ''"),
            ('source', "ALTER TABLE leads ADD COLUMN source TEXT DEFAULT ''", "ALTER TABLE leads ADD COLUMN source VARCHAR(255) DEFAULT ''"),
            ('assigned_to', "ALTER TABLE leads ADD COLUMN assigned_to TEXT DEFAULT ''", "ALTER TABLE leads ADD COLUMN assigned_to VARCHAR(255) DEFAULT ''"),
            ('tags', "ALTER TABLE leads ADD COLUMN tags TEXT DEFAULT ''", "ALTER TABLE leads ADD COLUMN tags TEXT"),
            ('score', "ALTER TABLE leads ADD COLUMN score INTEGER DEFAULT 0", "ALTER TABLE leads ADD COLUMN score INTEGER DEFAULT 0"),
        ]:
            if col not in lead_cols:
                conn.execute(text(sqlite_sql) if dialect in {"sqlite", "postgresql"} else text(mysql_sql))
        if 'target_username' not in notif_cols:
            conn.execute(text("ALTER TABLE notifications ADD COLUMN target_username TEXT DEFAULT ''") if dialect in {"sqlite", "postgresql"} else text("ALTER TABLE notifications ADD COLUMN target_username VARCHAR(255) DEFAULT ''"))
        if 'actor_username' not in notif_cols:
            conn.execute(text("ALTER TABLE notifications ADD COLUMN actor_username TEXT DEFAULT ''") if dialect in {"sqlite", "postgresql"} else text("ALTER TABLE notifications ADD COLUMN actor_username VARCHAR(255) DEFAULT ''"))

        user_cols = get_table_columns(conn, "users", dialect)
        for col, sqlite_sql, mysql_sql in [
            ('job_title', "ALTER TABLE users ADD COLUMN job_title TEXT DEFAULT ''", "ALTER TABLE users ADD COLUMN job_title VARCHAR(255) DEFAULT ''"),
            ('avatar_path', "ALTER TABLE users ADD COLUMN avatar_path TEXT DEFAULT ''", "ALTER TABLE users ADD COLUMN avatar_path VARCHAR(255) DEFAULT ''"),
            ('password_change_requested', "ALTER TABLE users ADD COLUMN password_change_requested INTEGER DEFAULT 0", "ALTER TABLE users ADD COLUMN password_change_requested INTEGER DEFAULT 0"),
            ('requested_password_hash', "ALTER TABLE users ADD COLUMN requested_password_hash TEXT DEFAULT ''", "ALTER TABLE users ADD COLUMN requested_password_hash TEXT"),
        ]:
            if col not in user_cols:
                conn.execute(text(sqlite_sql) if dialect in {"sqlite", "postgresql"} else text(mysql_sql))

        chat_cols = get_table_columns(conn, "chat_messages", dialect)
        for col, sqlite_sql, mysql_sql in [
            ('channel_id', "ALTER TABLE chat_messages ADD COLUMN channel_id INTEGER DEFAULT 1", "ALTER TABLE chat_messages ADD COLUMN channel_id INTEGER DEFAULT 1"),
            ('reply_to_id', "ALTER TABLE chat_messages ADD COLUMN reply_to_id INTEGER DEFAULT 0", "ALTER TABLE chat_messages ADD COLUMN reply_to_id INTEGER DEFAULT 0"),
            ('file_path', "ALTER TABLE chat_messages ADD COLUMN file_path TEXT DEFAULT ''", "ALTER TABLE chat_messages ADD COLUMN file_path TEXT"),
            ('file_name', "ALTER TABLE chat_messages ADD COLUMN file_name TEXT DEFAULT ''", "ALTER TABLE chat_messages ADD COLUMN file_name VARCHAR(255) DEFAULT ''"),
            ('file_type', "ALTER TABLE chat_messages ADD COLUMN file_type TEXT DEFAULT ''", "ALTER TABLE chat_messages ADD COLUMN file_type VARCHAR(50) DEFAULT ''"),
            ('edited_at', "ALTER TABLE chat_messages ADD COLUMN edited_at REAL DEFAULT 0", "ALTER TABLE chat_messages ADD COLUMN edited_at DOUBLE DEFAULT 0"),
            ('pinned', "ALTER TABLE chat_messages ADD COLUMN pinned INTEGER DEFAULT 0", "ALTER TABLE chat_messages ADD COLUMN pinned INTEGER DEFAULT 0"),
        ]:
            if col not in chat_cols:
                conn.execute(text(sqlite_sql) if dialect in {"sqlite", "postgresql"} else text(mysql_sql))

        channel_cols = get_table_columns(conn, "chat_channels", dialect)
        for col, sqlite_sql, mysql_sql in [
            ('password_hash', "ALTER TABLE chat_channels ADD COLUMN password_hash TEXT DEFAULT ''", "ALTER TABLE chat_channels ADD COLUMN password_hash TEXT"),
            ('only_members', "ALTER TABLE chat_channels ADD COLUMN only_members INTEGER DEFAULT 0", "ALTER TABLE chat_channels ADD COLUMN only_members INTEGER DEFAULT 0"),
            ('room_type', "ALTER TABLE chat_channels ADD COLUMN room_type TEXT DEFAULT 'channel'", "ALTER TABLE chat_channels ADD COLUMN room_type VARCHAR(50) DEFAULT 'channel'"),
        ]:
            if col not in channel_cols:
                conn.execute(text(sqlite_sql) if dialect in {"sqlite", "postgresql"} else text(mysql_sql))
        if not conn.execute(text("SELECT id FROM chat_channels WHERE id=1")).fetchone():
            conn.execute(text("INSERT INTO chat_channels (id,name,description,created_by,created_at,is_active) VALUES (1,'General','Default team channel','system',:t,1)"), {'t': time.time()})

        campaign_job_cols = get_table_columns(conn, "campaign_jobs", dialect)
        for col, sqlite_sql, mysql_sql in [
            ('assigned_username', "ALTER TABLE campaign_jobs ADD COLUMN assigned_username TEXT DEFAULT ''", "ALTER TABLE campaign_jobs ADD COLUMN assigned_username VARCHAR(255) DEFAULT ''"),
            ('assigned_device_token', "ALTER TABLE campaign_jobs ADD COLUMN assigned_device_token TEXT DEFAULT ''", "ALTER TABLE campaign_jobs ADD COLUMN assigned_device_token VARCHAR(255) DEFAULT ''"),
            ('subject_snapshot', "ALTER TABLE campaign_jobs ADD COLUMN subject_snapshot TEXT DEFAULT ''", "ALTER TABLE campaign_jobs ADD COLUMN subject_snapshot TEXT"),
            ('body_snapshot', "ALTER TABLE campaign_jobs ADD COLUMN body_snapshot TEXT DEFAULT ''", "ALTER TABLE campaign_jobs ADD COLUMN body_snapshot LONGTEXT"),
            ('from_email', "ALTER TABLE campaign_jobs ADD COLUMN from_email TEXT DEFAULT ''", "ALTER TABLE campaign_jobs ADD COLUMN from_email VARCHAR(255) DEFAULT ''"),
            ('provider_message_id', "ALTER TABLE campaign_jobs ADD COLUMN provider_message_id TEXT DEFAULT ''", "ALTER TABLE campaign_jobs ADD COLUMN provider_message_id VARCHAR(255) DEFAULT ''"),
            ('last_attempt_at', "ALTER TABLE campaign_jobs ADD COLUMN last_attempt_at REAL DEFAULT 0", "ALTER TABLE campaign_jobs ADD COLUMN last_attempt_at DOUBLE DEFAULT 0"),
        ]:
            if col not in campaign_job_cols:
                conn.execute(text(sqlite_sql) if dialect in {"sqlite", "postgresql"} else text(mysql_sql))

        admin = conn.execute(text("SELECT id, username, password FROM users WHERE role='admin' ORDER BY id LIMIT 1")).fetchone()
        if not admin:
            if os.getenv("RENDER") and not INITIAL_ADMIN_PASSWORD:
                raise RuntimeError("INITIAL_ADMIN_PASSWORD is required when bootstrapping a new Render database.")
            bootstrap_password = resolve_bootstrap_admin_password()
            conn.execute(text("""
            INSERT INTO users (username,password,display_name,role,is_active)
            VALUES (:username,:password,'Administrator','admin',1)
            """), {"username": INITIAL_ADMIN_USERNAME, "password": hash_password(bootstrap_password)})
        elif admin.password == 'Admin@123':
            if os.getenv("RENDER") and not INITIAL_ADMIN_PASSWORD:
                raise RuntimeError("Rotate the default admin password before running on Render.")
            bootstrap_password = resolve_bootstrap_admin_password()
            conn.execute(text("UPDATE users SET password=:password, username=:username WHERE id=:id"), {"password": hash_password(bootstrap_password), "username": INITIAL_ADMIN_USERNAME, "id": admin.id})

        default_tpl = conn.execute(text("SELECT id FROM templates WHERE name='Intro Offer'" )).fetchone()
        if not default_tpl:
            conn.execute(text("""
            INSERT INTO templates (name,subject,body,created_at)
            VALUES (:n,:s,:b,:t)
            """), {
                "n": "Intro Offer",
                "s": "Altahhan Dates Offer for {{company}}",
                "b": """Dear {{company}}

We are pleased to share our offer for premium dates from Egypt.

Country: {{country}}
Type: {{type}}

Please let us know if you would like our catalog and quotation.

Best regards,
Altahhan Dates""",
                "t": time.time(),
            })


init_db()
maybe_seed_current_clients()
maybe_seed_trade_intelligence()


# ----------------------------
# Auth and shared helpers
# ----------------------------
def get_user_by_username(username: str):
    with engine.begin() as conn:
        return conn.execute(text("SELECT * FROM users WHERE username=:u"), {"u": username}).fetchone()


def current_user(request: Request):
    username = request.cookies.get("username")
    if not username:
        return None
    user = get_user_by_username(username)
    if not user or not user.is_active:
        return None
    return user


def require_login(request: Request):
    user = current_user(request)
    if not user:
        raise HTTPException(status_code=401)
    return user


def is_mobile_request(request: Request) -> bool:
    ua = (request.headers.get("user-agent") or "").lower()
    mobile_keywords = ["iphone", "android", "mobile", "ipad", "ipod", "opera mini", "iemobile"]
    return any(k in ua for k in mobile_keywords)


def ensure_desktop_only(request: Request):
    if is_mobile_request(request):
        raise HTTPException(status_code=403, detail="Desktop only feature")



def require_admin(request: Request):
    user = require_login(request)
    if user.role != 'admin':
        raise HTTPException(status_code=403)
    return user


ROLE_LABELS = {
    'admin': 'Admin',
    'manager': 'Manager',
    'staff': 'Staff',
    'user': 'Staff',
    'viewer': 'Viewer',
}

def normalized_role(role: str) -> str:
    role = (role or 'user').strip().lower()
    if role == 'user':
        return 'staff'
    if role not in {'admin', 'manager', 'staff', 'viewer'}:
        return 'staff'
    return role

def user_can_edit(user) -> bool:
    return normalized_role(getattr(user, 'role', 'staff')) in {'admin', 'manager', 'staff'}

def user_can_manage_devices(user) -> bool:
    return normalized_role(getattr(user, 'role', 'staff')) in {'admin', 'manager', 'staff'}

def require_editor(request: Request):
    user = require_login(request)
    if not user_can_edit(user):
        raise HTTPException(status_code=403, detail='Read-only account')
    return user

def ensure_due_notifications(username: str):
    now_ts = time.time()
    day_key = time.strftime('%Y-%m-%d')
    with engine.begin() as conn:
        followups = conn.execute(text("""
            SELECT id, entity_type, entity_id, title, followup_at, assigned_to
            FROM followups
            WHERE status IN ('Open','Waiting')
              AND followup_at > 0
              AND followup_at <= :soon
              AND (COALESCE(assigned_to,'')='' OR assigned_to=:u)
            ORDER BY followup_at ASC
            LIMIT 20
        """), {'soon': now_ts + 86400, 'u': username}).fetchall()
        for row in followups:
            when = time.strftime('%Y-%m-%d', time.localtime(float(row.followup_at or 0))) if row.followup_at else day_key
            add_notification(
                f"Follow-up due: {row.title or row.entity_type + ' #' + str(row.entity_id)} ({when})",
                kind='reminder', related_type=row.entity_type, related_id=int(row.entity_id or 0),
                conn=conn, target_username=username, actor_username='', dedupe_seconds=12*3600
            )

        lead_rows = conn.execute(text("""
            SELECT id, company, contact_person, next_followup_at
            FROM leads
            WHERE next_followup_at > 0 AND next_followup_at <= :soon
            ORDER BY next_followup_at ASC
            LIMIT 20
        """), {'soon': now_ts + 86400}).fetchall()
        for row in lead_rows:
            label = (row.company or row.contact_person or f'Lead #{row.id}')
            when = time.strftime('%Y-%m-%d', time.localtime(float(row.next_followup_at or 0))) if row.next_followup_at else day_key
            add_notification(
                f"Lead follow-up due: {label} ({when})",
                kind='reminder', related_type='lead', related_id=int(row.id or 0),
                conn=conn, target_username=username, actor_username='', dedupe_seconds=12*3600
            )

def get_bridge_agents_for_user(username: str):
    with engine.begin() as conn:
        return conn.execute(text("SELECT * FROM bridge_agents WHERE username=:u ORDER BY id DESC"), {"u": username}).fetchall()


def get_bridge_agent_by_token(conn, device_token: str):
    return conn.execute(text("SELECT * FROM bridge_agents WHERE device_token=:t AND is_active=1"), {"t": (device_token or '').strip()}).fetchone()


def queue_campaign_for_outlook_bridge(conn, campaign, jobs, fallback_username: str, tracking_base: str):
    queued = 0
    notified_users = set()
    for job in jobs:
        lead = {
            "company": job.company,
            "contact_person": job.contact_person,
            "phone": job.phone,
            "email": job.email,
            "country": job.country,
            "city": job.city,
            "source": job.source,
            "assigned_to": job.assigned_to,
            "tags": job.tags,
            "type": job.type,
            "status": job.lead_status,
            "stage": job.stage,
        }
        sub, html_body = render_template_text(campaign.subject or '', campaign.body or '', lead)
        tracking_pixel = f'<img src="{tracking_base}/track/open?lead_id={job.lead_id}&campaign_id={campaign.id}&email={urllib.parse.quote(job.email or "")}" width="1" height="1">'
        final_body = (html_body or '') + '<br><br>' + tracking_pixel
        assigned_username = resolve_campaign_sender_username(conn, job.assigned_to or '', fallback_username)
        conn.execute(text("""
            UPDATE campaign_jobs
            SET status='queued',
                assigned_username=:assigned_username,
                assigned_device_token='',
                subject_snapshot=:subject_snapshot,
                body_snapshot=:body_snapshot,
                error_message='',
                provider_message_id='',
                last_attempt_at=0
            WHERE id=:id
        """), {
            "assigned_username": assigned_username,
            "subject_snapshot": sub,
            "body_snapshot": final_body,
            "id": job.id,
        })
        conn.execute(text("INSERT INTO campaign_job_logs (job_id,status,message,created_at) VALUES (:job_id,'queued',:message,:created_at)"), {
            "job_id": job.id,
            "message": f"Queued for Outlook Desktop user: {assigned_username}",
            "created_at": time.time(),
        })
        if assigned_username:
            notified_users.add(assigned_username)
        queued += 1
    conn.execute(text("UPDATE campaigns SET status=:s WHERE id=:id"), {"s": 'queued-desktop', "id": campaign.id})
    for username in notified_users:
        add_notification(
            f"Campaign #{campaign.id} has queued emails for your Outlook Desktop bridge",
            kind='campaign',
            related_type='campaign',
            related_id=campaign.id,
            conn=conn,
            target_username=username,
            actor_username=(campaign.created_by or fallback_username),
        )
    return queued



def add_notification(message: str, kind: str = 'general', related_type: str = '', related_id: int = 0, conn=None, target_username: str = '', actor_username: str = '', dedupe_seconds: int = 2):
    target_username = (target_username or '').strip()
    actor_username = (actor_username or '').strip()
    if target_username and actor_username and target_username == actor_username:
        return
    params = {
        "m": message[:500],
        "k": kind,
        "rt": related_type,
        "rid": related_id,
        "t": time.time(),
        "tu": target_username,
        "au": actor_username,
        "cutoff": time.time() - max(dedupe_seconds, 0),
    }
    dedupe_sql = text("""
        SELECT id FROM notifications
        WHERE message=:m AND kind=:k AND related_type=:rt AND related_id=:rid
          AND COALESCE(target_username,'')=:tu
          AND created_at >= :cutoff
        ORDER BY id DESC LIMIT 1
    """)
    insert_sql = text("""
        INSERT INTO notifications (message,kind,related_type,related_id,created_at,is_read,target_username,actor_username)
        VALUES (:m,:k,:rt,:rid,:t,0,:tu,:au)
    """)
    if conn is not None:
        if conn.execute(dedupe_sql, params).fetchone():
            return
        conn.execute(insert_sql, params)
        return
    with engine.begin() as conn2:
        if conn2.execute(dedupe_sql, params).fetchone():
            return
        conn2.execute(insert_sql, params)


def notify_many(usernames, message: str, kind: str = 'general', related_type: str = '', related_id: int = 0, conn=None, actor_username: str = ''):
    seen = set()
    for username in usernames or []:
        username = (username or '').strip()
        if not username or username in seen:
            continue
        seen.add(username)
        add_notification(message, kind=kind, related_type=related_type, related_id=related_id, conn=conn, target_username=username, actor_username=actor_username)



def active_usernames(conn):
    return [r.username for r in conn.execute(text("SELECT username FROM users WHERE is_active=1 ORDER BY username ASC")).fetchall()]

def task_assignee_targets(conn, assigned_to: str, created_by: str = '', exclude=None):
    exclude = exclude or set()
    assigned_to = (assigned_to or '').strip()
    targets = set()
    if assigned_to == '__all__':
        targets.update(active_usernames(conn))
    elif assigned_to:
        targets.add(assigned_to)
    if created_by:
        targets.add(created_by)
    return {u for u in targets if u and u not in exclude}

def notification_where_sql():
    # Bell/popup is intentionally limited to chat, task, and announcement updates.
    return "kind IN ('chat','task','announcement') AND (COALESCE(target_username,'')='' OR target_username=:username)"


def latest_notification_id_for_user(username: str):
    with engine.begin() as conn:
        return conn.execute(text(f"SELECT COALESCE(MAX(id),0) FROM notifications WHERE {notification_where_sql()}"), {"username": username}).scalar() or 0


def notification_target_url(row) -> str:
    related_type = (getattr(row, 'related_type', '') or '').strip().lower()
    kind = (getattr(row, 'kind', '') or '').strip().lower()
    related_id = safe_int(getattr(row, 'related_id', 0), 0)
    if related_type == 'task' and related_id:
        return f"/task/{related_id}"
    if related_type == 'announcement' and related_id:
        return f"/announcement/{related_id}"
    if related_type == 'lead' and related_id:
        return f"/lead/{related_id}"
    if related_type == 'current_client' and related_id:
        return f"/current-client/{related_id}"
    if related_type == 'campaign' and related_id:
        return f"/campaign/{related_id}"
    if related_type == 'user':
        return "/users"
    if kind == 'chat':
        return "/chat"
    if kind == 'current_client':
        return "/current-clients"
    if kind == 'announcement':
        return "/announcements"
    if kind == 'task':
        return "/tasks"
    if kind in {'lead', 'pipeline', 'followup', 'note'}:
        return "/leads"
    return "/dashboard"


def log_activity(username: str, action: str, entity_type: str = '', entity_id: int = 0, details: str = '', conn=None):
    params = {"u": username, "a": action, "et": entity_type, "eid": entity_id, "d": details[:1000], "t": time.time()}
    sql = text("""
        INSERT INTO activity_logs (username,action,entity_type,entity_id,details,created_at)
        VALUES (:u,:a,:et,:eid,:d,:t)
    """)
    if conn is not None:
        conn.execute(sql, params)
        return
    with engine.begin() as conn2:
        conn2.execute(sql, params)


def render_template_text(subject: str, body: str, lead: dict):
    values = {
        "company": lead.get("company", ""),
        "contact_person": lead.get("contact_person", ""),
        "phone": lead.get("phone", ""),
        "email": lead.get("email", ""),
        "country": lead.get("country", ""),
        "city": lead.get("city", ""),
        "source": lead.get("source", ""),
        "assigned_to": lead.get("assigned_to", ""),
        "tags": lead.get("tags", ""),
        "type": lead.get("type", ""),
        "status": lead.get("status", ""),
        "stage": lead.get("stage", ""),
    }
    for k, v in values.items():
        subject = subject.replace("{{" + k + "}}", str(v or ""))
        body = body.replace("{{" + k + "}}", str(v or ""))
    return subject, body.replace("\n", "<br>")


class BridgeRegisterPayload(BaseModel):
    device_name: str = ''
    outlook_account_email: str = ''


class BridgeHeartbeatPayload(BaseModel):
    device_token: str


class BridgeJobResultPayload(BaseModel):
    device_token: str
    job_id: int
    status: str
    message: str = ''
    provider_message_id: str = ''


def send_via_outlook(to_email: str, subject: str, html_body: str):
    try:
        import win32com.client  # type: ignore
    except Exception as exc:
        raise RuntimeError("pywin32 is not installed. Install pywin32 on the Windows server.") from exc
    outlook = win32com.client.Dispatch("Outlook.Application")
    mail = outlook.CreateItem(0)
    mail.To = to_email
    mail.Subject = subject
    mail.HTMLBody = html_body
    mail.Send()


# -----------------------------------------------------------------------------
# Email sending helpers
#
# The original implementation relied on the Windows‑only pywin32 library to
# dispatch messages through a locally installed Outlook client.  This works
# well on a developer's machine but makes the application impossible to deploy
# on Linux servers or in containers.  To support cross‑platform deployments and
# allow sending email through any SMTP provider (including Microsoft 365),
# additional helper functions are defined below.

def send_via_smtp(to_email: str, subject: str, html_body: str) -> None:
    """
    Send an email using standard SMTP.  Connection parameters are taken from
    environment variables:

      SMTP_HOST:     hostname of the SMTP server (e.g. smtp.office365.com)
      SMTP_PORT:     port number (default 587)
      SMTP_USERNAME: login username or from address
      SMTP_PASSWORD: password or application secret
      SMTP_USE_TLS:  "true"/"false" whether to use STARTTLS (default true)

    The message is sent as HTML.  If any required setting is missing the
    function will raise a RuntimeError to be handled by the caller.
    """
    host = os.getenv("SMTP_HOST") or ""
    port = int(os.getenv("SMTP_PORT") or 587)
    username = os.getenv("SMTP_USERNAME") or ""
    password = os.getenv("SMTP_PASSWORD") or ""
    use_tls = (os.getenv("SMTP_USE_TLS") or "true").lower() in {"1", "true", "yes"}
    if not host or not username or not password:
        raise RuntimeError("SMTP settings are not configured")
    msg = MIMEMultipart("alternative")
    msg["From"] = username
    msg["To"] = to_email
    msg["Subject"] = subject
    part = MIMEText(html_body or "", "html")
    msg.attach(part)
    with smtplib.SMTP(host, port) as server:
        # Try to upgrade the connection to TLS if requested
        if use_tls:
            server.starttls()
        server.login(username, password)
        server.sendmail(username, [to_email], msg.as_string())


def send_via_graph(access_token: str, to_email: str, subject: str, html_body: str) -> None:
    access_token = (access_token or '').strip()
    if not access_token:
        raise RuntimeError('Microsoft Graph access token is not configured')
    payload = {
        'message': {
            'subject': subject,
            'body': {'contentType': 'HTML', 'content': html_body or ''},
            'toRecipients': [{'emailAddress': {'address': to_email}}],
        },
        'saveToSentItems': True,
    }
    req = urllib.request.Request(
        'https://graph.microsoft.com/v1.0/me/sendMail',
        data=json.dumps(payload).encode('utf-8'),
        headers={'Authorization': f'Bearer {access_token}', 'Content-Type': 'application/json'},
        method='POST',
    )
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            if resp.status not in (200, 202):
                raise RuntimeError(f'Graph send failed with status {resp.status}')
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode('utf-8', errors='ignore')
        raise RuntimeError(f'Graph send failed: {detail or exc.reason}') from exc



def graph_oauth_config() -> dict:
    return {
        'tenant': (os.getenv('MICROSOFT_TENANT_ID') or 'common').strip() or 'common',
        'client_id': (os.getenv('MICROSOFT_CLIENT_ID') or '').strip(),
        'client_secret': (os.getenv('MICROSOFT_CLIENT_SECRET') or '').strip(),
        'redirect_uri': (os.getenv('MICROSOFT_REDIRECT_URI') or '').strip(),
        'scopes': ['offline_access', 'openid', 'profile', 'User.Read', 'Mail.Send', 'Mail.Read'],
    }


def refresh_graph_access_token(refresh_token: str) -> dict | None:
    cfg = graph_oauth_config()
    if not (cfg['client_id'] and cfg['client_secret'] and refresh_token):
        return None
    data = urllib.parse.urlencode({
        'client_id': cfg['client_id'],
        'client_secret': cfg['client_secret'],
        'grant_type': 'refresh_token',
        'refresh_token': refresh_token,
        'redirect_uri': cfg['redirect_uri'],
        'scope': ' '.join(cfg['scopes']),
    }).encode('utf-8')
    req = urllib.request.Request(
        f"https://login.microsoftonline.com/{cfg['tenant']}/oauth2/v2.0/token",
        data=data,
        headers={'Content-Type': 'application/x-www-form-urlencoded'},
        method='POST',
    )
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            return json.loads(resp.read().decode('utf-8'))
    except Exception:
        return None


def resolve_profile_graph_token(profile) -> str:
    token = (getattr(profile, 'graph_access_token', '') or '').strip()
    if token:
        return token
    oauth_account_id = int(getattr(profile, 'oauth_account_id', 0) or 0)
    if not oauth_account_id:
        return (os.getenv('MICROSOFT_GRAPH_ACCESS_TOKEN', '') or '').strip()
    with engine.begin() as conn:
        account = conn.execute(text('SELECT * FROM oauth_accounts WHERE id=:id'), {'id': oauth_account_id}).fetchone()
        if not account:
            return (os.getenv('MICROSOFT_GRAPH_ACCESS_TOKEN', '') or '').strip()
        access_token = (getattr(account, 'access_token', '') or '').strip()
        expires_at = float(getattr(account, 'token_expires_at', 0) or 0)
        if access_token and expires_at > time.time() + 60:
            return access_token
        refresh_token = (getattr(account, 'refresh_token', '') or '').strip()
        refreshed = refresh_graph_access_token(refresh_token)
        if refreshed and refreshed.get('access_token'):
            new_access = refreshed.get('access_token', '')
            new_refresh = refreshed.get('refresh_token') or refresh_token
            expires_in = int(refreshed.get('expires_in', 3600) or 3600)
            conn.execute(text('UPDATE oauth_accounts SET access_token=:a, refresh_token=:r, token_expires_at=:e, updated_at=:t WHERE id=:id'), {
                'a': new_access,
                'r': new_refresh,
                'e': time.time() + expires_in,
                't': time.time(),
                'id': oauth_account_id,
            })
            return new_access
        return access_token


def send_email_with_profile(profile, to_email: str, subject: str, html_body: str) -> None:
    provider = (getattr(profile, 'provider', '') or 'smtp').strip().lower()
    if provider == 'microsoft_graph':
        token = resolve_profile_graph_token(profile)
        if token:
            return send_via_graph(token, to_email, subject, html_body)
        raise RuntimeError('Graph profile needs an access token or connected Microsoft account')
    if provider == 'smtp':
        host = (getattr(profile, 'smtp_host', '') or '').strip() or (os.getenv('SMTP_HOST') or '').strip()
        port = int(getattr(profile, 'smtp_port', 0) or os.getenv('SMTP_PORT') or 587)
        username = (getattr(profile, 'smtp_username', '') or '').strip() or (os.getenv('SMTP_USERNAME') or '').strip()
        password = (getattr(profile, 'smtp_password', '') or '') or (os.getenv('SMTP_PASSWORD') or '')
        use_tls = int(getattr(profile, 'smtp_use_tls', 1) or 0) == 1
        from_email = (getattr(profile, 'from_email', '') or '').strip() or username
        if not host or not username or not password:
            raise RuntimeError('SMTP profile is incomplete: host, username and password are required')
        msg = MIMEMultipart('alternative')
        msg['From'] = from_email
        msg['To'] = to_email
        msg['Subject'] = subject
        reply_to = (getattr(profile, 'reply_to', '') or '').strip()
        if reply_to:
            msg['Reply-To'] = reply_to
        msg.attach(MIMEText(html_body or '', 'html'))
        with smtplib.SMTP(host, port) as server:
            if use_tls:
                server.starttls()
            server.login(username, password)
            server.sendmail(from_email, [to_email], msg.as_string())
        return
    # outlook_desktop and any unknown provider fall back to current behavior
    send_email(to_email, subject, html_body, 'outlook' if provider == 'outlook_desktop' else 'smtp')


def generate_agent_content(company: str, contact_name: str, country: str, goal: str) -> dict:
    company = (company or '').strip()
    contact_name = (contact_name or '').strip()
    country = (country or '').strip()
    goal = (goal or 'Start a conversation for export sales').strip()
    api_key = os.getenv('OPENAI_API_KEY', '').strip()
    prompt = (
        'You are a B2B export sales assistant for a dates manufacturer. '
        'Write one short email subject and one HTML email body. '
        'Keep it practical, professional, and easy to send. '
        f'Company: {company or "Unknown"}. Contact: {contact_name or "Team"}. Country: {country or "Unknown"}. Goal: {goal}. '
        'Return JSON with keys subject and body.'
    )
    if api_key:
        try:
            payload = {
                'model': os.getenv('OPENAI_MODEL', 'gpt-4.1-mini'),
                'input': prompt,
            }
            req = urllib.request.Request(
                'https://api.openai.com/v1/responses',
                data=json.dumps(payload).encode('utf-8'),
                headers={'Authorization': f'Bearer {api_key}', 'Content-Type': 'application/json'},
                method='POST',
            )
            with urllib.request.urlopen(req, timeout=45) as resp:
                data = json.loads(resp.read().decode('utf-8'))
            output_text = ''
            for item in data.get('output', []):
                for content in item.get('content', []):
                    if content.get('type') in {'output_text', 'text'}:
                        output_text += content.get('text', '')
            parsed = json.loads(output_text.strip()) if output_text.strip().startswith('{') else None
            if isinstance(parsed, dict) and parsed.get('subject') and parsed.get('body'):
                return {'subject': parsed['subject'], 'body': parsed['body']}
        except Exception:
            pass
    intro = f"Hello {contact_name or 'team'},"
    subject = f"Export opportunity for {company or 'your company'}"
    body = (
        f"<p>{intro}</p>"
        f"<p>We would like to introduce Al Tahhan Golden Dates and explore a possible cooperation with {company or 'your company'}"
        f"{' in ' + country if country else ''}.</p>"
        f"<p>{goal}.</p>"
        "<p>We can share product specs, packing options, and price ideas based on your market needs.</p>"
        "<p>Best regards,<br>Sales Team</p>"
    )
    return {'subject': subject, 'body': body}


def send_email(to_email: str, subject: str, html_body: str, mode: str = "outlook") -> None:
    """
    Unified email sending entry point.  The caller can specify a preferred
    mode ("outlook" or "smtp").  If the specified mode fails, this function
    attempts to fall back gracefully to SMTP.

    :param to_email: recipient address
    :param subject: message subject
    :param html_body: HTML body of the message
    :param mode: "outlook" (default) or "smtp"
    """
    mode = (mode or "outlook").lower()
    if mode == "outlook":
        try:
            send_via_outlook(to_email, subject, html_body)
            return
        except Exception:
            pass
    send_via_smtp(to_email, subject, html_body)


# ----------------------------
# Views / API
# ----------------------------
@app.get("/healthz")
def healthz():
    return {"ok": True, "app": APP_TITLE, "lan_ip": get_preferred_lan_ip()}


@app.get("/", response_class=HTMLResponse)
def landing_page(request: Request):
    return templates.TemplateResponse("landing.html", {"request": request})


@app.get("/intro")
def intro_redirect():
    return redirect("/")


@app.get("/tour", response_class=HTMLResponse)
def product_tour_page(request: Request):
    return templates.TemplateResponse("tour.html", {"request": request})


@app.get("/outreach")
def outreach_alias():
    return redirect("/email-tool")


@app.get("/login", response_class=HTMLResponse)
def login_page(request: Request):
    return templates.TemplateResponse("login.html", {"request": request, "error": request.query_params.get("error", "")})


@app.get("/workspace-search", response_class=HTMLResponse)
def workspace_search(request: Request, q: str = ""):
    user = require_login(request)
    q_clean = (q or '').strip()
    results = {
        'leads': [],
        'clients': [],
        'shipments': [],
        'deals': [],
        'tasks': [],
        'followups': [],
        'agreements': [],
        'campaigns': [],
        'export_products': [],
        'export_markets': [],
    }
    counters = {k: 0 for k in results}
    if q_clean:
        like = f"%{q_clean}%"
        with engine.begin() as conn:
            results['leads'] = conn.execute(text("""
                SELECT id, company, contact_person, email, country, stage, score
                FROM leads
                WHERE company LIKE :q OR contact_person LIKE :q OR email LIKE :q OR country LIKE :q OR notes LIKE :q OR tags LIKE :q
                ORDER BY score DESC, id DESC LIMIT 12
            """), {'q': like}).fetchall()
            results['clients'] = conn.execute(text("""
                SELECT id, customer_code, company_en, company_ar, country
                FROM current_clients
                WHERE customer_code LIKE :q OR company_en LIKE :q OR company_ar LIKE :q OR country LIKE :q OR bank_name LIKE :q
                ORDER BY id DESC LIMIT 12
            """), {'q': like}).fetchall()
            results['shipments'] = conn.execute(text("""
                SELECT id, shipment_no, company, supplier, product_name, current_status, destination_port
                FROM shipments
                WHERE shipment_no LIKE :q OR company LIKE :q OR supplier LIKE :q OR product_name LIKE :q OR destination_port LIKE :q OR origin_port LIKE :q OR notes LIKE :q
                ORDER BY id DESC LIMIT 12
            """), {'q': like}).fetchall()
            results['deals'] = conn.execute(text("""
                SELECT d.id, d.invoice_no, d.client_name, d.product_name, d.status, s.shipment_no
                FROM shipment_deals d
                LEFT JOIN shipments s ON s.id=d.shipment_id
                WHERE d.invoice_no LIKE :q OR d.client_name LIKE :q OR d.product_name LIKE :q OR s.shipment_no LIKE :q
                ORDER BY d.id DESC LIMIT 12
            """), {'q': like}).fetchall()
            results['tasks'] = conn.execute(text("""
                SELECT id, title, assigned_to, status, priority, due_at
                FROM tasks
                WHERE title LIKE :q OR description LIKE :q OR assigned_to LIKE :q OR created_by LIKE :q
                ORDER BY CASE WHEN status='Done' THEN 1 ELSE 0 END, id DESC LIMIT 12
            """), {'q': like}).fetchall()
            results['followups'] = conn.execute(text("""
                SELECT id, entity_type, entity_id, title, assigned_to, status, priority, followup_at
                FROM followups
                WHERE title LIKE :q OR notes LIKE :q OR assigned_to LIKE :q OR entity_type LIKE :q
                ORDER BY CASE WHEN status='Done' THEN 1 ELSE 0 END, id DESC LIMIT 12
            """), {'q': like}).fetchall()
            results['agreements'] = conn.execute(text("""
                SELECT id, title, category, region, partner_countries, status
                FROM trade_reference_items
                WHERE title LIKE :q OR region LIKE :q OR partner_countries LIKE :q OR tags LIKE :q OR summary LIKE :q
                ORDER BY sort_order ASC, id DESC LIMIT 12
            """), {'q': like}).fetchall()
            results['campaigns'] = conn.execute(text("""
                SELECT id, name, status, sender_mode, created_by, created_at FROM (
                    SELECT id, name, status, sender_mode, created_by, created_at
                    FROM campaigns
                    WHERE name LIKE :q OR subject LIKE :q OR body LIKE :q OR created_by LIKE :q
                    UNION ALL
                    SELECT id, campaign_name AS name, status, 'campaign_v2' AS sender_mode, created_by, created_at
                    FROM campaigns_v2
                    WHERE campaign_name LIKE :q OR notes LIKE :q OR created_by LIKE :q
                    UNION ALL
                    SELECT id, sequence_name AS name, status, 'sequence_v2' AS sender_mode, created_by, created_at
                    FROM campaign_sequences_v2
                    WHERE sequence_name LIKE :q OR notes LIKE :q OR created_by LIKE :q
                ) all_campaigns
                ORDER BY id DESC LIMIT 12
            """), {'q': like}).fetchall()
            try:
                results['export_products'] = conn.execute(text("""
                    SELECT id, product_name, category, origin_country, packaging, target_markets
                    FROM export_products
                    WHERE product_name LIKE :q OR category LIKE :q OR origin_country LIKE :q OR packaging LIKE :q OR notes LIKE :q OR target_markets LIKE :q
                    ORDER BY sort_order ASC, id DESC LIMIT 12
                """), {'q': like}).fetchall()
                results['export_markets'] = conn.execute(text("""
                    SELECT id, country_name, region_name, demand_level, price_position, preferred_products
                    FROM export_markets
                    WHERE country_name LIKE :q OR region_name LIKE :q OR demand_level LIKE :q OR price_position LIKE :q OR preferred_products LIKE :q OR market_notes LIKE :q
                    ORDER BY sort_order ASC, id DESC LIMIT 12
                """), {'q': like}).fetchall()
            except Exception:
                results['export_products'] = []
                results['export_markets'] = []
        counters = {k: len(v) for k, v in results.items()}
    return templates.TemplateResponse('search_results.html', {
        'request': request,
        'username': user.username,
        'user': user,
        'query': q_clean,
        'results': results,
        'counters': counters,
        'total_found': sum(counters.values()),
        'latest_notification_id': latest_notification_id_for_user(user.username)
    })


@app.post("/login")
def login(request: Request, username: str = Form(...), password: str = Form(...)):
    with engine.begin() as conn:
        user = conn.execute(text("SELECT * FROM users WHERE username=:u AND is_active=1"), {"u": username.strip()}).fetchone()
        if not user or not verify_password(password, user.password):
            return templates.TemplateResponse("login.html", {"request": request, "error": "Invalid login"})
        if not str(user.password).startswith("pbkdf2_sha256$"):
            conn.execute(text("UPDATE users SET password=:p WHERE id=:id"), {"p": hash_password(password), "id": user.id})

    resp = redirect("/dashboard")
    resp.set_cookie("username", user.username, httponly=True, samesite="lax", secure=COOKIE_SECURE, max_age=60*60*12)
    return resp


@app.get("/logout")
def logout():
    resp = redirect("/login")
    resp.delete_cookie("username", secure=COOKIE_SECURE, samesite="lax")
    return resp


@app.get("/server-info")
def server_info(request: Request):
    user = require_login(request)
    lan_ip = get_preferred_lan_ip()
    app_url = build_tracking_base_url(request)
    return {
        "username": user.username,
        "role": user.role,
        "lan_ip": lan_ip,
        "app_url": app_url,
        "tracking_base_url": app_url,
        "image_base_url": build_image_base_url(request),
        "db_url": "configured" if ACTIVE_DB_URL else "",
    }


@app.get("/settings/database", response_class=HTMLResponse)
def database_settings_page(request: Request):
    user = require_admin(request)
    settings = load_runtime_settings()
    message = request.query_params.get("message", "")
    error = request.query_params.get("error", "")
    return templates.TemplateResponse("settings_database.html", {
        "request": request,
        "user": user,
        "username": user.username,
        "title": "Database Settings",
        "settings": settings,
        "active_db_url": ACTIVE_DB_URL,
        "message": message,
        "error": error,
        "latest_notification_id": latest_notification_id_for_user(user.username),
    })


@app.post("/settings/database")
def save_database_settings(request: Request, db_driver: str = Form("sqlite"), db_server: str = Form(""), db_name: str = Form(""), db_user: str = Form(""), db_password: str = Form("")):
    user = require_admin(request)
    payload = {
        "db_driver": (db_driver or "sqlite").strip().lower(),
        "db_server": db_server.strip(),
        "db_name": db_name.strip(),
        "db_user": db_user.strip(),
        "db_password": db_password,
    }
    try:
        test_url = build_db_url_from_settings({**load_runtime_settings(), **payload})
        test_engine = make_engine(test_url)
        with test_engine.begin() as conn:
            conn.execute(text("SELECT 1"))
        try:
            test_engine.dispose()
        except Exception:
            pass
        save_runtime_settings(payload)
        reload_engine_from_settings()
        log_activity(user.username, "save_db_settings", "settings", 0, payload.get("db_driver", "sqlite"))
        return redirect("/settings/database?message=Database+settings+saved")
    except Exception as exc:
        return redirect("/settings/database?error=" + urllib.parse.quote(str(exc)[:180]))


@app.get("/settings/site", response_class=HTMLResponse)
def site_settings_page(request: Request):
    user = require_admin(request)
    settings = load_runtime_settings()
    image_base = build_image_base_url(request)
    files = []
    try:
        files = sorted([f.name for f in UPLOADS_DIR.iterdir() if f.is_file()], reverse=True)[:20]
    except Exception:
        pass
    return templates.TemplateResponse("settings_site.html", {
        "request": request,
        "user": user,
        "username": user.username,
        "title": "Site & Domain Settings",
        "settings": settings,
        "domain_base": build_tracking_base_url(request),
        "image_base": image_base,
        "files": files,
        "message": request.query_params.get("message", ""),
        "error": request.query_params.get("error", ""),
        "latest_notification_id": latest_notification_id_for_user(user.username),
    })


@app.post("/settings/site")
def save_site_settings(request: Request, domain: str = Form(""), image_prefix: str = Form("/uploads")):
    user = require_admin(request)
    payload = {
        "domain": domain.strip(),
        "image_prefix": normalize_prefix(image_prefix),
    }
    save_runtime_settings(payload)
    reload_engine_from_settings()
    log_activity(user.username, "save_site_settings", "settings", 0, payload.get("domain", ""))
    return redirect("/settings/site?message=Site+settings+saved")


@app.post("/settings/site/upload-image")
async def upload_site_image(request: Request, image: UploadFile = File(...)):
    user = require_admin(request)
    filename = Path(image.filename or "upload.bin").name
    ext = Path(filename).suffix.lower()
    if ext not in {".png", ".jpg", ".jpeg", ".webp", ".gif", ".svg"}:
        return redirect("/settings/site?error=Only+image+files+are+allowed")
    safe_name = f"{int(time.time())}_{secrets.token_hex(4)}{ext}"
    target = UPLOADS_DIR / safe_name
    data = await image.read()
    target.write_bytes(data)
    log_activity(user.username, "upload_image", "settings", 0, safe_name)
    return redirect("/settings/site?message=" + urllib.parse.quote(public_upload_url(safe_name, request)))


@app.get("/dashboard", response_class=HTMLResponse)
def dashboard(request: Request):
    user = require_login(request)
    ensure_due_notifications(user.username)
    now = time.time()
    is_admin = user.role == 'admin'
    lead_filter = '' if is_admin else ' AND created_by=:username'
    activity_sql = 'SELECT * FROM activity_logs ORDER BY id DESC LIMIT 12' if is_admin else 'SELECT * FROM activity_logs WHERE username=:username ORDER BY id DESC LIMIT 12'
    params = {"t": now, "username": user.username}
    with engine.begin() as conn:
        stats = {
            "leads": conn.execute(text(f"SELECT COUNT(*) FROM leads WHERE 1=1{lead_filter}"), params).scalar() or 0,
            "notifications": conn.execute(text(f"SELECT COUNT(*) FROM notifications WHERE is_read=0 AND {notification_where_sql()}"), {"username": user.username}).scalar() or 0,
            "followup_due": conn.execute(text(f"SELECT COUNT(*) FROM leads WHERE next_followup_at > 0 AND next_followup_at <= :t{lead_filter}"), params).scalar() or 0,
            "won": conn.execute(text(f"SELECT COUNT(*) FROM leads WHERE stage='Won'{lead_filter}"), params).scalar() or 0,
            "lost": conn.execute(text(f"SELECT COUNT(*) FROM leads WHERE stage='Lost'{lead_filter}"), params).scalar() or 0,
            "pipeline_value": conn.execute(text(f"SELECT COALESCE(SUM(estimated_value),0) FROM leads WHERE stage NOT IN ('Lost','Won'){lead_filter}"), params).scalar() or 0,
            "tasks_open": conn.execute(text("SELECT COUNT(*) FROM tasks WHERE status != 'Done' AND (assigned_to=:username OR created_by=:username)"), params).scalar() or 0,
            "my_tasks_open": conn.execute(text("SELECT COUNT(*) FROM tasks WHERE status != 'Done' AND (assigned_to=:username OR assigned_to='__all__')"), params).scalar() or 0,
            "overdue_tasks": conn.execute(text("SELECT COUNT(*) FROM tasks WHERE status != 'Done' AND due_at > 0 AND due_at < :t AND (assigned_to=:username OR created_by=:username)"), params).scalar() or 0,
            "avg_score": conn.execute(text(f"SELECT COALESCE(AVG(score),0) FROM leads WHERE 1=1{lead_filter}"), params).scalar() or 0,
            "shipments_open": conn.execute(text("SELECT COUNT(*) FROM shipments WHERE current_status NOT IN ('Delivered','Closed')"), {}).scalar() or 0,
            "shipments_delayed": conn.execute(text("SELECT COUNT(*) FROM shipments WHERE eta_at > 0 AND eta_at < :t AND current_status NOT IN ('Delivered','Closed')"), params).scalar() or 0,
            "arriving_soon": conn.execute(text("SELECT COUNT(*) FROM shipments WHERE eta_at > :t AND eta_at <= :soon AND current_status NOT IN ('Delivered','Closed')"), {"t": now, "soon": now + 7*86400}).scalar() or 0,
            "missing_docs": conn.execute(text("SELECT COUNT(*) FROM shipments s WHERE current_status NOT IN ('Delivered','Closed') AND (SELECT COUNT(*) FROM shipment_documents d WHERE d.shipment_id=s.id AND d.doc_type IN ('invoice','packing_list')) < 2"), {}).scalar() or 0,
            "followups_total_due": conn.execute(text("SELECT COUNT(*) FROM followups WHERE status IN ('Open','Waiting') AND followup_at > 0 AND followup_at <= :t"), params).scalar() or 0,
        }
        stats["conversion_rate"] = round((stats["won"] / stats["leads"] * 100), 1) if stats["leads"] else 0
        stage_summary = conn.execute(text(f"SELECT stage, COUNT(*) AS cnt, COALESCE(SUM(estimated_value),0) AS total FROM leads WHERE 1=1{lead_filter} GROUP BY stage ORDER BY cnt DESC, stage ASC"), params).fetchall()
        hot = conn.execute(text(f"""
        SELECT
            l.id,
            l.company,
            l.stage,
            COALESCE((SELECT COUNT(*)*5 FROM tracking_events te WHERE te.lead_id=l.id AND te.event_type='open'),0) +
            COALESCE((SELECT COUNT(*)*20 FROM tracking_events te WHERE te.lead_id=l.id AND te.event_type='click'),0) +
            CASE WHEN l.stage='Negotiation' THEN 15 WHEN l.stage='Quotation' THEN 25 WHEN l.stage='Won' THEN 40 ELSE 0 END AS score
        FROM leads l WHERE 1=1{lead_filter} ORDER BY score DESC, l.id DESC LIMIT 5
        """), params).fetchall()
        recent_notifications = conn.execute(text(f"SELECT * FROM notifications WHERE {notification_where_sql()} ORDER BY id DESC LIMIT 8"), {"username": user.username}).fetchall()
        recent_activity = conn.execute(text(activity_sql), {"username": user.username}).fetchall()
        upcoming_tasks = conn.execute(text("SELECT * FROM tasks WHERE assigned_to=:username OR created_by=:username ORDER BY CASE WHEN status='Done' THEN 1 ELSE 0 END, CASE WHEN due_at > 0 THEN due_at ELSE 32503680000 END ASC, id DESC LIMIT 8"), params).fetchall()
        my_tasks = conn.execute(text("SELECT * FROM tasks WHERE (assigned_to=:username OR assigned_to='__all__') ORDER BY CASE WHEN status='Done' THEN 1 ELSE 0 END, CASE WHEN due_at > 0 THEN due_at ELSE 32503680000 END ASC, id DESC LIMIT 6"), params).fetchall()
        shipment_rows = conn.execute(text("SELECT * FROM shipments ORDER BY CASE WHEN current_status IN ('Delivered','Closed') THEN 1 ELSE 0 END, CASE WHEN eta_at > 0 THEN eta_at ELSE 32503680000 END ASC, id DESC LIMIT 6"), {}).fetchall()
        overdue_followups = conn.execute(text("SELECT * FROM followups WHERE status IN ('Open','Waiting') AND followup_at > 0 AND followup_at <= :t ORDER BY followup_at ASC LIMIT 8"), params).fetchall()
    return templates.TemplateResponse("dashboard.html", {"request": request, "username": user.username, "user": user, "stats": stats, "hot": hot, "recent_notifications": recent_notifications, "recent_activity": recent_activity, "stage_summary": stage_summary, "upcoming_tasks": upcoming_tasks, "my_tasks": my_tasks, "shipment_rows": shipment_rows, "overdue_followups": overdue_followups, "latest_notification_id": latest_notification_id_for_user(user.username)})


@app.get("/leads", response_class=HTMLResponse)
def leads_page(request: Request, q: str = '', stage: str = '', assigned_to: str = '', source: str = ''):
    user = require_login(request)
    params = {"q": f"%{q.strip()}%", "stage": stage.strip(), "assigned_to": assigned_to.strip(), "source": source.strip()}
    sql = "SELECT * FROM leads WHERE 1=1"
    if q.strip():
        sql += " AND (company LIKE :q OR contact_person LIKE :q OR email LIKE :q OR phone LIKE :q OR country LIKE :q OR city LIKE :q OR type LIKE :q OR source LIKE :q OR status LIKE :q OR tags LIKE :q)"
    if stage.strip():
        sql += " AND stage=:stage"
    if assigned_to.strip():
        sql += " AND assigned_to=:assigned_to"
    if source.strip():
        sql += " AND source=:source"
    sql += " ORDER BY id DESC"
    with engine.begin() as conn:
        leads = conn.execute(text(sql), params).fetchall()
        users = conn.execute(text("SELECT username, display_name FROM users WHERE is_active=1 ORDER BY username ASC")).fetchall()
        summary = {
            "all": conn.execute(text("SELECT COUNT(*) FROM leads")).scalar() or 0,
            "hot": conn.execute(text("SELECT COUNT(*) FROM leads WHERE score >= 60")).scalar() or 0,
            "won": conn.execute(text("SELECT COUNT(*) FROM leads WHERE stage='Won'" )).scalar() or 0,
            "value": conn.execute(text("SELECT COALESCE(SUM(estimated_value),0) FROM leads WHERE stage NOT IN ('Lost','Won')")).scalar() or 0,
        }
        sources = conn.execute(text("SELECT DISTINCT source FROM leads WHERE COALESCE(source,'') != '' ORDER BY source ASC")).fetchall()
    return templates.TemplateResponse("leads.html", {"request": request, "username": user.username, "user": user, "leads": leads, "q": q, "stage_filter": stage, "assigned_filter": assigned_to, "source_filter": source, "users": users, "summary": summary, "sources": sources})


@app.post("/leads/add")
def add_lead(request: Request, company: str = Form(...), contact_person: str = Form(''), phone: str = Form(''), email: str = Form(''), country: str = Form(''), city: str = Form(''), source: str = Form(''), assigned_to: str = Form(''), tags: str = Form(''), type: str = Form(''), status: str = Form('New'), stage: str = Form('Lead'), estimated_value: float = Form(0), notes: str = Form('')):
    user = require_login(request)
    payload = {
        "company": company.strip(),
        "contact_person": contact_person.strip(),
        "phone": phone.strip(),
        "email": email.strip(),
        "country": country.strip(),
        "city": city.strip(),
        "source": source.strip(),
        "assigned_to": assigned_to.strip(),
        "tags": tags.strip(),
        "type": type.strip(),
        "status": status.strip(),
        "stage": normalize_stage(stage),
        "estimated_value": estimated_value,
        "notes": notes.strip(),
    }
    payload["score"] = calculate_lead_score(payload)
    with engine.begin() as conn:
        lead_id = insert_and_get_id(conn, """
        INSERT INTO leads (company,contact_person,phone,email,country,city,source,assigned_to,tags,type,status,notes,stage,estimated_value,score,last_activity,created_by)
        VALUES (:company,:contact_person,:phone,:email,:country,:city,:source,:assigned_to,:tags,:type,:status,:notes,:stage,:estimated_value,:score,:t,:u)
        """, {**payload, "t": time.time(), "u": user.username})
        log_activity(user.username, "add_lead", "lead", lead_id, company.strip(), conn=conn)
    add_notification(f"{user.username} added lead {company}", kind='lead', related_type='lead', related_id=lead_id)
    return redirect("/leads")


@app.post("/leads/{lead_id}/stage")
def update_stage(request: Request, lead_id: int, stage: str = Form(...)):
    user = require_login(request)
    with engine.begin() as conn:
        conn.execute(text("UPDATE leads SET stage=:s, last_activity=:t WHERE id=:id"), {"s": stage, "t": time.time(), "id": lead_id})
        log_activity(user.username, "move_stage", "lead", lead_id, stage, conn=conn)
    add_notification(f"{user.username} moved lead #{lead_id} to {stage}", kind='pipeline', related_type='lead', related_id=lead_id)
    return redirect("/pipeline")


@app.post("/leads/{lead_id}/followup")
def set_followup(request: Request, lead_id: int, days: int = Form(...)):
    user = require_login(request)
    ts = time.time() + max(days, 0) * 86400
    with engine.begin() as conn:
        conn.execute(text("UPDATE leads SET next_followup_at=:ts WHERE id=:id"), {"ts": ts, "id": lead_id})
        log_activity(user.username, "set_followup", "lead", lead_id, str(days), conn=conn)
    add_notification(f"{user.username} set follow-up for lead #{lead_id}", kind='followup', related_type='lead', related_id=lead_id)
    return redirect("/leads")


@app.post("/leads/{lead_id}/note")
def add_lead_note(request: Request, lead_id: int, body: str = Form(...)):
    user = require_login(request)
    with engine.begin() as conn:
        conn.execute(text("INSERT INTO lead_notes (lead_id,author,body,created_at) VALUES (:lid,:a,:b,:t)"), {"lid": lead_id, "a": user.username, "b": body.strip(), "t": time.time()})
        log_activity(user.username, "add_note", "lead", lead_id, body.strip()[:120], conn=conn)
    add_notification(f"{user.username} added a note to lead #{lead_id}", kind='note', related_type='lead', related_id=lead_id)
    return redirect(f"/lead/{lead_id}")


@app.post("/leads/{lead_id}/upload")
async def upload_lead_attachment(request: Request, lead_id: int, file: UploadFile = File(...)):
    user = require_login(request)
    original_name = Path(file.filename or '').name
    if not original_name:
        return redirect(f"/lead/{lead_id}?error=invalid_file")
    suffix = Path(original_name).suffix.lower()
    safe_name = f"lead_{lead_id}_{int(time.time())}_{secrets.token_hex(4)}{suffix}"
    target = UPLOADS_DIR / safe_name
    target.write_bytes(await file.read())
    with engine.begin() as conn:
        conn.execute(text("INSERT INTO lead_attachments (lead_id, filename, original_name, created_at) VALUES (:lid,:f,:o,:t)"), {"lid": lead_id, "f": safe_name, "o": original_name, "t": time.time()})
        log_activity(user.username, "upload_attachment", "lead", lead_id, original_name, conn=conn)
    add_notification(f"{user.username} uploaded an attachment to lead #{lead_id}", kind='lead', related_type='lead', related_id=lead_id)
    return redirect(f"/lead/{lead_id}")


@app.get("/lead/{lead_id}", response_class=HTMLResponse)
def lead_detail(request: Request, lead_id: int):
    user = require_login(request)
    with engine.begin() as conn:
        lead = conn.execute(text("SELECT * FROM leads WHERE id=:id"), {"id": lead_id}).fetchone()
        if not lead:
            return redirect("/leads")
        notes = conn.execute(text("SELECT * FROM lead_notes WHERE lead_id=:id ORDER BY id DESC"), {"id": lead_id}).fetchall()
        attachments = conn.execute(text("SELECT * FROM lead_attachments WHERE lead_id=:id ORDER BY id DESC"), {"id": lead_id}).fetchall()
        tracking = conn.execute(text("SELECT * FROM tracking_events WHERE lead_id=:id ORDER BY id DESC LIMIT 20"), {"id": lead_id}).fetchall()
        users = conn.execute(text("SELECT username, display_name FROM users WHERE is_active=1 ORDER BY username ASC")).fetchall()
    return templates.TemplateResponse("lead_detail.html", {"request": request, "username": user.username, "user": user, "lead": lead, "notes": notes, "attachments": attachments, "tracking": tracking, "users": users})


@app.post("/lead/{lead_id}/update")
def lead_update(request: Request, lead_id: int, company: str = Form(...), contact_person: str = Form(''), phone: str = Form(''), email: str = Form(''), country: str = Form(''), city: str = Form(''), source: str = Form(''), assigned_to: str = Form(''), tags: str = Form(''), type: str = Form(''), status: str = Form('New'), stage: str = Form('Lead'), estimated_value: float = Form(0), notes: str = Form('')):
    user = require_login(request)
    payload = {
        "company": company.strip(),
        "contact_person": contact_person.strip(),
        "phone": phone.strip(),
        "email": email.strip(),
        "country": country.strip(),
        "city": city.strip(),
        "source": source.strip(),
        "assigned_to": assigned_to.strip(),
        "tags": tags.strip(),
        "type": type.strip(),
        "status": status.strip(),
        "stage": normalize_stage(stage),
        "estimated_value": estimated_value,
        "notes": notes.strip(),
    }
    payload["score"] = calculate_lead_score(payload)
    with engine.begin() as conn:
        existing = conn.execute(text("SELECT id FROM leads WHERE id=:id"), {"id": lead_id}).fetchone()
        if not existing:
            return redirect("/leads?error=lead_not_found")
        conn.execute(text("UPDATE leads SET company=:company, contact_person=:contact_person, phone=:phone, email=:email, country=:country, city=:city, source=:source, assigned_to=:assigned_to, tags=:tags, type=:type, status=:status, notes=:notes, stage=:stage, estimated_value=:estimated_value, score=:score, last_activity=:t WHERE id=:id"), {**payload, "t": time.time(), "id": lead_id})
        log_activity(user.username, "update_lead", "lead", lead_id, company.strip(), conn=conn)
    add_notification(f"{user.username} updated lead #{lead_id}", kind='lead', related_type='lead', related_id=lead_id)
    return redirect(f"/lead/{lead_id}?success=updated")


@app.post("/lead/{lead_id}/delete")
def lead_delete(request: Request, lead_id: int):
    user = require_admin(request)
    with engine.begin() as conn:
        lead = conn.execute(text("SELECT company FROM leads WHERE id=:id"), {"id": lead_id}).fetchone()
        if not lead:
            return redirect("/leads?error=lead_not_found")
        conn.execute(text("DELETE FROM lead_notes WHERE lead_id=:id"), {"id": lead_id})
        conn.execute(text("DELETE FROM campaign_jobs WHERE lead_id=:id"), {"id": lead_id})
        conn.execute(text("DELETE FROM tracking_events WHERE lead_id=:id"), {"id": lead_id})
        conn.execute(text("DELETE FROM leads WHERE id=:id"), {"id": lead_id})
        log_activity(user.username, "delete_lead", "lead", lead_id, lead.company or '', conn=conn)
    add_notification(f"{user.username} deleted lead #{lead_id}", kind='lead')
    return redirect("/leads?success=lead_deleted")


@app.post("/import_excel")
async def import_excel(request: Request, file: UploadFile = File(...)):
    user = require_login(request)
    wb = openpyxl.load_workbook(file.file)
    ws = wb.active
    headers = [str(c.value).strip() if c.value is not None else '' for c in ws[1]]
    idx = {h.lower(): i for i, h in enumerate(headers)}

    def cell(row, name):
        if name.lower() not in idx:
            return ''
        v = row[idx[name.lower()]].value
        return '' if v is None else str(v).strip()

    added = 0
    with engine.begin() as conn:
        for row in ws.iter_rows(min_row=2):
            company = cell(row, 'Company')
            email = cell(row, 'Email')
            if not company and not email:
                continue
            payload = {
                "company": company,
                "contact_person": cell(row, 'Contact Person'),
                "phone": cell(row, 'Phone'),
                "email": email,
                "country": cell(row, 'Country'),
                "city": cell(row, 'City'),
                "source": cell(row, 'Source'),
                "assigned_to": cell(row, 'Assigned To'),
                "tags": cell(row, 'Tags'),
                "type": cell(row, 'Type'),
                "status": cell(row, 'Status') or 'New',
                "notes": cell(row, 'Notes'),
                "stage": normalize_stage(cell(row, 'Stage') or 'Lead'),
                "estimated_value": float(cell(row, 'Estimated Value') or 0),
            }
            payload['score'] = calculate_lead_score(payload)
            lead_id = insert_and_get_id(conn, """
            INSERT INTO leads (company,contact_person,phone,email,country,city,source,assigned_to,tags,type,status,notes,stage,estimated_value,score,last_activity,created_by)
            VALUES (:company,:contact_person,:phone,:email,:country,:city,:source,:assigned_to,:tags,:type,:status,:notes,:stage,:estimated_value,:score,:t,:u)
            """, {**payload, "t": time.time(), "u": user.username})
            log_activity(user.username, "import_lead", "lead", lead_id, company, conn=conn)
            added += 1
    add_notification(f"{user.username} imported {added} leads", kind='import')
    return redirect("/leads")



@app.get("/leads/template")
def leads_template(request: Request):
    require_login(request)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Leads Template'
    headers = ['Company','Contact Person','Email','Phone','Country','City','Source','Assigned To','Tags','Type','Status','Stage','Estimated Value','Notes']
    ws.append(headers)
    ws.append(['Example Trading Co','Ahmed Ali','example@company.com','201234567890','Egypt','Cairo','Website','admin','importer, dates','Medjool','New','Lead',25000,'Interested in 5kg cartons'])
    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return StreamingResponse(buf, media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', headers={"Content-Disposition":"attachment; filename=leads_import_template.xlsx"})

@app.get("/export_leads")
def export_leads(request: Request):
    user = require_login(request)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Leads'
    ws.append(['ID','Company','Contact Person','Phone','Email','Country','City','Source','Assigned To','Tags','Type','Status','Stage','Estimated Value','Score','Notes'])
    with engine.begin() as conn:
        rows = conn.execute(text("SELECT id,company,contact_person,phone,email,country,city,source,assigned_to,tags,type,status,stage,estimated_value,score,notes FROM leads ORDER BY id DESC")).fetchall()
    for row in rows:
        ws.append(list(row))
    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    add_notification(f"{user.username} exported leads", kind='export')
    return StreamingResponse(buf, media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', headers={"Content-Disposition": "attachment; filename=leads_export.xlsx"})


@app.get("/current-clients", response_class=HTMLResponse)
def current_clients_page(request: Request, q: str = '', country: str = ''):
    user = require_login(request)
    params = {'q': f"%{q.strip()}%", 'country': country.strip()}
    sql = "SELECT * FROM current_clients WHERE 1=1"
    if q.strip():
        sql += " AND (customer_code LIKE :q OR company_en LIKE :q OR company_ar LIKE :q OR address LIKE :q OR bank_name LIKE :q OR swift_code LIKE :q)"
    if country.strip():
        sql += " AND country=:country"
    sql += " ORDER BY company_en ASC, id DESC"
    with engine.begin() as conn:
        rows = conn.execute(text(sql), params).fetchall()
        countries = conn.execute(text("SELECT DISTINCT country FROM current_clients WHERE COALESCE(country,'')<>'' ORDER BY country ASC")).fetchall()
        total_invoices = conn.execute(text("SELECT COUNT(*) FROM client_invoices")).scalar() or 0
    return templates.TemplateResponse("current_clients.html", {
        'request': request, 'username': user.username, 'user': user, 'rows': rows,
        'q': q, 'country_filter': country, 'countries': countries, 'total_invoices': total_invoices,
        'latest_notification_id': latest_notification_id_for_user(user.username)
    })


@app.get("/current-client/{client_id}", response_class=HTMLResponse)
def current_client_detail(request: Request, client_id: int):
    user = require_login(request)
    with engine.begin() as conn:
        client = conn.execute(text("SELECT * FROM current_clients WHERE id=:id"), {'id': client_id}).fetchone()
        if not client:
            return redirect('/current-clients')
        invoices = conn.execute(text("SELECT * FROM client_invoices WHERE client_id=:id ORDER BY id DESC"), {'id': client_id}).fetchall()
        shipments = conn.execute(text("""
            SELECT * FROM shipments
            WHERE client_id=:id OR (client_id=0 AND LOWER(COALESCE(company,''))=LOWER(:company))
            ORDER BY id DESC
        """), {'id': client_id, 'company': (client.company_en or '').strip()}).fetchall()
        stats = conn.execute(text("""
            SELECT
                COUNT(*) AS shipments_count,
                COALESCE(SUM(container_count),0) AS total_containers,
                COALESCE(SUM(cartons_count),0) AS total_cartons,
                COALESCE(SUM(invoice_amount),0) AS total_invoice_amount
            FROM shipments
            WHERE client_id=:id OR (client_id=0 AND LOWER(COALESCE(company,''))=LOWER(:company))
        """), {'id': client_id, 'company': (client.company_en or '').strip()}).fetchone()
        item_mix = conn.execute(text("""
            SELECT COALESCE(item_name, product_name, 'Unknown') AS item_label,
                   COALESCE(SUM(container_count),0) AS containers,
                   COALESCE(SUM(cartons_count),0) AS cartons,
                   COALESCE(SUM(invoice_amount),0) AS amount
            FROM shipments
            WHERE client_id=:id OR (client_id=0 AND LOWER(COALESCE(company,''))=LOWER(:company))
            GROUP BY COALESCE(item_name, product_name, 'Unknown')
            ORDER BY amount DESC, cartons DESC, containers DESC
            LIMIT 12
        """), {'id': client_id, 'company': (client.company_en or '').strip()}).fetchall()
    return templates.TemplateResponse("current_client_detail.html", {
        'request': request, 'username': user.username, 'user': user, 'client': client, 'invoices': invoices, 'shipments': shipments, 'stats': stats, 'item_mix': item_mix,
        'latest_notification_id': latest_notification_id_for_user(user.username)
    })


@app.get("/api/current-client/{client_id}")
def api_current_client(request: Request, client_id: int):
    require_login(request)
    with engine.begin() as conn:
        client = conn.execute(text("SELECT * FROM current_clients WHERE id=:id"), {'id': client_id}).fetchone()
        if not client:
            raise HTTPException(status_code=404, detail="Client not found")
        data = dict(client._mapping)
    return JSONResponse(content=data)


@app.post("/current-clients/add")
def current_clients_add(request: Request, customer_code: str = Form(''), country: str = Form(''), company_ar: str = Form(''), company_en: str = Form(...), address: str = Form(''), ice: str = Form(''), bank_name: str = Form(''), bank_address: str = Form(''), swift_code: str = Form(''), postal_code: str = Form(''), iban_account: str = Form('')):
    user = require_login(request)
    with engine.begin() as conn:
        client_id = insert_and_get_id(conn, """
            INSERT INTO current_clients (customer_code,country,company_ar,company_en,address,ice,bank_name,bank_address,swift_code,postal_code,iban_account,source_file,created_at,created_by)
            VALUES (:customer_code,:country,:company_ar,:company_en,:address,:ice,:bank_name,:bank_address,:swift_code,:postal_code,:iban_account,'manual',:created_at,:created_by)
        """, {'customer_code': customer_code.strip(), 'country': country.strip(), 'company_ar': company_ar.strip(), 'company_en': company_en.strip(), 'address': address.strip(), 'ice': ice.strip(), 'bank_name': bank_name.strip(), 'bank_address': bank_address.strip(), 'swift_code': swift_code.strip(), 'postal_code': postal_code.strip(), 'iban_account': iban_account.strip(), 'created_at': time.time(), 'created_by': user.username})
        log_activity(user.username, 'add_current_client', 'current_client', client_id, company_en.strip(), conn=conn)
    add_notification(f"{user.username} added current client {company_en.strip()}", kind='current_client', related_type='current_client', related_id=client_id)
    return redirect(f"/current-client/{client_id}")


@app.post("/current-clients/import-bundled")
def current_clients_import_bundled(request: Request):
    user = require_admin(request)
    added = sync_current_clients_from_excel(current_clients_seed_file(), actor_username=user.username)
    add_notification(f"{user.username} synced current clients file ({added} new)", kind='current_client')
    return redirect('/current-clients')


@app.post("/current-client/{client_id}/invoice")
async def current_client_add_invoice(request: Request, client_id: int, invoice_no: str = Form(''), invoice_date: str = Form(''), amount: float = Form(0), currency: str = Form('USD'), notes: str = Form(''), file: UploadFile | None = File(None)):
    user = require_editor(request)
    saved_name = ''
    original_name = ''
    if file and getattr(file, 'filename', ''):
        original_name = file.filename
        safe_name = f"invoice_{client_id}_{int(time.time())}_{secrets.token_hex(4)}_{safe_filename(original_name)}"
        saved_name = safe_name
        with open(CLIENT_UPLOADS_DIR / safe_name, 'wb') as f:
            f.write(await file.read())
    with engine.begin() as conn:
        invoice_id = insert_and_get_id(conn, """
            INSERT INTO client_invoices (client_id,invoice_no,invoice_date,amount,currency,notes,attachment_filename,attachment_original_name,created_at,created_by)
            VALUES (:client_id,:invoice_no,:invoice_date,:amount,:currency,:notes,:attachment_filename,:attachment_original_name,:created_at,:created_by)
        """, {'client_id': client_id, 'invoice_no': invoice_no.strip(), 'invoice_date': invoice_date.strip(), 'amount': amount, 'currency': currency.strip() or 'USD', 'notes': notes.strip(), 'attachment_filename': saved_name, 'attachment_original_name': original_name, 'created_at': time.time(), 'created_by': user.username})
        log_activity(user.username, 'add_client_invoice', 'client_invoice', invoice_id, invoice_no.strip(), conn=conn)
    add_notification(f"{user.username} added invoice {invoice_no.strip() or '#'+str(invoice_id)}", kind='current_client', related_type='current_client', related_id=client_id)
    return redirect(f"/current-client/{client_id}")


@app.get("/agreements", response_class=HTMLResponse)
def agreements_page(request: Request, q: str = '', category: str = 'agreement', region: str = '', country: str = ''):
    user = require_login(request)
    q_clean = q.strip()
    category = (category or 'agreement').strip()
    region = (region or '').strip()
    country = (country or '').strip()
    sql = "SELECT * FROM trade_reference_items WHERE 1=1"
    params = {}
    if category in {'agreement', 'law', 'all'} and category != 'all':
        sql += " AND category=:category"
        params['category'] = category
    if region:
        sql += " AND region LIKE :region"
        params['region'] = f"%{region}%"
    if country:
        sql += " AND (partner_countries LIKE :country OR tags LIKE :country OR region LIKE :country)"
        params['country'] = f"%{country}%"
    if q_clean:
        sql += " AND (title LIKE :q OR summary LIKE :q OR benefits LIKE :q OR rules_of_origin LIKE :q OR tags LIKE :q OR partner_countries LIKE :q OR source_org LIKE :q)"
        params['q'] = f"%{q_clean}%"
    sql += " ORDER BY sort_order ASC, title ASC"
    with engine.begin() as conn:
        rows = conn.execute(text(sql), params).fetchall()
        regions = conn.execute(text("SELECT DISTINCT region FROM trade_reference_items WHERE COALESCE(region,'')<>'' ORDER BY region ASC")).fetchall()
        highlights = conn.execute(text("SELECT * FROM trade_reference_items WHERE category='agreement' ORDER BY sort_order ASC, id ASC LIMIT 8")).fetchall()
        suggested = []
        if country:
            scored = sorted(rows or highlights, key=lambda item: agreement_match_score(country, item), reverse=True)
            suggested = [x for x in scored if agreement_match_score(country, x) > 0][:5]
    return templates.TemplateResponse('agreements.html', {
        'request': request, 'username': user.username, 'user': user, 'rows': rows, 'regions': regions,
        'q': q, 'category': category, 'region_filter': region, 'country': country, 'highlights': highlights,
        'suggested': suggested, 'latest_notification_id': latest_notification_id_for_user(user.username)
    })


@app.get("/agreement/{item_id}", response_class=HTMLResponse)
def agreement_detail(request: Request, item_id: int):
    user = require_login(request)
    with engine.begin() as conn:
        item = conn.execute(text("SELECT * FROM trade_reference_items WHERE id=:id"), {'id': item_id}).fetchone()
        if not item:
            return redirect('/agreements')
        related = conn.execute(text("SELECT * FROM trade_reference_items WHERE id<>:id AND category=:category AND (region=:region OR source_org=:source_org) ORDER BY sort_order ASC, title ASC LIMIT 6"), {'id': item_id, 'category': item.category, 'region': item.region, 'source_org': item.source_org}).fetchall()
    return templates.TemplateResponse('agreement_detail.html', {
        'request': request, 'username': user.username, 'user': user, 'item': item, 'related': related,
        'latest_notification_id': latest_notification_id_for_user(user.username)
    })


@app.get("/reports", response_class=HTMLResponse)
def reports_page(request: Request, period: str = 'month'):
    user = require_login(request)
    days = 7 if period == 'week' else 30 if period == 'month' else 90 if period == 'quarter' else 365
    since = time.time() - days * 86400
    with engine.begin() as conn:
        summary = conn.execute(text("""
            SELECT
                COUNT(*) AS shipments_count,
                COALESCE(SUM(container_count),0) AS total_containers,
                COALESCE(SUM(cartons_count),0) AS total_cartons,
                COALESCE(SUM(invoice_amount),0) AS total_invoice_amount
            FROM shipments
            WHERE created_at >= :s
        """), {'s': since}).fetchone()
        customer_rows = conn.execute(text("""
            SELECT
                COALESCE(cc.company_en, s.company, 'Unknown') AS customer_name,
                COUNT(*) AS shipments_count,
                COALESCE(SUM(s.container_count),0) AS total_containers,
                COALESCE(SUM(s.cartons_count),0) AS total_cartons,
                COALESCE(SUM(s.invoice_amount),0) AS total_invoice_amount
            FROM shipments s
            LEFT JOIN current_clients cc ON cc.id=s.client_id
            WHERE s.created_at >= :s
            GROUP BY COALESCE(cc.company_en, s.company, 'Unknown')
            ORDER BY total_invoice_amount DESC, total_containers DESC
            LIMIT 20
        """), {'s': since}).fetchall()
        product_rows = conn.execute(text("""
            SELECT
                COALESCE(s.product_category,'Dates') AS product_category,
                COALESCE(NULLIF(s.item_name,''), NULLIF(s.product_name,''), 'Unknown') AS item_name,
                COUNT(*) AS shipments_count,
                COALESCE(SUM(s.container_count),0) AS total_containers,
                COALESCE(SUM(s.cartons_count),0) AS total_cartons,
                COALESCE(SUM(s.invoice_amount),0) AS total_invoice_amount
            FROM shipments s
            WHERE s.created_at >= :s
            GROUP BY COALESCE(s.product_category,'Dates'), COALESCE(NULLIF(s.item_name,''), NULLIF(s.product_name,''), 'Unknown')
            ORDER BY total_invoice_amount DESC, total_cartons DESC
            LIMIT 20
        """), {'s': since}).fetchall()
        booking_rows = conn.execute(text("""
            SELECT shipment_no, company, container_count, container_type, cartons_count, invoice_amount, currency, current_status, etd_at, eta_at
            FROM shipments
            WHERE created_at >= :s
            ORDER BY created_at DESC, id DESC
            LIMIT 50
        """), {'s': since}).fetchall()
    total_amount = float(getattr(summary, 'total_invoice_amount', 0) or 0)
    total_containers = float(getattr(summary, 'total_containers', 0) or 0)
    total_cartons = float(getattr(summary, 'total_cartons', 0) or 0)
    customers = []
    for r in customer_rows:
        m = dict(r._mapping)
        m['amount_pct'] = round((float(m.get('total_invoice_amount') or 0) / total_amount * 100), 2) if total_amount else 0
        m['containers_pct'] = round((float(m.get('total_containers') or 0) / total_containers * 100), 2) if total_containers else 0
        customers.append(m)
    products = []
    for r in product_rows:
        m = dict(r._mapping)
        m['amount_pct'] = round((float(m.get('total_invoice_amount') or 0) / total_amount * 100), 2) if total_amount else 0
        m['cartons_pct'] = round((float(m.get('total_cartons') or 0) / total_cartons * 100), 2) if total_cartons else 0
        products.append(m)
    return templates.TemplateResponse('reports.html', {
        'request': request, 'username': user.username, 'user': user, 'period': period,
        'summary': summary, 'customers': customers, 'products': products, 'bookings': booking_rows,
        'latest_notification_id': latest_notification_id_for_user(user.username)
    })


@app.get("/reports/export")
def reports_export(request: Request, period: str = 'week'):
    user = require_login(request)
    days = 7 if period == 'week' else 30
    since = time.time() - days * 86400
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Summary'
    ws.append(['Metric', f'Last {days} days'])
    with engine.begin() as conn:
        rows = [
            ['Tasks', conn.execute(text("SELECT COUNT(*) FROM tasks WHERE created_at >= :s"), {'s': since}).scalar() or 0],
            ['Leads', conn.execute(text("SELECT COUNT(*) FROM leads WHERE last_activity >= :s OR next_followup_at >= :s"), {'s': since}).scalar() or 0],
            ['Current Clients', conn.execute(text("SELECT COUNT(*) FROM current_clients WHERE created_at >= :s"), {'s': since}).scalar() or 0],
            ['Invoices', conn.execute(text("SELECT COUNT(*) FROM client_invoices WHERE created_at >= :s"), {'s': since}).scalar() or 0],
            ['Chat Messages', conn.execute(text("SELECT COUNT(*) FROM chat_messages WHERE created_at >= :s"), {'s': since}).scalar() or 0],
            ['Activity Logs', conn.execute(text("SELECT COUNT(*) FROM activity_logs WHERE created_at >= :s"), {'s': since}).scalar() or 0],
            ['Trade Agreements', conn.execute(text("SELECT COUNT(*) FROM trade_reference_items WHERE category='agreement'")).scalar() or 0],
            ['Trade Laws & Resources', conn.execute(text("SELECT COUNT(*) FROM trade_reference_items WHERE category='law'")).scalar() or 0],
        ]
        for r in rows:
            ws.append(r)
        ws2 = wb.create_sheet('Team')
        ws2.append(['Username', 'Tasks', 'Leads', 'Current Clients', 'Invoices', 'Messages', 'Actions'])
        team = conn.execute(text("""
            SELECT u.username,
                   COALESCE((SELECT COUNT(*) FROM tasks t WHERE t.created_at >= :s AND (t.assigned_to=u.username OR t.created_by=u.username)),0) AS tasks_count,
                   COALESCE((SELECT COUNT(*) FROM leads l WHERE l.created_by=u.username AND l.last_activity >= :s),0) AS leads_count,
                   COALESCE((SELECT COUNT(*) FROM current_clients c WHERE c.created_by=u.username AND c.created_at >= :s),0) AS clients_count,
                   COALESCE((SELECT COUNT(*) FROM client_invoices ci WHERE ci.created_by=u.username AND ci.created_at >= :s),0) AS invoices_count,
                   COALESCE((SELECT COUNT(*) FROM chat_messages m WHERE m.username=u.username AND m.created_at >= :s),0) AS messages_count,
                   COALESCE((SELECT COUNT(*) FROM activity_logs a WHERE a.username=u.username AND a.created_at >= :s),0) AS actions_count
            FROM users u WHERE u.is_active=1 ORDER BY u.username ASC
        """), {'s': since}).fetchall()
        for r in team:
            ws2.append([r.username, r.tasks_count, r.leads_count, r.clients_count, r.invoices_count, r.messages_count, r.actions_count])
    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return StreamingResponse(buf, media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', headers={'Content-Disposition': f'attachment; filename=altahhan_report_{period}.xlsx'})



@app.get("/users", response_class=HTMLResponse)
def users_page(request: Request):
    user = require_admin(request)
    with engine.begin() as conn:
        users = conn.execute(text("SELECT * FROM users ORDER BY id ASC")).fetchall()
    return templates.TemplateResponse("users.html", {"request": request, "username": user.username, "user": user, "users": users})


@app.post("/users/add")
def add_user(request: Request, username: str = Form(...), password: str = Form(...), role: str = Form(...), display_name: str = Form(''), job_title: str = Form('')):
    current = require_admin(request)
    username = username.strip()
    display_name = display_name.strip()
    role = role.strip().lower()
    if not username or not password:
        return redirect("/users?error=missing_fields")
    role = normalized_role(role)
    try:
        with engine.begin() as conn:
            conn.execute(text("INSERT INTO users (username,password,display_name,job_title,role,is_active) VALUES (:u,:p,:d,:j,:r,1)"), {"u": username, "p": hash_password(password), "d": display_name, "j": (job_title or '').strip(), "r": role})
            add_notification(f"{current.username} added user {username}", kind='user', conn=conn)
            log_activity(current.username, "add_user", "user", 0, username, conn=conn)
    except IntegrityError:
        return redirect("/users?error=username_exists")
    return redirect("/users?success=user_added")


@app.post("/users/{user_id}/toggle")
def toggle_user(request: Request, user_id: int):
    current = require_admin(request)
    with engine.begin() as conn:
        u = conn.execute(text("SELECT id, username, is_active FROM users WHERE id=:id"), {"id": user_id}).fetchone()
        if not u:
            return redirect("/users?error=user_not_found")
        if u.username == current.username:
            return redirect("/users?error=self_toggle")
        new_val = 0 if safe_int(u.is_active, 1) else 1
        conn.execute(text("UPDATE users SET is_active=:a WHERE id=:id"), {"a": new_val, "id": user_id})
        add_notification(f"{current.username} {'disabled' if new_val == 0 else 'enabled'} user {u.username}", kind='user', conn=conn)
        log_activity(current.username, "toggle_user", "user", user_id, f"{u.username}:{new_val}", conn=conn)
    return redirect(f"/users?success={'user_disabled' if new_val == 0 else 'user_enabled'}")


@app.post("/users/{user_id}/password")
def reset_password(request: Request, user_id: int, password: str = Form(...)):
    current = require_admin(request)
    password = password.strip()
    if len(password) < 4:
        return redirect("/users?error=password_short")
    with engine.begin() as conn:
        u = conn.execute(text("SELECT username FROM users WHERE id=:id"), {"id": user_id}).fetchone()
        if not u:
            return redirect("/users?error=user_not_found")
        conn.execute(text("UPDATE users SET password=:p WHERE id=:id"), {"p": hash_password(password), "id": user_id})
        add_notification(f"{current.username} reset password for {u.username}", kind='user', conn=conn)
        log_activity(current.username, "reset_password", "user", user_id, u.username, conn=conn)
    return redirect("/users?success=password_reset")


@app.get("/users/{user_id}/edit")
def user_edit_redirect(request: Request, user_id: int):
    require_admin(request)
    return redirect(f"/users?edit={user_id}")


@app.post("/users/{user_id}/edit")
def edit_user(request: Request, user_id: int, display_name: str = Form(''), job_title: str = Form(''), role: str = Form('user'), is_active: int = Form(1)):
    current = require_admin(request)
    role = (role or 'user').strip().lower()
    role = normalized_role(role)
    active_value = 1 if str(is_active) == '1' else 0
    with engine.begin() as conn:
        u = conn.execute(text("SELECT id, username, role, is_active, display_name FROM users WHERE id=:id"), {"id": user_id}).fetchone()
        if not u:
            return redirect("/users?error=user_not_found")
        if u.username == current.username and active_value == 0:
            return redirect("/users?error=self_toggle")
        conn.execute(text("UPDATE users SET display_name=:d, job_title=:j, role=:r, is_active=:a WHERE id=:id"), {"d": (display_name or '').strip(), "j": (job_title or '').strip(), "r": role, "a": active_value, "id": user_id})
        add_notification(f"{current.username} updated user {u.username}", kind='user', related_type='user', related_id=user_id, conn=conn)
        log_activity(current.username, "edit_user", "user", user_id, u.username, conn=conn)
    return redirect("/users?success=user_updated")


@app.post("/users/{user_id}/delete")
def delete_user(request: Request, user_id: int):
    current = require_admin(request)
    with engine.begin() as conn:
        u = conn.execute(text("SELECT id, username FROM users WHERE id=:id"), {"id": user_id}).fetchone()
        if not u:
            return redirect("/users?error=user_not_found")
        if u.username == current.username:
            return redirect("/users?error=self_delete")
        conn.execute(text("DELETE FROM notifications WHERE target_username=:u OR actor_username=:u"), {"u": u.username})
        conn.execute(text("DELETE FROM task_comments WHERE author=:u"), {"u": u.username})
        conn.execute(text("UPDATE tasks SET assigned_to='' WHERE assigned_to=:u"), {"u": u.username})
        conn.execute(text("UPDATE tasks SET created_by=:admin WHERE created_by=:u"), {"u": u.username, "admin": current.username})
        conn.execute(text("UPDATE leads SET created_by=:admin WHERE created_by=:u"), {"u": u.username, "admin": current.username})
        conn.execute(text("DELETE FROM users WHERE id=:id"), {"id": user_id})
        add_notification(f"{current.username} deleted user {u.username}", kind='user', related_type='user', related_id=user_id, conn=conn)
        log_activity(current.username, "delete_user", "user", user_id, u.username, conn=conn)
    return redirect("/users?success=user_deleted")




@app.get("/profile", response_class=HTMLResponse)
def profile_page(request: Request):
    user = require_login(request)
    with engine.begin() as conn:
        fresh = conn.execute(text("SELECT * FROM users WHERE username=:u"), {"u": user.username}).fetchone()
    return templates.TemplateResponse("profile.html", {"request": request, "username": user.username, "user": fresh, "profile_user": fresh})


@app.post("/profile/save")
async def profile_save(request: Request, display_name: str = Form(''), job_title: str = Form(''), avatar: UploadFile | None = File(None)):
    user = require_login(request)
    avatar_name = ''
    if avatar and getattr(avatar, 'filename', ''):
        safe_name = f"avatar_{user.username}_{int(time.time())}_{secrets.token_hex(4)}_{safe_filename(avatar.filename)}"
        with open(AVATAR_UPLOADS_DIR / safe_name, 'wb') as f:
            f.write(await avatar.read())
        avatar_name = safe_name
    with engine.begin() as conn:
        params = {"d": (display_name or '').strip(), "j": (job_title or '').strip(), "u": user.username}
        if avatar_name:
            conn.execute(text("UPDATE users SET display_name=:d, job_title=:j, avatar_path=:a WHERE username=:u"), {**params, "a": avatar_name})
        else:
            conn.execute(text("UPDATE users SET display_name=:d, job_title=:j WHERE username=:u"), params)
        log_activity(user.username, 'update_profile', 'user', 0, user.username, conn=conn)
    return redirect('/profile?success=profile_saved')


@app.post("/profile/request-password")
def profile_request_password(request: Request, new_password: str = Form(...)):
    user = require_login(request)
    new_password = (new_password or '').strip()
    if len(new_password) < 4:
        return redirect('/profile?error=password_short')
    with engine.begin() as conn:
        conn.execute(text("UPDATE users SET password_change_requested=1, requested_password_hash=:p WHERE username=:u"), {"p": hash_password(new_password), "u": user.username})
        admins = [r.username for r in conn.execute(text("SELECT username FROM users WHERE role='admin' AND is_active=1" )).fetchall()]
        notify_many(admins, f"{user.username} requested a password change approval", kind='user', actor_username=user.username, conn=conn)
        log_activity(user.username, 'request_password_change', 'user', 0, user.username, conn=conn)
    return redirect('/profile?success=password_requested')


@app.post("/users/{user_id}/approve-password-request")
def approve_password_request(request: Request, user_id: int):
    current = require_admin(request)
    with engine.begin() as conn:
        u = conn.execute(text("SELECT username, requested_password_hash, password_change_requested FROM users WHERE id=:id"), {"id": user_id}).fetchone()
        if not u or not safe_int(getattr(u, 'password_change_requested', 0), 0) or not (u.requested_password_hash or '').strip():
            return redirect('/users?error=user_not_found')
        conn.execute(text("UPDATE users SET password=:p, password_change_requested=0, requested_password_hash='' WHERE id=:id"), {"p": u.requested_password_hash, "id": user_id})
        add_notification(f"{current.username} approved password change for {u.username}", kind='user', conn=conn)
        notify_many([u.username], 'Your password change request was approved.', kind='user', actor_username=current.username, conn=conn)
        log_activity(current.username, 'approve_password_request', 'user', user_id, u.username, conn=conn)
    return redirect('/users?success=password_request_approved')


@app.post("/users/{user_id}/reject-password-request")
def reject_password_request(request: Request, user_id: int):
    current = require_admin(request)
    with engine.begin() as conn:
        u = conn.execute(text("SELECT username FROM users WHERE id=:id"), {"id": user_id}).fetchone()
        if not u:
            return redirect('/users?error=user_not_found')
        conn.execute(text("UPDATE users SET password_change_requested=0, requested_password_hash='' WHERE id=:id"), {"id": user_id})
        notify_many([u.username], 'Your password change request was declined.', kind='user', actor_username=current.username, conn=conn)
        log_activity(current.username, 'reject_password_request', 'user', user_id, u.username, conn=conn)
    return redirect('/users?success=password_request_rejected')


@app.post("/chat/{message_id}/delete")
def delete_chat_message(request: Request, message_id: int):
    user = require_login(request)
    with engine.begin() as conn:
        msg = conn.execute(text("SELECT * FROM chat_messages WHERE id=:id"), {"id": message_id}).fetchone()
        if not msg:
            return redirect('/chat')
        if user.role != 'admin' and msg.username != user.username:
            return redirect(f'/chat?error=not_allowed')
        channel_id = safe_int(getattr(msg, 'channel_id', 1), 1)
        conn.execute(text("DELETE FROM chat_messages WHERE id=:id"), {"id": message_id})
        log_activity(user.username, 'delete_chat_message', 'chat', message_id, '', conn=conn)
    return redirect(f'/chat?channel_id={channel_id}')


@app.post("/chat/clear")
def clear_chat_channel(request: Request, channel_id: int = Form(1)):
    user = require_admin(request)
    with engine.begin() as conn:
        conn.execute(text("DELETE FROM chat_messages WHERE COALESCE(channel_id,1)=:cid"), {"cid": channel_id})
        log_activity(user.username, 'clear_chat_channel', 'chat_channel', channel_id, '', conn=conn)
    return redirect(f'/chat?channel_id={channel_id}&success=chat_cleared')


@app.get("/notifications/dropdown")
def notifications_dropdown(request: Request):
    user = require_login(request)
    with engine.begin() as conn:
        unread = conn.execute(text(f"SELECT COUNT(*) FROM notifications WHERE is_read=0 AND {notification_where_sql()}"), {"username": user.username}).scalar() or 0
        rows = conn.execute(text(f"SELECT * FROM notifications WHERE {notification_where_sql()} ORDER BY id DESC LIMIT 12"), {"username": user.username}).fetchall()
        max_id = conn.execute(text(f"SELECT COALESCE(MAX(id),0) FROM notifications WHERE {notification_where_sql()}"), {"username": user.username}).scalar() or 0
    prepared_rows = []
    for r in rows:
        item = dict(r._mapping)
        item["target_url"] = notification_target_url(r)
        prepared_rows.append(item)
    return JSONResponse({"unread": unread, "rows": prepared_rows, "max_id": max_id})


@app.post("/notifications/{notification_id}/read")
def notifications_mark_read(request: Request, notification_id: int):
    user = require_login(request)
    with engine.begin() as conn:
        conn.execute(text(f"UPDATE notifications SET is_read=1 WHERE id=:id AND {notification_where_sql()}"), {"id": notification_id, "username": user.username})
    return JSONResponse({"ok": True})


@app.post("/notifications/read_visible")
def notifications_read_visible(request: Request, ids: str = Form('')):
    user = require_login(request)
    only_ids = [int(x) for x in ids.split(',') if x.strip().isdigit()]
    if only_ids:
        with engine.begin() as conn:
            conn.execute(text("UPDATE notifications SET is_read=1 WHERE id IN (%s) AND (COALESCE(target_username,'')='' OR target_username=:username)" % ",".join(str(x) for x in only_ids)), {"username": user.username})
    return redirect("/dashboard")


@app.get("/notifications/go/{notification_id}")
def notification_go(request: Request, notification_id: int):
    user = require_login(request)
    with engine.begin() as conn:
        row = conn.execute(text(f"SELECT * FROM notifications WHERE id=:id AND {notification_where_sql()}"), {"id": notification_id, "username": user.username}).fetchone()
        if not row:
            return redirect('/dashboard')
        conn.execute(text(f"UPDATE notifications SET is_read=1 WHERE id=:id AND {notification_where_sql()}"), {"id": notification_id, "username": user.username})
    return redirect(notification_target_url(row))


@app.get("/notifications", response_class=HTMLResponse)
def notifications_page(request: Request):
    user = require_login(request)
    with engine.begin() as conn:
        rows = conn.execute(text(f"SELECT * FROM notifications WHERE {notification_where_sql()} ORDER BY id DESC LIMIT 200"), {"username": user.username}).fetchall()
    return templates.TemplateResponse("notifications.html", {"request": request, "username": user.username, "user": user, "rows": rows, "latest_notification_id": latest_notification_id_for_user(user.username)})


@app.get("/notifications/poll")
def notifications_poll(request: Request, last_id: int = 0):
    user = require_login(request)
    with engine.begin() as conn:
        rows = conn.execute(text(f"SELECT * FROM notifications WHERE id > :id AND is_read=0 AND {notification_where_sql()} ORDER BY id ASC"), {"id": last_id, "username": user.username}).fetchall()
    data = [dict(r._mapping) for r in rows]
    return JSONResponse({"rows": data})


@app.post("/notifications/read_all")
def notifications_read_all(request: Request):
    user = require_login(request)
    with engine.begin() as conn:
        conn.execute(text(f"UPDATE notifications SET is_read=1 WHERE is_read=0 AND {notification_where_sql()}"), {"username": user.username})
    return redirect("/notifications")



@app.get("/chat", response_class=HTMLResponse)
def chat_page(request: Request, channel_id: int = 1):
    user = require_login(request)
    with engine.begin() as conn:
        channels = list_accessible_channels(conn, user)
        channel_ids = [c.id for c in channels] or [1]
        if channel_id not in channel_ids:
            channel_id = channel_ids[0]
        active_channel = conn.execute(text("SELECT * FROM chat_channels WHERE id=:id"), {'id': channel_id}).fetchone()
        if not active_channel or not user_can_access_channel(conn, user.username, user.role, channel_id):
            return redirect('/chat?error=room_not_allowed')
        needs_password = bool((getattr(active_channel, 'password_hash', '') or '').strip()) and not is_room_password_verified(request, channel_id)
        can_write_room = False if needs_password else user_can_write_channel(conn, user.username, user.role, channel_id)
        rows = []
        if not needs_password:
            rows = conn.execute(text("""
                SELECT m.*, COALESCE(r.username,'') AS reply_username, COALESCE(r.message,'') AS reply_message,
                       COALESCE(u.display_name,'') AS display_name, COALESCE(u.job_title,'') AS job_title, COALESCE(u.avatar_path,'') AS avatar_path
                FROM chat_messages m
                LEFT JOIN chat_messages r ON r.id=m.reply_to_id
                LEFT JOIN users u ON u.username=m.username
                WHERE COALESCE(m.channel_id,1)=:cid
                ORDER BY m.pinned DESC, m.id DESC LIMIT 150
            """), {'cid': channel_id}).fetchall()
        room_members = room_member_rows(conn, channel_id)
        all_users = conn.execute(text("SELECT username, display_name FROM users WHERE is_active=1 ORDER BY username ASC")).fetchall() if user.role == 'admin' else []
    return templates.TemplateResponse("chat.html", {
        "request": request, "username": user.username, "user": user,
        "rows": list(reversed(rows)), "channels": channels, "active_channel_id": channel_id,
        "active_channel": active_channel, "room_members": room_members, "all_users": all_users,
        "needs_room_password": needs_password, "can_write_room": can_write_room,
        "latest_notification_id": latest_notification_id_for_user(user.username),
        "avatar_url": avatar_url, "render_chat_text": render_chat_text
    })


@app.post("/chat/unlock")
def chat_unlock_room(request: Request, channel_id: int = Form(...), room_password: str = Form('')):
    user = require_login(request)
    with engine.begin() as conn:
        if not user_can_access_channel(conn, user.username, user.role, channel_id):
            return redirect('/chat?error=room_not_allowed')
        if not verify_channel_password(conn, channel_id, room_password or ''):
            return redirect(f'/chat?channel_id={channel_id}&error=room_password_invalid')
    response = RedirectResponse(f'/chat?channel_id={channel_id}', status_code=302)
    response.set_cookie(channel_cookie_key(channel_id), '1', max_age=60*60*12, httponly=True, samesite='lax', secure=COOKIE_SECURE)
    return response


@app.post("/chat/send")
async def chat_send(request: Request, channel_id: int = Form(1), message: str = Form(''), sticker: str = Form(''), reply_to_id: int = Form(0), file: UploadFile | None = File(None)):
    user = require_login(request)
    msg = (message or '').strip()
    sticker = (sticker or '').strip().lower()
    if sticker and not msg:
        msg = f'[[sticker:{sticker}]]'
    saved_name = ''
    original_name = ''
    file_type = ''
    if file and getattr(file, 'filename', ''):
        original_name = file.filename
        file_type = chat_file_kind(original_name)
        if file_type:
            safe_name = f"chat_{int(time.time())}_{secrets.token_hex(4)}_{safe_filename(original_name)}"
            saved_name = safe_name
            with open(CHAT_UPLOADS_DIR / safe_name, 'wb') as f:
                f.write(await file.read())
    if not msg and not saved_name:
        return redirect(f"/chat?channel_id={channel_id}")
    with engine.begin() as conn:
        if not conn.execute(text("SELECT id FROM chat_channels WHERE id=:id AND is_active=1"), {'id': channel_id}).fetchone():
            channel_id = 1
        if not user_can_access_channel(conn, user.username, user.role, channel_id):
            return redirect('/chat?error=room_not_allowed')
        if not user_can_write_channel(conn, user.username, user.role, channel_id):
            return redirect(f'/chat?channel_id={channel_id}&error=read_only_room')
        conn.execute(text("""
            INSERT INTO chat_messages (username,message,created_at,channel_id,reply_to_id,file_path,file_name,file_type)
            VALUES (:u,:m,:t,:channel_id,:reply_to_id,:file_path,:file_name,:file_type)
        """), {"u": user.username, "m": msg, "t": time.time(), "channel_id": channel_id,
                "reply_to_id": reply_to_id or 0, "file_path": saved_name, "file_name": original_name, "file_type": file_type})
        log_activity(user.username, "chat_message", "chat", channel_id, (msg or original_name)[:120], conn=conn)
        channel = conn.execute(text("SELECT name FROM chat_channels WHERE id=:id"), {'id': channel_id}).fetchone()
        recipients = [r.username for r in conn.execute(text("SELECT username FROM users WHERE is_active=1 AND username != :u"), {"u": user.username}).fetchall()]
        notify_many(recipients, f"New chat message in #{channel.name if channel else 'General'} from {user.username}", kind='chat', actor_username=user.username, conn=conn)
    return redirect(f"/chat?channel_id={channel_id}")


@app.post("/chat/channel/create")
def create_chat_channel(request: Request, name: str = Form(...), description: str = Form(''), room_password: str = Form(''), member_usernames: list[str] = Form([]), only_members: int = Form(0)):
    user = require_admin(request)
    clean_name = name.strip()
    with engine.begin() as conn:
        existing = conn.execute(text("SELECT id FROM chat_channels WHERE lower(name)=lower(:n)"), {'n': clean_name}).fetchone()
        if existing:
            return redirect(f"/chat?channel_id={existing.id}")
        password_hash = hash_password(room_password.strip()) if (room_password or '').strip() else ''
        channel_id = insert_and_get_id(conn, """
            INSERT INTO chat_channels (name,description,created_by,created_at,is_active,password_hash,only_members,room_type)
            VALUES (:n,:d,:u,:t,1,:p,:only_members,'room')
        """, {'n': clean_name, 'd': description.strip(), 'u': user.username, 't': time.time(), 'p': password_hash, 'only_members': 1 if safe_int(only_members, 0) else 0})
        selected = sorted({(x or '').strip() for x in (member_usernames or []) if (x or '').strip()})
        if user.username not in selected:
            selected.append(user.username)
        for uname in selected:
            conn.execute(text("INSERT INTO chat_channel_members (channel_id, username, role, can_write, joined_at) VALUES (:cid,:u,:r,:w,:t)"), {'cid': channel_id, 'u': uname, 'r': 'owner' if uname == user.username else 'member', 'w': 1, 't': time.time()})
        log_activity(user.username, 'create_channel', 'chat_channel', channel_id, clean_name, conn=conn)
    add_notification(f"{user.username} created room #{clean_name}", kind='chat', related_type='chat_channel', related_id=channel_id)
    return redirect(f"/chat?channel_id={channel_id}")


@app.get("/chat/poll")
def chat_poll(request: Request, channel_id: int = 1):
    user = require_login(request)
    with engine.begin() as conn:
        if not user_can_access_channel(conn, user.username, user.role, channel_id):
            return JSONResponse({"rows": []})
        rows = conn.execute(text("""
            SELECT m.*, COALESCE(r.username,'') AS reply_username, COALESCE(r.message,'') AS reply_message,
                   COALESCE(u.display_name,'') AS display_name, COALESCE(u.job_title,'') AS job_title, COALESCE(u.avatar_path,'') AS avatar_path
            FROM chat_messages m
            LEFT JOIN chat_messages r ON r.id=m.reply_to_id
            LEFT JOIN users u ON u.username=m.username
            WHERE COALESCE(m.channel_id,1)=:cid
            ORDER BY m.pinned DESC, m.id DESC LIMIT 150
        """), {'cid': channel_id}).fetchall()
    return JSONResponse({"rows": [dict(r._mapping) for r in reversed(rows)]})


@app.post("/chat/channel/member/add")
def chat_channel_member_add(request: Request, channel_id: int = Form(...), username: str = Form(...), can_write: int = Form(1)):
    user = require_admin(request)
    with engine.begin() as conn:
        conn.execute(text("DELETE FROM chat_channel_members WHERE channel_id=:cid AND username=:u"), {'cid': channel_id, 'u': username.strip()})
        conn.execute(text("INSERT INTO chat_channel_members (channel_id, username, role, can_write, joined_at) VALUES (:cid,:u,'member',:w,:t)"), {'cid': channel_id, 'u': username.strip(), 'w': 1 if safe_int(can_write,1) else 0, 't': time.time()})
        log_activity(user.username, 'add_chat_room_member', 'chat_channel', channel_id, username.strip(), conn=conn)
    return redirect(f"/chat?channel_id={channel_id}")


@app.post("/chat/channel/member/remove")
def chat_channel_member_remove(request: Request, channel_id: int = Form(...), username: str = Form(...)):
    user = require_admin(request)
    if username.strip() == user.username:
        return redirect(f"/chat?channel_id={channel_id}&error=not_allowed")
    with engine.begin() as conn:
        conn.execute(text("DELETE FROM chat_channel_members WHERE channel_id=:cid AND username=:u"), {'cid': channel_id, 'u': username.strip()})
        log_activity(user.username, 'remove_chat_room_member', 'chat_channel', channel_id, username.strip(), conn=conn)
    return redirect(f"/chat?channel_id={channel_id}")


@app.post("/chat/channel/password")
def chat_channel_password(request: Request, channel_id: int = Form(...), room_password: str = Form('')):
    user = require_admin(request)
    with engine.begin() as conn:
        conn.execute(text("UPDATE chat_channels SET password_hash=:p WHERE id=:id"), {'p': hash_password(room_password.strip()) if (room_password or '').strip() else '', 'id': channel_id})
        log_activity(user.username, 'update_chat_room_password', 'chat_channel', channel_id, '', conn=conn)
    response = RedirectResponse(f"/chat?channel_id={channel_id}", status_code=302)
    if not (room_password or '').strip():
        response.delete_cookie(channel_cookie_key(channel_id))
    return response


@app.get("/pipeline", response_class=HTMLResponse)

def pipeline_page(request: Request):
    user = require_login(request)
    stages = ['Lead', 'Contacted', 'Negotiation', 'Quotation', 'Won', 'Lost']
    buckets = {}
    stage_meta = {}
    with engine.begin() as conn:
        for stage in stages:
            buckets[stage] = conn.execute(text("SELECT * FROM leads WHERE stage=:s ORDER BY id DESC"), {"s": stage}).fetchall()
            stage_meta[stage] = {
                'count': len(buckets[stage]),
                'value': sum(float(getattr(x, 'estimated_value', 0) or 0) for x in buckets[stage]),
            }
    return templates.TemplateResponse("pipeline.html", {"request": request, "username": user.username, "user": user, "stages": stages, "buckets": buckets, 'stage_meta': stage_meta})


@app.get("/analytics", response_class=HTMLResponse)
def analytics_page(request: Request):
    user = require_admin(request)
    with engine.begin() as conn:
        lead_counts = conn.execute(text("SELECT stage, COUNT(*) AS cnt, COALESCE(SUM(estimated_value),0) AS total FROM leads GROUP BY stage ORDER BY cnt DESC, stage ASC")).fetchall()
        task_counts = conn.execute(text("SELECT status, COUNT(*) AS cnt FROM tasks GROUP BY status ORDER BY cnt DESC, status ASC")).fetchall()
        campaign_counts = conn.execute(text("SELECT status, COUNT(*) AS cnt FROM campaigns GROUP BY status ORDER BY cnt DESC, status ASC")).fetchall()
        assignee_summary = conn.execute(text("SELECT COALESCE(assigned_to,'') AS assigned_to, COUNT(*) AS cnt, COALESCE(SUM(estimated_value),0) AS total FROM leads GROUP BY assigned_to ORDER BY cnt DESC, assigned_to ASC")).fetchall()
    return templates.TemplateResponse("analytics.html", {"request": request, "username": user.username, "user": user, "lead_counts": lead_counts, "task_counts": task_counts, "campaign_counts": campaign_counts, "assignee_summary": assignee_summary})


@app.get("/templates", response_class=HTMLResponse)
def templates_page(request: Request):
    user = require_login(request)
    with engine.begin() as conn:
        rows = conn.execute(text("SELECT * FROM templates ORDER BY id DESC")).fetchall()
    return templates.TemplateResponse("templates.html", {"request": request, "username": user.username, "user": user, "rows": rows, "variables": ['{{company}}','{{contact_person}}','{{phone}}','{{email}}','{{country}}','{{city}}','{{source}}','{{assigned_to}}','{{tags}}','{{type}}','{{status}}','{{stage}}']})


@app.post("/templates/save")
def save_template(request: Request, name: str = Form(...), subject: str = Form(...), body: str = Form(...), template_id: int = Form(0)):
    user = require_login(request)
    with engine.begin() as conn:
        if template_id:
            conn.execute(text("UPDATE templates SET name=:n, subject=:s, body=:b WHERE id=:id"), {"n": name.strip(), "s": subject, "b": body, "id": template_id})
            add_notification(f"{user.username} updated template {name}", kind='template')
            log_activity(user.username, "update_template", "template", template_id, name.strip(), conn=conn)
        else:
            conn.execute(text("INSERT INTO templates (name,subject,body,created_at) VALUES (:n,:s,:b,:t)"), {"n": name.strip(), "s": subject, "b": body, "t": time.time()})
            add_notification(f"{user.username} created template {name}", kind='template')
            log_activity(user.username, "create_template", "template", 0, name.strip(), conn=conn)
    return redirect("/templates")


@app.post("/templates/{template_id}/delete")
def delete_template(request: Request, template_id: int):
    user = require_login(request)
    with engine.begin() as conn:
        conn.execute(text("DELETE FROM templates WHERE id=:id"), {"id": template_id})
        log_activity(user.username, "delete_template", "template", template_id, '', conn=conn)
    add_notification(f"{user.username} deleted a template", kind='template')
    return redirect("/templates")


@app.get("/bridge", response_class=HTMLResponse)
def bridge_page(request: Request, token: str = ''):
    user = require_login(request)
    mobile = is_mobile_request(request)
    with engine.begin() as conn:
        agents = conn.execute(text("SELECT * FROM bridge_agents WHERE username=:u ORDER BY id DESC"), {"u": user.username}).fetchall()
        pending_count = conn.execute(text("SELECT COUNT(*) FROM campaign_jobs WHERE assigned_username=:u AND status IN ('queued','sending')"), {"u": user.username}).scalar() or 0
    return templates.TemplateResponse("bridge.html", {
        "request": request,
        "username": user.username,
        "user": user,
        "agents": agents,
        "new_token": token.strip(),
        "pending_count": pending_count,
        "render_base_url": build_tracking_base_url(request),
        "is_mobile": mobile,
    })


@app.post("/bridge/register")
def bridge_register_browser(request: Request, device_name: str = Form(''), outlook_account_email: str = Form('')):
    user = require_editor(request)
    ensure_desktop_only(request)
    device_name = (device_name or '').strip() or socket.gethostname()
    token = secrets.token_urlsafe(32)
    with engine.begin() as conn:
        conn.execute(text("""
            INSERT INTO bridge_agents (username, device_name, device_token, outlook_account_email, is_active, last_seen_at, created_at)
            VALUES (:username, :device_name, :device_token, :outlook_account_email, 1, 0, :created_at)
        """), {
            "username": user.username,
            "device_name": device_name[:255],
            "device_token": token,
            "outlook_account_email": (outlook_account_email or '').strip()[:255],
            "created_at": time.time(),
        })
        log_activity(user.username, 'register_bridge_agent', 'bridge_agent', 0, device_name, conn=conn)
    return redirect(f"/bridge?token={urllib.parse.quote(token)}")


@app.post("/bridge/{agent_id}/deactivate")
def bridge_deactivate_browser(request: Request, agent_id: int):
    user = require_login(request)
    ensure_desktop_only(request)
    with engine.begin() as conn:
        conn.execute(text("UPDATE bridge_agents SET is_active=0 WHERE id=:id AND username=:u"), {"id": agent_id, "u": user.username})
        log_activity(user.username, 'deactivate_bridge_agent', 'bridge_agent', agent_id, '', conn=conn)
    return redirect('/bridge')


@app.post("/api/bridge/heartbeat")
def api_bridge_heartbeat(payload: BridgeHeartbeatPayload):
    device_token = (payload.device_token or '').strip()
    with engine.begin() as conn:
        agent = get_bridge_agent_by_token(conn, device_token)
        if not agent:
            raise HTTPException(status_code=401, detail="Invalid device token")
        conn.execute(text("UPDATE bridge_agents SET last_seen_at=:t WHERE id=:id"), {"t": time.time(), "id": agent.id})
    return {"ok": True}


@app.get("/api/bridge/jobs")
def api_bridge_jobs(x_device_token: str = Header(...)):
    with engine.begin() as conn:
        agent = get_bridge_agent_by_token(conn, x_device_token)
        if not agent:
            raise HTTPException(status_code=401, detail="Invalid device token")
        rows = conn.execute(text("""
            SELECT id, email, subject_snapshot, body_snapshot, from_email
            FROM campaign_jobs
            WHERE assigned_username=:username
              AND status='queued'
              AND COALESCE(email,'') != ''
            ORDER BY id ASC
            LIMIT 10
        """), {"username": agent.username}).fetchall()
        job_ids = [row.id for row in rows]
        if job_ids:
            conn.execute(text(f"UPDATE campaign_jobs SET status='sending', assigned_device_token=:device_token, last_attempt_at=:last_attempt_at WHERE id IN ({','.join(str(int(i)) for i in job_ids)})"), {
                "device_token": agent.device_token,
                "last_attempt_at": time.time(),
            })
            for job_id in job_ids:
                conn.execute(text("INSERT INTO campaign_job_logs (job_id,status,message,created_at) VALUES (:job_id,'sending',:message,:created_at)"), {
                    "job_id": job_id,
                    "message": f"Picked by device {agent.device_name}",
                    "created_at": time.time(),
                })
        return {
            "jobs": [
                {
                    "id": row.id,
                    "to_email": row.email or '',
                    "subject": row.subject_snapshot or '',
                    "html_body": row.body_snapshot or '',
                    "from_email": row.from_email or '',
                }
                for row in rows
            ]
        }


@app.post("/api/bridge/job-result")
def api_bridge_job_result(payload: BridgeJobResultPayload):
    status = (payload.status or '').strip().lower()
    if status not in {'sent', 'failed'}:
        raise HTTPException(status_code=400, detail='Invalid status')
    with engine.begin() as conn:
        agent = get_bridge_agent_by_token(conn, payload.device_token)
        if not agent:
            raise HTTPException(status_code=401, detail="Invalid device token")
        job = conn.execute(text("SELECT * FROM campaign_jobs WHERE id=:id"), {"id": payload.job_id}).fetchone()
        if not job:
            raise HTTPException(status_code=404, detail='Job not found')
        if (job.assigned_username or '') != (agent.username or ''):
            raise HTTPException(status_code=403, detail='Job does not belong to this user')
        conn.execute(text("""
            UPDATE campaign_jobs
            SET status=:status,
                error_message=:error_message,
                provider_message_id=:provider_message_id,
                sent_at=:sent_at
            WHERE id=:id
        """), {
            "status": status,
            "error_message": (payload.message or '')[:1000],
            "provider_message_id": (payload.provider_message_id or '')[:255],
            "sent_at": time.time() if status == 'sent' else 0,
            "id": payload.job_id,
        })
        conn.execute(text("INSERT INTO campaign_job_logs (job_id,status,message,created_at) VALUES (:job_id,:status,:message,:created_at)"), {
            "job_id": payload.job_id,
            "status": status,
            "message": (payload.message or '')[:1000],
            "created_at": time.time(),
        })
        campaign = conn.execute(text("SELECT campaign_id FROM campaign_jobs WHERE id=:id"), {"id": payload.job_id}).fetchone()
        if campaign:
            remaining = conn.execute(text("SELECT COUNT(*) FROM campaign_jobs WHERE campaign_id=:cid AND status IN ('queued','sending','pending')"), {"cid": campaign.campaign_id}).scalar() or 0
            if remaining == 0:
                conn.execute(text("UPDATE campaigns SET status='processed' WHERE id=:id"), {"id": campaign.campaign_id})
    return {"ok": True}


@app.get("/campaigns", response_class=HTMLResponse)
def campaigns_page(request: Request):
    user = require_login(request)
    with engine.begin() as conn:
        campaigns = conn.execute(text("SELECT * FROM campaigns ORDER BY id DESC")).fetchall()
        templates_rows = conn.execute(text("SELECT * FROM templates ORDER BY name ASC")).fetchall()
        leads = conn.execute(text("SELECT * FROM leads ORDER BY id DESC LIMIT 500")).fetchall()
    return templates.TemplateResponse("campaigns.html", {"request": request, "username": user.username, "user": user, "campaigns": campaigns, "templates_rows": templates_rows, "leads": leads, "variables": ['{{company}}','{{email}}','{{country}}','{{type}}','{{status}}','{{stage}}'], "is_mobile": is_mobile_request(request)})


@app.get("/template/{template_id}")
def get_template(request: Request, template_id: int):
    require_login(request)
    with engine.begin() as conn:
        row = conn.execute(text("SELECT * FROM templates WHERE id=:id"), {"id": template_id}).fetchone()
    if not row:
        return JSONResponse({"ok": False}, status_code=404)
    return JSONResponse({"ok": True, "template": dict(row._mapping)})


@app.post("/campaigns/create")
def create_campaign(request: Request, name: str = Form(...), sender_mode: str = Form('outlook'), template_id: int = Form(0), subject: str = Form(...), body: str = Form(...), lead_ids: list[str] = Form([])):
    user = require_login(request)
    if is_mobile_request(request):
        return redirect('/campaigns?mobile=1')
    selected = [int(x) for x in (lead_ids or []) if str(x).isdigit()]
    with engine.begin() as conn:
        campaign_id = insert_and_get_id(conn, """
        INSERT INTO campaigns (name,sender_mode,template_id,subject,body,status,created_by,created_at)
        VALUES (:n,:m,:tid,:s,:b,'queued',:u,:t)
        """, {"n": name.strip(), "m": sender_mode, "tid": template_id or None, "s": subject, "b": body, "u": user.username, "t": time.time()})
        if selected:
            rows = conn.execute(text("SELECT id,email FROM leads WHERE id IN (%s)" % ",".join(str(i) for i in selected))).fetchall()
            for row in rows:
                conn.execute(text("INSERT INTO campaign_jobs (campaign_id,lead_id,email,status,created_at) VALUES (:cid,:lid,:e,'pending',:t)"), {"cid": campaign_id, "lid": row.id, "e": row.email or '', "t": time.time()})
        log_activity(user.username, "create_campaign", "campaign", campaign_id, name.strip(), conn=conn)
    add_notification(f"{user.username} queued campaign {name}", kind='campaign', related_type='campaign', related_id=campaign_id)
    return redirect(f"/campaign/{campaign_id}")


@app.get("/campaign/{campaign_id}", response_class=HTMLResponse)
def campaign_detail(request: Request, campaign_id: int):
    user = require_login(request)
    mobile = is_mobile_request(request)
    with engine.begin() as conn:
        campaign = conn.execute(text("SELECT * FROM campaigns WHERE id=:id"), {"id": campaign_id}).fetchone()
        jobs = conn.execute(text("SELECT cj.*, l.company FROM campaign_jobs cj LEFT JOIN leads l ON l.id=cj.lead_id WHERE cj.campaign_id=:id ORDER BY cj.id ASC"), {"id": campaign_id}).fetchall()
        device_rows = conn.execute(text("SELECT username, device_name, outlook_account_email, is_active, last_seen_at FROM bridge_agents WHERE is_active=1 ORDER BY username ASC, device_name ASC")).fetchall()
    return templates.TemplateResponse("campaign_detail.html", {"request": request, "username": user.username, "user": user, "campaign": campaign, "jobs": jobs, "device_rows": device_rows})


@app.post("/campaign/{campaign_id}/send")
def campaign_send(request: Request, campaign_id: int):
    user = require_login(request)
    ensure_desktop_only(request)
    tracking_base = build_tracking_base_url(request)
    with engine.begin() as conn:
        campaign = conn.execute(text("SELECT * FROM campaigns WHERE id=:id"), {"id": campaign_id}).fetchone()
        if not campaign:
            return redirect("/campaigns")
        jobs = conn.execute(text("SELECT cj.*, l.company, l.contact_person, l.phone, l.email, l.country, l.city, l.source, l.assigned_to, l.tags, l.type, l.status AS lead_status, l.stage FROM campaign_jobs cj LEFT JOIN leads l ON l.id=cj.lead_id WHERE cj.campaign_id=:id AND cj.status IN ('pending','failed')"), {"id": campaign_id}).fetchall()
        if (campaign.sender_mode or 'outlook').lower() == 'outlook':
            queued_count = queue_campaign_for_outlook_bridge(conn, campaign, jobs, user.username, tracking_base)
            log_activity(user.username, "queue_campaign_bridge", "campaign", campaign_id, campaign.name or '', conn=conn)
            add_notification(f"{user.username} queued campaign #{campaign_id} for Outlook Desktop delivery ({queued_count} emails)", kind='campaign', related_type='campaign', related_id=campaign_id)
            return redirect(f"/campaign/{campaign_id}")

        for job in jobs:
            lead = {
                "company": job.company,
                "contact_person": job.contact_person,
                "phone": job.phone,
                "email": job.email,
                "country": job.country,
                "city": job.city,
                "source": job.source,
                "assigned_to": job.assigned_to,
                "tags": job.tags,
                "type": job.type,
                "status": job.lead_status,
                "stage": job.stage,
            }
            sub, html_body = render_template_text(campaign.subject or '', campaign.body or '', lead)
            tracking_pixel = f'<img src="{tracking_base}/track/open?lead_id={job.lead_id}&campaign_id={campaign_id}&email={urllib.parse.quote(job.email or "")}" width="1" height="1">'
            final_body = html_body + '<br><br>' + tracking_pixel
            try:
                send_email(job.email or '', sub, final_body, campaign.sender_mode or 'smtp')
                conn.execute(text("UPDATE campaign_jobs SET status='sent', sent_at=:t, error_message='' WHERE id=:id"), {"t": time.time(), "id": job.id})
                conn.execute(text("INSERT INTO tracking_events (lead_id,email,campaign_id,event_type,details,created_at) VALUES (:lid,:e,:cid,'queued-send',:d,:t)"), {"lid": job.lead_id, "e": job.email or '', "cid": campaign_id, "d": sub, "t": time.time()})
            except Exception as exc:
                conn.execute(text("UPDATE campaign_jobs SET status='failed', error_message=:e WHERE id=:id"), {"e": str(exc), "id": job.id})
        conn.execute(text("UPDATE campaigns SET status='processed' WHERE id=:id"), {"id": campaign_id})
        log_activity(user.username, "process_campaign", "campaign", campaign_id, campaign.name or '', conn=conn)
    add_notification(f"{user.username} processed campaign #{campaign_id}", kind='campaign', related_type='campaign', related_id=campaign_id)
    return redirect(f"/campaign/{campaign_id}")


@app.get("/tracking", response_class=HTMLResponse)
def tracking_page(request: Request):
    user = require_login(request)
    with engine.begin() as conn:
        rows = conn.execute(text("SELECT * FROM tracking_events ORDER BY id DESC LIMIT 200")).fetchall()
        leads = conn.execute(text("SELECT id,company,email FROM leads ORDER BY id DESC LIMIT 100")).fetchall()
        campaigns = conn.execute(text("SELECT id,name FROM campaigns ORDER BY id DESC LIMIT 100")).fetchall()
    return templates.TemplateResponse("tracking.html", {"request": request, "username": user.username, "user": user, "rows": rows, "leads": leads, "campaigns": campaigns, "tracking_base_url": build_tracking_base_url(request)})


@app.get("/track/open")
def track_open(lead_id: int = 0, campaign_id: int = 0, email: str = ''):
    with engine.begin() as conn:
        conn.execute(text("INSERT INTO tracking_events (lead_id,email,campaign_id,event_type,details,created_at) VALUES (:lid,:e,:cid,'open','pixel',:t)"), {"lid": lead_id, "e": email, "cid": campaign_id, "t": time.time()})
    tiny_gif = b'GIF89a\x01\x00\x01\x00\x80\x00\x00\x00\x00\x00\xff\xff\xff!\xf9\x04\x01\x00\x00\x00\x00,\x00\x00\x00\x00\x01\x00\x01\x00\x00\x02\x02D\x01\x00;'
    return Response(content=tiny_gif, media_type='image/gif')


@app.get("/track/click")
def track_click(target: str, lead_id: int = 0, campaign_id: int = 0, email: str = ''):
    with engine.begin() as conn:
        conn.execute(text("INSERT INTO tracking_events (lead_id,email,campaign_id,event_type,details,created_at) VALUES (:lid,:e,:cid,'click',:d,:t)"), {"lid": lead_id, "e": email, "cid": campaign_id, "d": target, "t": time.time()})
    return redirect(target)


@app.get("/announcements", response_class=HTMLResponse)
def announcements_page(request: Request):
    user = require_login(request)
    with engine.begin() as conn:
        rows = conn.execute(text("SELECT * FROM announcements ORDER BY id DESC")).fetchall()
    return templates.TemplateResponse("announcements.html", {"request": request, "username": user.username, "user": user, "rows": rows})


@app.post("/announcements/add")
def add_announcement(request: Request, title: str = Form(...), body: str = Form(...)):
    user = require_login(request)
    with engine.begin() as conn:
        conn.execute(text("INSERT INTO announcements (title,body,author,created_at) VALUES (:t,:b,:a,:c)"), {"t": title.strip(), "b": body.strip(), "a": user.username, "c": time.time()})
        log_activity(user.username, "post_announcement", "announcement", 0, title.strip(), conn=conn)
    with engine.begin() as conn:
        recipients = [r.username for r in conn.execute(text("SELECT username FROM users WHERE is_active=1 AND username != :u"), {"u": user.username}).fetchall()]
        notify_many(recipients, f"{user.username} posted an announcement: {title.strip()}", kind='announcement', actor_username=user.username, conn=conn)
    return redirect("/announcements")


@app.get("/announcement/{announcement_id}", response_class=HTMLResponse)
def announcement_detail(request: Request, announcement_id: int):
    user = require_login(request)
    with engine.begin() as conn:
        row = conn.execute(text("SELECT * FROM announcements WHERE id=:id"), {"id": announcement_id}).fetchone()
        replies = conn.execute(text("SELECT * FROM announcement_replies WHERE announcement_id=:id ORDER BY id ASC"), {"id": announcement_id}).fetchall()
    return templates.TemplateResponse("announcement_detail.html", {"request": request, "username": user.username, "user": user, "row": row, "replies": replies})


@app.post("/announcement/{announcement_id}/reply")
def announcement_reply(request: Request, announcement_id: int, body: str = Form(...)):
    user = require_login(request)
    with engine.begin() as conn:
        conn.execute(text("INSERT INTO announcement_replies (announcement_id,author,body,created_at) VALUES (:id,:a,:b,:t)"), {"id": announcement_id, "a": user.username, "b": body.strip(), "t": time.time()})
        log_activity(user.username, "reply_announcement", "announcement", announcement_id, body.strip()[:120], conn=conn)
    with engine.begin() as conn:
        recipients = [r.username for r in conn.execute(text("SELECT username FROM users WHERE is_active=1 AND username != :u"), {"u": user.username}).fetchall()]
        notify_many(recipients, f"{user.username} replied to announcement #{announcement_id}", kind='announcement', actor_username=user.username, conn=conn)
    return redirect(f"/announcement/{announcement_id}")


@app.get("/tasks", response_class=HTMLResponse)
def tasks_page(request: Request, status: str = '', assigned_to: str = '', q: str = ''):
    user = require_login(request)
    sql = "SELECT * FROM tasks WHERE 1=1"
    params = {"status": status.strip(), "assigned_to": assigned_to.strip(), "q": f"%{q.strip()}%"}
    if status.strip():
        sql += " AND status=:status"
    if assigned_to.strip():
        sql += " AND assigned_to=:assigned_to"
    if q.strip():
        sql += " AND (title LIKE :q OR description LIKE :q OR created_by LIKE :q OR assigned_to LIKE :q)"
    sql += " ORDER BY CASE WHEN status='Done' THEN 1 ELSE 0 END, CASE WHEN due_at > 0 THEN due_at ELSE 32503680000 END ASC, id DESC"
    with engine.begin() as conn:
        rows = conn.execute(text(sql), params).fetchall()
        users = conn.execute(text("SELECT username, display_name FROM users WHERE is_active=1 ORDER BY username ASC")).fetchall()
        my_open = conn.execute(text("SELECT COUNT(*) FROM tasks WHERE status != 'Done' AND (assigned_to=:u OR assigned_to='__all__')"), {"u": user.username}).scalar() or 0
        overdue = conn.execute(text("SELECT COUNT(*) FROM tasks WHERE status != 'Done' AND due_at > 0 AND due_at < :t"), {"t": time.time()}).scalar() or 0
        done_today = conn.execute(text("SELECT COUNT(*) FROM tasks WHERE status='Done' AND completed_at >= :t"), {"t": time.time()-86400}).scalar() or 0
    return templates.TemplateResponse("tasks.html", {"request": request, "username": user.username, "user": user, "rows": rows, "users": users, "status_filter": status, "assigned_filter": assigned_to, "q": q, "my_open": my_open, "overdue": overdue, "done_today": done_today})


@app.get("/task/{task_id}", response_class=HTMLResponse)
def task_detail(request: Request, task_id: int):
    user = require_login(request)
    with engine.begin() as conn:
        task = conn.execute(text("SELECT * FROM tasks WHERE id=:id"), {"id": task_id}).fetchone()
        if not task:
            return redirect('/tasks?error=task_not_found')
        users = conn.execute(text("SELECT username, display_name FROM users WHERE is_active=1 ORDER BY username ASC")).fetchall()
        comments = conn.execute(text("SELECT * FROM task_comments WHERE task_id=:id ORDER BY id DESC"), {"id": task_id}).fetchall()
        attachments = conn.execute(text("SELECT * FROM entity_attachments WHERE entity_type='task' AND entity_id=:id ORDER BY id DESC"), {"id": task_id}).fetchall()
        followups = conn.execute(text("SELECT * FROM followups WHERE entity_type='task' AND entity_id=:id ORDER BY followup_at ASC, id DESC"), {"id": task_id}).fetchall()
    return templates.TemplateResponse("task_detail.html", {"request": request, "username": user.username, "user": user, "task": task, "users": users, "comments": comments, "attachments": attachments, "followups": followups})


@app.post("/tasks/{task_id}/update")
def update_task(request: Request, task_id: int, title: str = Form(...), description: str = Form(''), assigned_to: str = Form(''), priority: str = Form('Medium'), due_date: str = Form('')):
    user = require_login(request)
    title = title.strip()
    if not title:
        return redirect(f'/task/{task_id}?error=missing_title')
    due_at = 0
    if due_date.strip():
        try:
            due_at = time.mktime(time.strptime(due_date.strip(), '%Y-%m-%d')) + 43200
        except Exception:
            due_at = 0
    priority = priority if priority in {'Low','Medium','High','Urgent'} else 'Medium'
    with engine.begin() as conn:
        old = conn.execute(text("SELECT assigned_to, title FROM tasks WHERE id=:id"), {"id": task_id}).fetchone()
        if not old:
            return redirect('/tasks?error=task_not_found')
        assigned_clean = (assigned_to or '').strip()
        conn.execute(text("UPDATE tasks SET title=:t, description=:d, assigned_to=:a, priority=:p, due_at=:due WHERE id=:id"), {"t": title, "d": description.strip(), "a": assigned_clean, "p": priority, "due": due_at, "id": task_id})
        log_activity(user.username, 'update_task', 'task', task_id, title, conn=conn)
        notify_many(task_assignee_targets(conn, old.assigned_to or '', exclude={user.username}) | task_assignee_targets(conn, assigned_clean, exclude={user.username}), f"Task updated: {title}", kind='task', related_type='task', related_id=task_id, actor_username=user.username, conn=conn)
    return redirect(f'/task/{task_id}?success=task_updated')


@app.post("/tasks/{task_id}/comment")
def add_task_comment(request: Request, task_id: int, body: str = Form(...)):
    user = require_login(request)
    body = body.strip()
    if not body:
        return redirect(f'/task/{task_id}')
    with engine.begin() as conn:
        task = conn.execute(text("SELECT title FROM tasks WHERE id=:id"), {"id": task_id}).fetchone()
        if not task:
            return redirect('/tasks?error=task_not_found')
        conn.execute(text("INSERT INTO task_comments (task_id,author,body,created_at) VALUES (:tid,:a,:b,:t)"), {"tid": task_id, "a": user.username, "b": body, "t": time.time()})
        log_activity(user.username, 'task_comment', 'task', task_id, body[:120], conn=conn)
        full_task = conn.execute(text("SELECT assigned_to, created_by, title FROM tasks WHERE id=:id"), {"id": task_id}).fetchone()
        notify_many(task_assignee_targets(conn, full_task.assigned_to or '', created_by=full_task.created_by or '', exclude={user.username}), f"{user.username} commented on task: {full_task.title}", kind='task', related_type='task', related_id=task_id, actor_username=user.username, conn=conn)
    return redirect(f'/task/{task_id}?success=comment_added')


@app.post("/tasks/add")
def add_task(request: Request, title: str = Form(...), description: str = Form(''), assigned_to: str = Form(''), priority: str = Form('Medium'), due_date: str = Form('')):
    user = require_login(request)
    title = title.strip()
    if not title:
        return redirect('/tasks?error=missing_title')
    due_at = 0
    if due_date.strip():
        try:
            due_at = time.mktime(time.strptime(due_date.strip(), '%Y-%m-%d')) + 43200
        except Exception:
            due_at = 0
    priority = priority if priority in {'Low','Medium','High','Urgent'} else 'Medium'
    with engine.begin() as conn:
        assigned_clean = (assigned_to or '').strip()
        task_id = insert_and_get_id(conn, """
        INSERT INTO tasks (title,description,assigned_to,status,priority,due_at,created_by,created_at)
        VALUES (:t,:d,:a,'Open',:p,:due,:u,:c)
        """, {"t": title, "d": description.strip(), "a": assigned_clean, "p": priority, "due": due_at, "u": user.username, "c": time.time()})
        log_activity(user.username, 'create_task', 'task', task_id, title, conn=conn)
        notify_many(task_assignee_targets(conn, assigned_clean, created_by=user.username, exclude={user.username}), f"Task assigned: {title}", kind='task', related_type='task', related_id=task_id, actor_username=user.username, conn=conn)
    return redirect('/tasks?success=task_created')


@app.post("/tasks/{task_id}/status")
def update_task_status(request: Request, task_id: int, status: str = Form(...)):
    user = require_login(request)
    status = status.strip()
    if status not in {'Open','In Progress','Waiting','Done'}:
        status = 'Open'
    with engine.begin() as conn:
        row = conn.execute(text("SELECT title, assigned_to, created_by FROM tasks WHERE id=:id"), {"id": task_id}).fetchone()
        if not row:
            return redirect('/tasks?error=task_not_found')
        completed_at = time.time() if status == 'Done' else 0
        conn.execute(text("UPDATE tasks SET status=:s, completed_at=:c WHERE id=:id"), {"s": status, "c": completed_at, "id": task_id})
        log_activity(user.username, 'task_status', 'task', task_id, status, conn=conn)
        notify_many(task_assignee_targets(conn, row.assigned_to or '', created_by=row.created_by or '', exclude={user.username}), f"{user.username} changed task '{row.title}' to {status}", kind='task', related_type='task', related_id=task_id, actor_username=user.username, conn=conn)
    return redirect('/tasks?success=task_updated')


@app.post("/tasks/{task_id}/delete")
def delete_task(request: Request, task_id: int):
    user = require_login(request)
    with engine.begin() as conn:
        row = conn.execute(text("SELECT title, assigned_to, created_by FROM tasks WHERE id=:id"), {"id": task_id}).fetchone()
        if not row:
            return redirect('/tasks?error=task_not_found')
        conn.execute(text("DELETE FROM task_comments WHERE task_id=:id"), {"id": task_id})
        conn.execute(text("DELETE FROM tasks WHERE id=:id"), {"id": task_id})
        log_activity(user.username, 'delete_task', 'task', task_id, row.title, conn=conn)
        notify_many(task_assignee_targets(conn, row.assigned_to or '', created_by=row.created_by or '', exclude={user.username}), f"{user.username} deleted task: {row.title}", kind='task', related_type='task', related_id=task_id, actor_username=user.username, conn=conn)
    return redirect('/tasks?success=task_deleted')


@app.get("/activity", response_class=HTMLResponse)
def activity_page(request: Request):
    user = require_admin(request)
    with engine.begin() as conn:
        rows = conn.execute(text("SELECT * FROM activity_logs ORDER BY id DESC LIMIT 300")).fetchall()
    return templates.TemplateResponse("activity.html", {"request": request, "username": user.username, "user": user, "rows": rows})


@app.get("/pi/{lead_id}", response_class=HTMLResponse)
def pi_preview(request: Request, lead_id: int):
    user = require_login(request)
    with engine.begin() as conn:
        lead = conn.execute(text("SELECT * FROM leads WHERE id=:id"), {"id": lead_id}).fetchone()
    return templates.TemplateResponse("pi_preview.html", {"request": request, "username": user.username, "user": user, "lead": lead})


# Campaign V2 module initialization
init_email_campaign_module(app, engine, templates, require_login, is_mobile_request, send_email, build_tracking_base_url, add_notification, log_activity, send_with_profile=send_email_with_profile, generate_agent_content=generate_agent_content)


# -----------------------------
# Export follow-up / shipments module
# -----------------------------
SHIPMENT_UPLOADS_DIR = UPLOADS_DIR / "shipments"
SHIPMENT_UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
TASK_UPLOADS_DIR = UPLOADS_DIR / "tasks"
TASK_UPLOADS_DIR.mkdir(parents=True, exist_ok=True)


def ensure_export_tables():
    dialect = engine.dialect.name

    if dialect == 'postgresql':
        ddl = [
            """
            CREATE TABLE IF NOT EXISTS shipments (
                id SERIAL PRIMARY KEY,
                shipment_no TEXT UNIQUE,
                trade_type TEXT DEFAULT 'Export',
                company TEXT DEFAULT '',
                client_id INTEGER DEFAULT 0,
                supplier TEXT DEFAULT '',
                contact_person TEXT DEFAULT '',
                product_category TEXT DEFAULT 'Dates',
                item_name TEXT DEFAULT '',
                product_name TEXT DEFAULT '',
                quantity DOUBLE PRECISION DEFAULT 0,
                quantity_unit TEXT DEFAULT 'MT',
                container_count INTEGER DEFAULT 0,
                container_type TEXT DEFAULT '40ft',
                cartons_count INTEGER DEFAULT 0,
                invoice_amount DOUBLE PRECISION DEFAULT 0,
                currency TEXT DEFAULT 'USD',
                origin_port TEXT DEFAULT '',
                destination_port TEXT DEFAULT '',
                vessel_name TEXT DEFAULT '',
                etd_at DOUBLE PRECISION DEFAULT 0,
                eta_at DOUBLE PRECISION DEFAULT 0,
                current_status TEXT DEFAULT 'Booked',
                notes TEXT DEFAULT '',
                missing_items TEXT DEFAULT '',
                assigned_to TEXT DEFAULT '',
                created_by TEXT DEFAULT '',
                created_at DOUBLE PRECISION DEFAULT 0,
                updated_at DOUBLE PRECISION DEFAULT 0
            );
            """,
            """
            CREATE TABLE IF NOT EXISTS shipment_documents (
                id SERIAL PRIMARY KEY,
                shipment_id INTEGER NOT NULL,
                doc_type TEXT DEFAULT 'other',
                title TEXT DEFAULT '',
                entry_mode TEXT DEFAULT 'manual',
                manual_text TEXT DEFAULT '',
                filename TEXT DEFAULT '',
                original_name TEXT DEFAULT '',
                notes TEXT DEFAULT '',
                created_by TEXT DEFAULT '',
                created_at DOUBLE PRECISION DEFAULT 0
            );
            """,
            """
            CREATE TABLE IF NOT EXISTS followups (
                id SERIAL PRIMARY KEY,
                entity_type TEXT NOT NULL,
                entity_id INTEGER NOT NULL,
                title TEXT DEFAULT '',
                notes TEXT DEFAULT '',
                followup_at DOUBLE PRECISION DEFAULT 0,
                status TEXT DEFAULT 'Open',
                priority TEXT DEFAULT 'Medium',
                assigned_to TEXT DEFAULT '',
                created_by TEXT DEFAULT '',
                created_at DOUBLE PRECISION DEFAULT 0,
                completed_at DOUBLE PRECISION DEFAULT 0
            );
            """,
            """
            CREATE TABLE IF NOT EXISTS entity_attachments (
                id SERIAL PRIMARY KEY,
                entity_type TEXT NOT NULL,
                entity_id INTEGER NOT NULL,
                filename TEXT NOT NULL,
                original_name TEXT DEFAULT '',
                note TEXT DEFAULT '',
                created_by TEXT DEFAULT '',
                created_at DOUBLE PRECISION DEFAULT 0
            );
            """,
        ]
    else:
        ddl = [
            """
            CREATE TABLE IF NOT EXISTS shipments (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                shipment_no TEXT UNIQUE,
                trade_type TEXT DEFAULT 'Export',
                company TEXT DEFAULT '',
                client_id INTEGER DEFAULT 0,
                supplier TEXT DEFAULT '',
                contact_person TEXT DEFAULT '',
                product_category TEXT DEFAULT 'Dates',
                item_name TEXT DEFAULT '',
                product_name TEXT DEFAULT '',
                quantity REAL DEFAULT 0,
                quantity_unit TEXT DEFAULT 'MT',
                container_count INTEGER DEFAULT 0,
                container_type TEXT DEFAULT '40ft',
                cartons_count INTEGER DEFAULT 0,
                invoice_amount REAL DEFAULT 0,
                currency TEXT DEFAULT 'USD',
                origin_port TEXT DEFAULT '',
                destination_port TEXT DEFAULT '',
                vessel_name TEXT DEFAULT '',
                etd_at REAL DEFAULT 0,
                eta_at REAL DEFAULT 0,
                current_status TEXT DEFAULT 'Booked',
                notes TEXT DEFAULT '',
                missing_items TEXT DEFAULT '',
                assigned_to TEXT DEFAULT '',
                created_by TEXT DEFAULT '',
                created_at REAL DEFAULT 0,
                updated_at REAL DEFAULT 0
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS shipment_documents (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                shipment_id INTEGER NOT NULL,
                doc_type TEXT DEFAULT 'other',
                title TEXT DEFAULT '',
                entry_mode TEXT DEFAULT 'manual',
                manual_text TEXT DEFAULT '',
                filename TEXT DEFAULT '',
                original_name TEXT DEFAULT '',
                notes TEXT DEFAULT '',
                created_by TEXT DEFAULT '',
                created_at REAL DEFAULT 0
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS followups (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                entity_type TEXT NOT NULL,
                entity_id INTEGER NOT NULL,
                title TEXT DEFAULT '',
                notes TEXT DEFAULT '',
                followup_at REAL DEFAULT 0,
                status TEXT DEFAULT 'Open',
                priority TEXT DEFAULT 'Medium',
                assigned_to TEXT DEFAULT '',
                created_by TEXT DEFAULT '',
                created_at REAL DEFAULT 0,
                completed_at REAL DEFAULT 0
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS entity_attachments (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                entity_type TEXT NOT NULL,
                entity_id INTEGER NOT NULL,
                filename TEXT NOT NULL,
                original_name TEXT DEFAULT '',
                note TEXT DEFAULT '',
                created_by TEXT DEFAULT '',
                created_at REAL DEFAULT 0
            )
            """,
        ]

    with engine.begin() as conn:
        for stmt in ddl:
            conn.execute(text(stmt))

        ensure_column(conn, 'shipments', 'client_id',
                      "ALTER TABLE shipments ADD COLUMN client_id INTEGER DEFAULT 0",
                      "ALTER TABLE shipments ADD COLUMN client_id BIGINT DEFAULT 0",
                      dialect)
        ensure_column(conn, 'shipments', 'product_category',
                      "ALTER TABLE shipments ADD COLUMN product_category TEXT DEFAULT 'Dates'",
                      "ALTER TABLE shipments ADD COLUMN product_category TEXT DEFAULT 'Dates'",
                      dialect)
        ensure_column(conn, 'shipments', 'item_name',
                      "ALTER TABLE shipments ADD COLUMN item_name TEXT DEFAULT ''",
                      "ALTER TABLE shipments ADD COLUMN item_name TEXT DEFAULT ''",
                      dialect)
        ensure_column(conn, 'shipments', 'container_count',
                      "ALTER TABLE shipments ADD COLUMN container_count INTEGER DEFAULT 0",
                      "ALTER TABLE shipments ADD COLUMN container_count BIGINT DEFAULT 0",
                      dialect)
        ensure_column(conn, 'shipments', 'container_type',
                      "ALTER TABLE shipments ADD COLUMN container_type TEXT DEFAULT '40ft'",
                      "ALTER TABLE shipments ADD COLUMN container_type TEXT DEFAULT '40ft'",
                      dialect)
        ensure_column(conn, 'shipments', 'cartons_count',
                      "ALTER TABLE shipments ADD COLUMN cartons_count INTEGER DEFAULT 0",
                      "ALTER TABLE shipments ADD COLUMN cartons_count BIGINT DEFAULT 0",
                      dialect)
        ensure_column(conn, 'shipments', 'invoice_amount',
                      "ALTER TABLE shipments ADD COLUMN invoice_amount REAL DEFAULT 0",
                      "ALTER TABLE shipments ADD COLUMN invoice_amount DOUBLE PRECISION DEFAULT 0",
                      dialect)
        ensure_column(conn, 'shipments', 'currency',
                      "ALTER TABLE shipments ADD COLUMN currency TEXT DEFAULT 'USD'",
                      "ALTER TABLE shipments ADD COLUMN currency TEXT DEFAULT 'USD'",
                      dialect)


def status_steps():
    return ['Booked', 'Loaded', 'Sailed', 'At Sea', 'Arrived', 'Customs', 'Delivered', 'Closed']


def shipment_progress(status_value: str) -> int:
    steps = status_steps()
    try:
        idx = steps.index((status_value or 'Booked').strip())
    except ValueError:
        idx = 0
    return int(((idx + 1) / len(steps)) * 100)


def shipment_status_meta(row, now_ts: float | None = None):
    now_ts = now_ts or time.time()
    status_value = (getattr(row, 'current_status', '') or 'Booked').strip()
    progress = shipment_progress(status_value)
    eta = float(getattr(row, 'eta_at', 0) or 0)
    etd = float(getattr(row, 'etd_at', 0) or 0)
    if status_value in {'Delivered', 'Closed'}:
        badge = 'done'
        summary = 'Shipment completed'
    elif eta and eta < now_ts:
        badge = 'danger'
        summary = 'Delayed / ETA passed'
    elif status_value == 'At Sea':
        badge = 'sea'
        summary = 'At sea now'
    elif etd and etd <= now_ts:
        badge = 'active'
        summary = 'Moved from Egypt'
    else:
        badge = 'pending'
        summary = status_value or 'Booked'
    remaining_days = max(0, int((eta - now_ts) // 86400)) if eta else None
    return {'progress': progress, 'badge': badge, 'summary': summary, 'remaining_days': remaining_days}


def save_entity_file(upload: UploadFile, prefix: str, folder: Path):
    original_name = Path(upload.filename or '').name
    if not original_name:
        return '', ''
    ext = Path(original_name).suffix.lower()
    safe_name = f"{prefix}_{int(time.time())}_{secrets.token_hex(4)}{ext}"
    data = upload.file.read() if hasattr(upload, 'file') else b''
    if hasattr(upload, 'file'):
        upload.file.seek(0)
        data = upload.file.read()
    else:
        data = b''
    (folder / safe_name).write_bytes(data)
    return safe_name, original_name


def shipment_doc_counts(conn, shipment_id: int):
    rows = conn.execute(text("SELECT doc_type, COUNT(*) AS cnt FROM shipment_documents WHERE shipment_id=:id GROUP BY doc_type"), {"id": shipment_id}).fetchall()
    return {r.doc_type: r.cnt for r in rows}


@app.get('/email-tool', response_class=HTMLResponse)
def email_tool_page(request: Request):
    user = require_login(request)
    cards = [
        {'title': 'Leads & Lists', 'desc': 'Review target accounts, segments and follow-up readiness before sending.', 'href': '/leads'},
        {'title': 'Campaigns', 'desc': 'Create, queue and monitor one-off outreach campaigns.', 'href': '/campaign-v2/campaigns'},
        {'title': 'Sequences', 'desc': 'Run step-by-step auto email follow-up with stop-on-reply controls.', 'href': '/campaign-v2/sequences'},
        {'title': 'Replies & Suppression', 'desc': 'Log replies, bounces and unsubscribes so automation stops correctly.', 'href': '/campaign-v2/replies'},
        {'title': 'Sending Profiles', 'desc': 'Choose which mailbox or Outlook device sends the emails.', 'href': '/campaign-v2/profiles'},
        {'title': 'Template Library', 'desc': 'Prepare reusable templates for export offers and follow-ups.', 'href': '/campaign-v2/templates'},
        {'title': 'AI Agent', 'desc': 'Generate outreach drafts and follow-up copy directly from lead context.', 'href': '/campaign-v2/agent'},
        {'title': 'Activity Logs', 'desc': 'See sent, failed, opened and sequence activity in one place.', 'href': '/campaign-v2/logs'},
        {'title': 'Outlook Bridge', 'desc': 'Register desktop Outlook devices and generate device tokens.', 'href': '/bridge'},
        {'title': 'Bridge Monitor', 'desc': 'Check online devices, heartbeat and pending jobs.', 'href': '/bridge/monitor'},
        {'title': 'Tracking & Signals', 'desc': 'Review opens, clicks and recipient events.', 'href': '/tracking'},
    ]
    return templates.TemplateResponse('email_tool.html', {
        'request': request, 'username': user.username, 'user': user, 'cards': cards,
        'latest_notification_id': latest_notification_id_for_user(user.username)
    })


@app.get('/shipments/dashboard', response_class=HTMLResponse)
def shipments_dashboard(request: Request):
    user = require_login(request)
    ensure_due_notifications(user.username)
    now_ts = time.time()
    with engine.begin() as conn:
        stats = {
            'open_shipments': conn.execute(text("SELECT COUNT(*) FROM shipments WHERE current_status NOT IN ('Delivered','Closed')")).scalar() or 0,
            'arriving_this_week': conn.execute(text("SELECT COUNT(*) FROM shipments WHERE eta_at > :t AND eta_at <= :n AND current_status NOT IN ('Delivered','Closed')"), {'t': now_ts, 'n': now_ts + 7*86400}).scalar() or 0,
            'delayed': conn.execute(text("SELECT COUNT(*) FROM shipments WHERE eta_at > 0 AND eta_at < :t AND current_status NOT IN ('Delivered','Closed')"), {'t': now_ts}).scalar() or 0,
            'missing_docs': conn.execute(text("SELECT COUNT(*) FROM shipments s WHERE current_status NOT IN ('Delivered','Closed') AND (SELECT COUNT(*) FROM shipment_documents d WHERE d.shipment_id=s.id AND d.doc_type IN ('invoice','packing_list')) < 2")).scalar() or 0,
            'followups_today': conn.execute(text("SELECT COUNT(*) FROM followups WHERE status IN ('Open','Waiting') AND followup_at > 0 AND followup_at <= :t"), {'t': now_ts}).scalar() or 0,
            'overdue_tasks': conn.execute(text("SELECT COUNT(*) FROM tasks WHERE status != 'Done' AND due_at > 0 AND due_at < :t"), {'t': now_ts}).scalar() or 0,
        }
        rows = conn.execute(text("SELECT * FROM shipments ORDER BY CASE WHEN current_status IN ('Delivered','Closed') THEN 1 ELSE 0 END, CASE WHEN eta_at > 0 THEN eta_at ELSE 32503680000 END ASC, id DESC LIMIT 24")).fetchall()
        shipments = [{**dict(r._mapping), 'meta': shipment_status_meta(r, now_ts), 'docs': shipment_doc_counts(conn, r.id)} for r in rows]
        followups = conn.execute(text("SELECT * FROM followups WHERE status IN ('Open','Waiting') ORDER BY CASE WHEN followup_at > 0 THEN followup_at ELSE 32503680000 END ASC, id DESC LIMIT 12")).fetchall()
    return templates.TemplateResponse('shipments_dashboard.html', {'request': request, 'username': user.username, 'user': user, 'stats': stats, 'shipments': shipments, 'followups': followups, 'status_steps': status_steps(), 'latest_notification_id': latest_notification_id_for_user(user.username)})


@app.get('/shipments', response_class=HTMLResponse)
def shipments_page(request: Request, q: str = '', trade_type: str = '', status: str = ''):
    user = require_login(request)
    sql = "SELECT * FROM shipments WHERE 1=1"
    params = {'q': f"%{q.strip()}%", 'trade_type': trade_type.strip(), 'status': status.strip()}
    if q.strip():
        sql += " AND (shipment_no LIKE :q OR company LIKE :q OR supplier LIKE :q OR product_name LIKE :q OR origin_port LIKE :q OR destination_port LIKE :q OR vessel_name LIKE :q)"
    if trade_type.strip():
        sql += " AND trade_type=:trade_type"
    if status.strip():
        sql += " AND current_status=:status"
    sql += " ORDER BY id DESC"
    with engine.begin() as conn:
        rows = conn.execute(text(sql), params).fetchall()
        shipments = [{**dict(r._mapping), 'meta': shipment_status_meta(r)} for r in rows]
        users = conn.execute(text("SELECT username, display_name FROM users WHERE is_active=1 ORDER BY username ASC")).fetchall()
        clients = conn.execute(text("SELECT id, company_en, company_ar, customer_code FROM current_clients ORDER BY company_en ASC")).fetchall()
        product_items = conn.execute(text("SELECT DISTINCT item_name FROM shipments WHERE COALESCE(item_name,'')<>'' ORDER BY item_name ASC")).fetchall()
    return templates.TemplateResponse('shipments.html', {'request': request, 'username': user.username, 'user': user, 'rows': shipments, 'users': users, 'clients': clients, 'product_items': product_items, 'q': q, 'trade_type_filter': trade_type, 'status_filter': status, 'status_steps': status_steps()})


@app.post('/shipments/add')
def add_shipment(request: Request, shipment_no: str = Form(...), trade_type: str = Form('Export'), client_id: int = Form(0), company: str = Form(''), supplier: str = Form(''), contact_person: str = Form(''), product_category: str = Form('Dates'), item_name: str = Form(''), product_name: str = Form(''), quantity: float = Form(0), quantity_unit: str = Form('MT'), container_count: int = Form(0), container_type: str = Form('40ft'), cartons_count: int = Form(0), invoice_amount: float = Form(0), currency: str = Form('USD'), origin_port: str = Form(''), destination_port: str = Form(''), vessel_name: str = Form(''), etd_date: str = Form(''), eta_date: str = Form(''), current_status: str = Form('Booked'), notes: str = Form(''), missing_items: str = Form(''), assigned_to: str = Form('')):
    user = require_editor(request)
    def parse_date(v):
        try:
            return time.mktime(time.strptime(v.strip(), '%Y-%m-%d')) + 43200 if v.strip() else 0
        except Exception:
            return 0
    with engine.begin() as conn:
        company_clean = company.strip()
        if int(client_id or 0) > 0:
            client_row = conn.execute(text("SELECT company_en FROM current_clients WHERE id=:id"), {'id': int(client_id or 0)}).fetchone()
            if client_row and (client_row.company_en or '').strip():
                company_clean = (client_row.company_en or '').strip()
        product_name_clean = product_name.strip() or item_name.strip()
        sid = insert_and_get_id(conn, """
        INSERT INTO shipments (shipment_no,trade_type,client_id,company,supplier,contact_person,product_category,item_name,product_name,quantity,quantity_unit,container_count,container_type,cartons_count,invoice_amount,currency,origin_port,destination_port,vessel_name,etd_at,eta_at,current_status,notes,missing_items,assigned_to,created_by,created_at,updated_at)
        VALUES (:shipment_no,:trade_type,:client_id,:company,:supplier,:contact_person,:product_category,:item_name,:product_name,:quantity,:quantity_unit,:container_count,:container_type,:cartons_count,:invoice_amount,:currency,:origin_port,:destination_port,:vessel_name,:etd_at,:eta_at,:current_status,:notes,:missing_items,:assigned_to,:created_by,:created_at,:updated_at)
        """, {'shipment_no': shipment_no.strip(), 'trade_type': trade_type.strip() or 'Export', 'client_id': int(client_id or 0), 'company': company_clean, 'supplier': supplier.strip(), 'contact_person': contact_person.strip(), 'product_category': product_category.strip() or 'Dates', 'item_name': item_name.strip(), 'product_name': product_name_clean, 'quantity': quantity or 0, 'quantity_unit': quantity_unit.strip() or 'MT', 'container_count': int(container_count or 0), 'container_type': container_type.strip() or '40ft', 'cartons_count': int(cartons_count or 0), 'invoice_amount': invoice_amount or 0, 'currency': currency.strip() or 'USD', 'origin_port': origin_port.strip(), 'destination_port': destination_port.strip(), 'vessel_name': vessel_name.strip(), 'etd_at': parse_date(etd_date), 'eta_at': parse_date(eta_date), 'current_status': current_status.strip() or 'Booked', 'notes': notes.strip(), 'missing_items': missing_items.strip(), 'assigned_to': assigned_to.strip(), 'created_by': user.username, 'created_at': time.time(), 'updated_at': time.time()})
        log_activity(user.username, 'create_shipment', 'shipment', sid, shipment_no.strip(), conn=conn)
        notify_many({assigned_to.strip()}, f"Shipment assigned: {shipment_no.strip()}", kind='shipment', related_type='shipment', related_id=sid, actor_username=user.username, conn=conn)
    return redirect('/shipments?success=shipment_created')


@app.get('/shipment/{shipment_id}', response_class=HTMLResponse)
def shipment_detail(request: Request, shipment_id: int):
    user = require_login(request)
    with engine.begin() as conn:
        shipment = conn.execute(text("SELECT * FROM shipments WHERE id=:id"), {'id': shipment_id}).fetchone()
        if not shipment:
            return redirect('/shipments?error=shipment_not_found')
        docs = conn.execute(text("SELECT * FROM shipment_documents WHERE shipment_id=:id ORDER BY id DESC"), {'id': shipment_id}).fetchall()
        deals = conn.execute(text("SELECT * FROM shipment_deals WHERE shipment_id=:id ORDER BY id DESC"), {'id': shipment_id}).fetchall()
        followups = conn.execute(text("SELECT * FROM followups WHERE entity_type='shipment' AND entity_id=:id ORDER BY CASE WHEN status='Done' THEN 1 ELSE 0 END, CASE WHEN followup_at > 0 THEN followup_at ELSE 32503680000 END ASC, id DESC"), {'id': shipment_id}).fetchall()
        users = conn.execute(text("SELECT username, display_name FROM users WHERE is_active=1 ORDER BY username ASC")).fetchall()
        clients = conn.execute(text("SELECT id, company_en, company_ar, customer_code FROM current_clients ORDER BY company_en ASC")).fetchall()
        attachment_rows = conn.execute(text("SELECT * FROM entity_attachments WHERE entity_type='shipment' AND entity_id=:id ORDER BY id DESC"), {'id': shipment_id}).fetchall()
    meta = shipment_status_meta(shipment)
    return templates.TemplateResponse('shipment_detail.html', {'request': request, 'username': user.username, 'user': user, 'shipment': shipment, 'meta': meta, 'status_steps': status_steps(), 'docs': docs, 'deals': deals, 'followups': followups, 'attachments': attachment_rows, 'users': users, 'clients': clients})


@app.post('/shipment/{shipment_id}/update')
def update_shipment(request: Request, shipment_id: int, trade_type: str = Form('Export'), client_id: int = Form(0), company: str = Form(''), supplier: str = Form(''), contact_person: str = Form(''), product_category: str = Form('Dates'), item_name: str = Form(''), product_name: str = Form(''), quantity: float = Form(0), quantity_unit: str = Form('MT'), container_count: int = Form(0), container_type: str = Form('40ft'), cartons_count: int = Form(0), invoice_amount: float = Form(0), currency: str = Form('USD'), origin_port: str = Form(''), destination_port: str = Form(''), vessel_name: str = Form(''), etd_date: str = Form(''), eta_date: str = Form(''), current_status: str = Form('Booked'), notes: str = Form(''), missing_items: str = Form(''), assigned_to: str = Form('')):
    user = require_editor(request)
    def parse_date(v):
        try:
            return time.mktime(time.strptime(v.strip(), '%Y-%m-%d')) + 43200 if v.strip() else 0
        except Exception:
            return 0
    with engine.begin() as conn:
        row = conn.execute(text("SELECT shipment_no FROM shipments WHERE id=:id"), {'id': shipment_id}).fetchone()
        if not row:
            return redirect('/shipments?error=shipment_not_found')
        company_clean = company.strip()
        if int(client_id or 0) > 0:
            client_row = conn.execute(text("SELECT company_en FROM current_clients WHERE id=:id"), {'id': int(client_id or 0)}).fetchone()
            if client_row and (client_row.company_en or '').strip():
                company_clean = (client_row.company_en or '').strip()
        product_name_clean = product_name.strip() or item_name.strip()
        conn.execute(text("UPDATE shipments SET trade_type=:trade_type, client_id=:client_id, company=:company, supplier=:supplier, contact_person=:contact_person, product_category=:product_category, item_name=:item_name, product_name=:product_name, quantity=:quantity, quantity_unit=:quantity_unit, container_count=:container_count, container_type=:container_type, cartons_count=:cartons_count, invoice_amount=:invoice_amount, currency=:currency, origin_port=:origin_port, destination_port=:destination_port, vessel_name=:vessel_name, etd_at=:etd_at, eta_at=:eta_at, current_status=:current_status, notes=:notes, missing_items=:missing_items, assigned_to=:assigned_to, updated_at=:updated_at WHERE id=:id"), {'id': shipment_id, 'trade_type': trade_type.strip() or 'Export', 'client_id': int(client_id or 0), 'company': company_clean, 'supplier': supplier.strip(), 'contact_person': contact_person.strip(), 'product_category': product_category.strip() or 'Dates', 'item_name': item_name.strip(), 'product_name': product_name_clean, 'quantity': quantity or 0, 'quantity_unit': quantity_unit.strip() or 'MT', 'container_count': int(container_count or 0), 'container_type': container_type.strip() or '40ft', 'cartons_count': int(cartons_count or 0), 'invoice_amount': invoice_amount or 0, 'currency': currency.strip() or 'USD', 'origin_port': origin_port.strip(), 'destination_port': destination_port.strip(), 'vessel_name': vessel_name.strip(), 'etd_at': parse_date(etd_date), 'eta_at': parse_date(eta_date), 'current_status': current_status.strip() or 'Booked', 'notes': notes.strip(), 'missing_items': missing_items.strip(), 'assigned_to': assigned_to.strip(), 'updated_at': time.time()})
        log_activity(user.username, 'update_shipment', 'shipment', shipment_id, row.shipment_no, conn=conn)
    return redirect(f'/shipment/{shipment_id}?success=updated')


@app.post('/shipment/{shipment_id}/document/add')
async def add_shipment_document(request: Request, shipment_id: int, doc_type: str = Form('other'), title: str = Form(''), entry_mode: str = Form('manual'), manual_text: str = Form(''), notes: str = Form(''), file: UploadFile | None = File(None)):
    user = require_editor(request)
    filename = ''
    original = ''
    if file and getattr(file, 'filename', ''):
        original = Path(file.filename or '').name
        ext = Path(original).suffix.lower()
        safe_name = f"shipment_{shipment_id}_{int(time.time())}_{secrets.token_hex(4)}{ext}"
        data = await file.read()
        (SHIPMENT_UPLOADS_DIR / safe_name).write_bytes(data)
        filename = safe_name
    mode = 'file' if filename else 'manual'
    with engine.begin() as conn:
        conn.execute(text("INSERT INTO shipment_documents (shipment_id,doc_type,title,entry_mode,manual_text,filename,original_name,notes,created_by,created_at) VALUES (:shipment_id,:doc_type,:title,:entry_mode,:manual_text,:filename,:original_name,:notes,:created_by,:created_at)"), {'shipment_id': shipment_id, 'doc_type': doc_type.strip() or 'other', 'title': title.strip(), 'entry_mode': mode if entry_mode != 'manual' else ('manual' if not filename else 'file'), 'manual_text': manual_text.strip(), 'filename': filename, 'original_name': original, 'notes': notes.strip(), 'created_by': user.username, 'created_at': time.time()})
        log_activity(user.username, 'add_shipment_document', 'shipment', shipment_id, f"{doc_type}:{title}"[:120], conn=conn)
    return redirect(f'/shipment/{shipment_id}?success=document_added')


@app.post('/shipment/{shipment_id}/followup/add')
def add_shipment_followup(request: Request, shipment_id: int, title: str = Form(...), notes: str = Form(''), followup_date: str = Form(''), priority: str = Form('Medium'), assigned_to: str = Form('')):
    user = require_editor(request)
    try:
        followup_at = time.mktime(time.strptime(followup_date.strip(), '%Y-%m-%d')) + 43200 if followup_date.strip() else 0
    except Exception:
        followup_at = 0
    with engine.begin() as conn:
        conn.execute(text("INSERT INTO followups (entity_type,entity_id,title,notes,followup_at,status,priority,assigned_to,created_by,created_at) VALUES ('shipment',:entity_id,:title,:notes,:followup_at,'Open',:priority,:assigned_to,:created_by,:created_at)"), {'entity_id': shipment_id, 'title': title.strip(), 'notes': notes.strip(), 'followup_at': followup_at, 'priority': priority.strip() or 'Medium', 'assigned_to': assigned_to.strip(), 'created_by': user.username, 'created_at': time.time()})
        log_activity(user.username, 'add_followup', 'shipment', shipment_id, title.strip(), conn=conn)
        notify_many({assigned_to.strip()}, f"Shipment follow-up: {title.strip()}", kind='followup', related_type='shipment', related_id=shipment_id, actor_username=user.username, conn=conn)
    return redirect(f'/shipment/{shipment_id}?success=followup_added')


@app.get('/followups', response_class=HTMLResponse)
def followups_page(request: Request, status: str = '', entity_type: str = ''):
    user = require_login(request)
    sql = "SELECT * FROM followups WHERE 1=1"
    params = {'status': status.strip(), 'entity_type': entity_type.strip()}
    if status.strip():
        sql += " AND status=:status"
    if entity_type.strip():
        sql += " AND entity_type=:entity_type"
    sql += " ORDER BY CASE WHEN status='Done' THEN 1 ELSE 0 END, CASE WHEN followup_at > 0 THEN followup_at ELSE 32503680000 END ASC, id DESC"
    with engine.begin() as conn:
        rows = conn.execute(text(sql), params).fetchall()
    return templates.TemplateResponse('followups.html', {'request': request, 'username': user.username, 'user': user, 'rows': rows, 'status_filter': status, 'entity_filter': entity_type})


@app.post('/followups/{followup_id}/status')
def update_followup_status(request: Request, followup_id: int, status: str = Form(...)):
    user = require_editor(request)
    status = status.strip() if status.strip() in {'Open','Waiting','Done','Cancelled'} else 'Open'
    with engine.begin() as conn:
        row = conn.execute(text("SELECT * FROM followups WHERE id=:id"), {'id': followup_id}).fetchone()
        if not row:
            return redirect('/followups?error=not_found')
        conn.execute(text("UPDATE followups SET status=:status, completed_at=:completed_at WHERE id=:id"), {'status': status, 'completed_at': time.time() if status == 'Done' else 0, 'id': followup_id})
        log_activity(user.username, 'followup_status', row.entity_type, row.entity_id, status, conn=conn)
    target = '/followups' if row.entity_type != 'shipment' else f'/shipment/{row.entity_id}'
    return redirect(target + '?success=followup_status_updated')


@app.post('/tasks/{task_id}/upload')
async def upload_task_attachment(request: Request, task_id: int, note: str = Form(''), file: UploadFile = File(...)):
    user = require_login(request)
    original = Path(file.filename or '').name
    if not original:
        return redirect(f'/task/{task_id}?error=invalid_file')
    ext = Path(original).suffix.lower()
    safe_name = f"task_{task_id}_{int(time.time())}_{secrets.token_hex(4)}{ext}"
    data = await file.read()
    (TASK_UPLOADS_DIR / safe_name).write_bytes(data)
    with engine.begin() as conn:
        conn.execute(text("INSERT INTO entity_attachments (entity_type,entity_id,filename,original_name,note,created_by,created_at) VALUES ('task',:entity_id,:filename,:original_name,:note,:created_by,:created_at)"), {'entity_id': task_id, 'filename': safe_name, 'original_name': original, 'note': note.strip(), 'created_by': user.username, 'created_at': time.time()})
        log_activity(user.username, 'task_attachment', 'task', task_id, original, conn=conn)
    return redirect(f'/task/{task_id}?success=attachment_added')


@app.post('/tasks/{task_id}/followup/add')
def add_task_followup(request: Request, task_id: int, title: str = Form(...), notes: str = Form(''), followup_date: str = Form(''), priority: str = Form('Medium'), assigned_to: str = Form('')):
    user = require_editor(request)
    try:
        followup_at = time.mktime(time.strptime(followup_date.strip(), '%Y-%m-%d')) + 43200 if followup_date.strip() else 0
    except Exception:
        followup_at = 0
    with engine.begin() as conn:
        conn.execute(text("INSERT INTO followups (entity_type,entity_id,title,notes,followup_at,status,priority,assigned_to,created_by,created_at) VALUES ('task',:entity_id,:title,:notes,:followup_at,'Open',:priority,:assigned_to,:created_by,:created_at)"), {'entity_id': task_id, 'title': title.strip(), 'notes': notes.strip(), 'followup_at': followup_at, 'priority': priority.strip() or 'Medium', 'assigned_to': assigned_to.strip(), 'created_by': user.username, 'created_at': time.time()})
        log_activity(user.username, 'task_followup', 'task', task_id, title.strip(), conn=conn)
    return redirect(f'/task/{task_id}?success=followup_added')




@app.get('/reminders', response_class=HTMLResponse)
def reminders_page(request: Request):
    user = require_login(request)
    ensure_due_notifications(user.username)
    now_ts = time.time()
    with engine.begin() as conn:
        followups_due = conn.execute(text("""
            SELECT * FROM followups
            WHERE status IN ('Open','Waiting')
              AND followup_at > 0
              AND (COALESCE(assigned_to,'')='' OR assigned_to=:u)
            ORDER BY followup_at ASC, id DESC
            LIMIT 100
        """), {'u': user.username}).fetchall()
        task_due = conn.execute(text("""
            SELECT * FROM tasks
            WHERE status != 'Done' AND due_at > 0
            ORDER BY due_at ASC, id DESC
            LIMIT 100
        """), {}).fetchall()
        lead_due = conn.execute(text("""
            SELECT id, company, contact_person, next_followup_at, stage, assigned_to
            FROM leads
            WHERE next_followup_at > 0
            ORDER BY next_followup_at ASC, id DESC
            LIMIT 100
        """), {}).fetchall()
    return templates.TemplateResponse('reminders.html', {
        'request': request, 'username': user.username, 'user': user,
        'followups_due': followups_due, 'task_due': task_due, 'lead_due': lead_due,
        'now_ts': now_ts, 'latest_notification_id': latest_notification_id_for_user(user.username)
    })

@app.get('/bridge/monitor', response_class=HTMLResponse)
def bridge_monitor(request: Request):
    user = require_login(request)
    with engine.begin() as conn:
        rows = conn.execute(text("SELECT * FROM bridge_agents ORDER BY username ASC, device_name ASC, id DESC")).fetchall() if normalized_role(user.role) in {'admin','manager'} else conn.execute(text("SELECT * FROM bridge_agents WHERE username=:u ORDER BY id DESC"), {'u': user.username}).fetchall()
        pending = {}
        try:
            pend = conn.execute(text("SELECT bridge_agent_id, COUNT(*) AS cnt FROM campaign_jobs WHERE status IN ('queued','pending') GROUP BY bridge_agent_id")).fetchall()
            pending = {int(r.bridge_agent_id or 0): int(r.cnt or 0) for r in pend}
        except Exception:
            pending = {}
    enriched = []
    now_ts = time.time()
    for r in rows:
        age = now_ts - float(getattr(r, 'last_seen_at', 0) or 0)
        online = age <= 90
        enriched.append({**dict(r._mapping), 'online': online, 'last_seen_age': int(age), 'pending_jobs': pending.get(int(r.id), 0)})
    return templates.TemplateResponse('bridge_monitor.html', {'request': request, 'username': user.username, 'user': user, 'rows': enriched, 'latest_notification_id': latest_notification_id_for_user(user.username)})


def client_label_from_mapping(mapping):
    company_label = (mapping.get('company_en') or mapping.get('company_ar') or '').strip()
    source_label = company_label or (mapping.get('customer_code') or '').strip() or f"Client #{mapping.get('source_id') or mapping.get('client_id') or ''}"
    return source_label, company_label

@app.get('/documents', response_class=HTMLResponse)
def documents_page(request: Request, doc_type: str = '', q: str = ''):
    user = require_login(request)
    doc_type = (doc_type or '').strip()
    q = (q or '').strip()
    rows = []
    with engine.begin() as conn:
        shipment_docs = conn.execute(text("""
            SELECT d.id, 'shipment' AS source_type, d.shipment_id AS source_id, d.doc_type, d.title, d.entry_mode,
                   d.manual_text, d.filename, d.original_name, d.notes, d.created_by, d.created_at,
                   s.shipment_no AS source_label, s.company AS company_label
            FROM shipment_documents d
            LEFT JOIN shipments s ON s.id=d.shipment_id
            ORDER BY d.id DESC
            LIMIT 300
        """), {}).fetchall()
        rows.extend([dict(r._mapping) for r in shipment_docs])

        try:
            invoice_docs = conn.execute(text("""
                SELECT ci.id, 'client_invoice' AS source_type, ci.client_id AS source_id, 'invoice' AS doc_type,
                       ci.invoice_no AS title, CASE WHEN COALESCE(ci.attachment_filename,'')='' THEN 'manual' ELSE 'file' END AS entry_mode,
                       ci.notes AS manual_text, ci.attachment_filename AS filename, ci.attachment_original_name AS original_name,
                       ci.notes, ci.created_by, ci.created_at, cc.customer_code, cc.company_ar, cc.company_en
                FROM client_invoices ci
                LEFT JOIN current_clients cc ON cc.id=ci.client_id
                ORDER BY ci.id DESC
                LIMIT 300
            """), {}).fetchall()
            for r in invoice_docs:
                row = dict(r._mapping)
                source_label, company_label = client_label_from_mapping(row)
                row['source_label'] = source_label
                row['company_label'] = company_label
                rows.append(row)
        except Exception:
            pass

        attach_docs = conn.execute(text("""
            SELECT ea.id, ea.entity_type AS source_type, ea.entity_id AS source_id, 'attachment' AS doc_type,
                   ea.original_name AS title, 'file' AS entry_mode, ea.note AS manual_text,
                   ea.filename AS filename, ea.original_name, ea.note AS notes, ea.created_by, ea.created_at,
                   CASE WHEN ea.entity_type='shipment' THEN (SELECT shipment_no FROM shipments s WHERE s.id=ea.entity_id) ELSE ('Entity #' || ea.entity_id) END AS source_label,
                   '' AS company_label
            FROM entity_attachments ea
            ORDER BY ea.id DESC
            LIMIT 300
        """), {}).fetchall()
        rows.extend([dict(r._mapping) for r in attach_docs])

    if doc_type:
        rows = [r for r in rows if (r.get('doc_type') or '').lower() == doc_type.lower()]
    if q:
        qq = q.lower()
        rows = [r for r in rows if qq in str(r.get('title') or '').lower() or qq in str(r.get('source_label') or '').lower() or qq in str(r.get('company_label') or '').lower() or qq in str(r.get('notes') or '').lower()]
    rows.sort(key=lambda r: float(r.get('created_at') or 0), reverse=True)
    return templates.TemplateResponse('documents.html', {'request': request, 'username': user.username, 'user': user, 'rows': rows[:300], 'doc_type_filter': doc_type, 'q': q, 'latest_notification_id': latest_notification_id_for_user(user.username)})


@app.get('/invoices-center', response_class=HTMLResponse)
def invoices_center(request: Request):
    user = require_login(request)
    rows = []
    with engine.begin() as conn:
        shipment_invoices = conn.execute(text("""
            SELECT d.id, 'shipment' AS source_type, d.shipment_id AS source_id, d.title, d.entry_mode, d.manual_text, d.filename, d.original_name, d.notes, d.created_by, d.created_at, s.shipment_no AS source_label
            FROM shipment_documents d LEFT JOIN shipments s ON s.id=d.shipment_id
            WHERE d.doc_type='invoice' ORDER BY d.id DESC
        """), {}).fetchall()
        rows.extend([dict(r._mapping) for r in shipment_invoices])

        try:
            client_invoices = conn.execute(text("""
                SELECT ci.id, 'client' AS source_type, ci.client_id AS source_id, ci.invoice_no AS title,
                       CASE WHEN COALESCE(ci.attachment_filename,'')='' THEN 'manual' ELSE 'file' END AS entry_mode,
                       ci.notes AS manual_text, ci.attachment_filename AS filename, ci.attachment_original_name AS original_name,
                       ci.notes, ci.created_by, ci.created_at, cc.customer_code, cc.company_ar, cc.company_en
                FROM client_invoices ci
                LEFT JOIN current_clients cc ON cc.id=ci.client_id
                ORDER BY ci.id DESC
            """), {}).fetchall()
            for r in client_invoices:
                row = dict(r._mapping)
                source_label, company_label = client_label_from_mapping(row)
                row['source_label'] = source_label
                row['company_label'] = company_label
                rows.append(row)
        except Exception:
            pass

    rows.sort(key=lambda r: float(r.get('created_at') or 0), reverse=True)
    return templates.TemplateResponse('invoices_center.html', {'request': request, 'username': user.username, 'user': user, 'rows': rows, 'latest_notification_id': latest_notification_id_for_user(user.username)})


@app.get('/packing-lists', response_class=HTMLResponse)
def packing_lists_center(request: Request):
    user = require_login(request)
    with engine.begin() as conn:
        rows = conn.execute(text("""
            SELECT d.*, s.shipment_no, s.company
            FROM shipment_documents d
            LEFT JOIN shipments s ON s.id=d.shipment_id
            WHERE d.doc_type='packing_list'
            ORDER BY d.id DESC
        """), {}).fetchall()
    return templates.TemplateResponse('packing_lists.html', {'request': request, 'username': user.username, 'user': user, 'rows': rows, 'latest_notification_id': latest_notification_id_for_user(user.username)})

@app.post('/bridge/{agent_id}/disable')
def bridge_disable(request: Request, agent_id: int):
    user = require_editor(request)
    ensure_desktop_only(request)
    with engine.begin() as conn:
        row = conn.execute(text("SELECT * FROM bridge_agents WHERE id=:id AND username=:u"), {'id': agent_id, 'u': user.username}).fetchone()
        if not row:
            return redirect('/bridge?error=device_not_found')
        conn.execute(text("UPDATE bridge_agents SET is_active=0 WHERE id=:id"), {'id': agent_id})
        log_activity(user.username, 'disable_bridge', 'bridge_agent', agent_id, row.device_name, conn=conn)
    return redirect('/bridge?success=device_disabled')


@app.post('/bridge/{agent_id}/delete')
def bridge_delete(request: Request, agent_id: int):
    user = require_editor(request)
    ensure_desktop_only(request)
    with engine.begin() as conn:
        row = conn.execute(text("SELECT * FROM bridge_agents WHERE id=:id AND username=:u"), {'id': agent_id, 'u': user.username}).fetchone()
        if not row:
            return redirect('/bridge?error=device_not_found')
        conn.execute(text("DELETE FROM bridge_agents WHERE id=:id"), {'id': agent_id})
        log_activity(user.username, 'delete_bridge', 'bridge_agent', agent_id, row.device_name, conn=conn)
    return redirect('/bridge?success=device_deleted')


@app.post('/bridge/{agent_id}/regenerate')
def bridge_regenerate(request: Request, agent_id: int):
    user = require_editor(request)
    ensure_desktop_only(request)
    new_token = secrets.token_urlsafe(32)
    with engine.begin() as conn:
        row = conn.execute(text("SELECT * FROM bridge_agents WHERE id=:id AND username=:u"), {'id': agent_id, 'u': user.username}).fetchone()
        if not row:
            return redirect('/bridge?error=device_not_found')
        conn.execute(text("UPDATE bridge_agents SET device_token=:t, is_active=1 WHERE id=:id"), {'t': new_token, 'id': agent_id})
        log_activity(user.username, 'regenerate_bridge', 'bridge_agent', agent_id, row.device_name, conn=conn)
    return redirect('/bridge?token=' + urllib.parse.quote(new_token))


ensure_export_tables()


# -----------------------------
# Product settings + Export engine
# -----------------------------

def save_product_profile(data: dict):
    cleaned = {k: v for k, v in (data or {}).items() if v is not None}
    PRODUCT_PROFILE.update(cleaned)
    PRODUCT_PROFILE_FILE.write_text(json.dumps(PRODUCT_PROFILE, ensure_ascii=False, indent=2), encoding='utf-8')
    templates.env.globals["branding"] = PRODUCT_PROFILE


def load_export_catalog() -> dict:
    default = {
        "products": [],
        "markets": [],
    }
    if EXPORT_CATALOG_FILE.exists():
        try:
            data = json.loads(EXPORT_CATALOG_FILE.read_text(encoding='utf-8'))
            if isinstance(data, dict):
                if isinstance(data.get('products'), list):
                    default['products'] = data.get('products')
                if isinstance(data.get('markets'), list):
                    default['markets'] = data.get('markets')
        except Exception:
            pass
    return default


def ensure_export_engine_tables():
    pg = engine.dialect.name == 'postgresql'
    id_col = 'SERIAL PRIMARY KEY' if pg else 'INTEGER PRIMARY KEY AUTOINCREMENT'
    real_col = 'DOUBLE PRECISION' if pg else 'REAL'
    with engine.begin() as conn:
        conn.execute(text(f"""
            CREATE TABLE IF NOT EXISTS export_products (
                id {id_col},
                product_name TEXT NOT NULL,
                category TEXT DEFAULT '',
                origin_country TEXT DEFAULT '',
                packaging TEXT DEFAULT '',
                price_band TEXT DEFAULT '',
                target_markets TEXT DEFAULT '',
                market_fit TEXT DEFAULT '',
                notes TEXT DEFAULT '',
                sort_order INTEGER DEFAULT 0,
                created_at {real_col} DEFAULT 0,
                updated_at {real_col} DEFAULT 0
            )
        """))
        conn.execute(text(f"""
            CREATE TABLE IF NOT EXISTS export_markets (
                id {id_col},
                country_name TEXT NOT NULL,
                region_name TEXT DEFAULT '',
                demand_level TEXT DEFAULT '',
                price_position TEXT DEFAULT '',
                preferred_products TEXT DEFAULT '',
                buyer_types TEXT DEFAULT '',
                entry_points TEXT DEFAULT '',
                market_notes TEXT DEFAULT '',
                sort_order INTEGER DEFAULT 0,
                created_at {real_col} DEFAULT 0,
                updated_at {real_col} DEFAULT 0
            )
        """))
        products_count = conn.execute(text('SELECT COUNT(*) FROM export_products')).scalar() or 0
        markets_count = conn.execute(text('SELECT COUNT(*) FROM export_markets')).scalar() or 0
        seed = load_export_catalog()
        now = time.time()
        if not products_count:
            for item in seed.get('products', []):
                payload = dict(item)
                payload.setdefault('created_at', now)
                payload.setdefault('updated_at', now)
                conn.execute(text("""
                    INSERT INTO export_products
                    (product_name, category, origin_country, packaging, price_band, target_markets, market_fit, notes, sort_order, created_at, updated_at)
                    VALUES
                    (:product_name, :category, :origin_country, :packaging, :price_band, :target_markets, :market_fit, :notes, :sort_order, :created_at, :updated_at)
                """), payload)
        if not markets_count:
            for item in seed.get('markets', []):
                payload = dict(item)
                payload.setdefault('created_at', now)
                payload.setdefault('updated_at', now)
                conn.execute(text("""
                    INSERT INTO export_markets
                    (country_name, region_name, demand_level, price_position, preferred_products, buyer_types, entry_points, market_notes, sort_order, created_at, updated_at)
                    VALUES
                    (:country_name, :region_name, :demand_level, :price_position, :preferred_products, :buyer_types, :entry_points, :market_notes, :sort_order, :created_at, :updated_at)
                """), payload)


def generate_export_agent_content(product_name: str, target_country: str, buyer_type: str, goal: str, tone: str = 'professional') -> dict:
    product_name = (product_name or '').strip() or 'Egyptian Dates'
    target_country = (target_country or '').strip() or 'your market'
    buyer_type = (buyer_type or '').strip() or 'importer'
    goal = (goal or '').strip() or 'Open a serious export discussion'
    tone = (tone or 'professional').strip()
    with engine.begin() as conn:
        matches = conn.execute(text("""
            SELECT country_name, region_name, demand_level, preferred_products, entry_points, market_notes
            FROM export_markets
            WHERE country_name LIKE :country OR preferred_products LIKE :product
            ORDER BY sort_order ASC, id DESC
            LIMIT 5
        """), {'country': f'%{target_country}%', 'product': f'%{product_name}%'}).fetchall()
        if not matches:
            matches = conn.execute(text("SELECT country_name, region_name, demand_level, preferred_products, entry_points, market_notes FROM export_markets ORDER BY sort_order ASC LIMIT 3")).fetchall()
    subject = f"{product_name} export offer for {target_country}"
    body = (
        f"<p>Hello,</p>"
        f"<p>We are reaching out from Al Tahhan Golden Dates regarding <b>{product_name}</b> for the <b>{target_country}</b> market.</p>"
        f"<p>We work with buyers such as {buyer_type} partners and can support with product specifications, packing options, and a practical export discussion based on your channel needs.</p>"
        f"<p>Main objective: {goal}.</p>"
        f"<p>Best regards,<br>Export Team</p>"
    )
    recommended_markets = [f"{r.country_name} ({r.region_name}, {r.demand_level})" for r in matches]
    next_actions = ['Send specs and packing options', 'Ask for target quantity and preferred Incoterm', 'Prepare CIF or FOB quotation path']
    offer_notes = '; '.join(filter(None, [getattr(r, 'entry_points', '') for r in matches[:2]])) or 'Lead with product-market fit and documentation readiness.'
    return {
        'summary': f"Focus on {product_name} for {target_country} with a {tone} tone and a buyer profile of {buyer_type}.",
        'recommended_markets': recommended_markets,
        'subject': subject,
        'body': body,
        'next_actions': next_actions,
        'offer_notes': offer_notes,
    }


@app.get('/product-settings', response_class=HTMLResponse)
def product_settings_page(request: Request):
    user = require_editor(request)
    return templates.TemplateResponse('product_settings.html', {
        'request': request,
        'username': user.username,
        'user': user,
        'profile': PRODUCT_PROFILE,
        'title': 'Product Settings',
        'latest_notification_id': latest_notification_id_for_user(user.username),
    })


@app.post('/product-settings')
def product_settings_save(request: Request,
                          product_name: str = Form(...),
                          edition_name: str = Form(...),
                          company_name: str = Form(...),
                          app_title: str = Form(...),
                          hero_tagline: str = Form(...),
                          login_title: str = Form(...),
                          login_subtitle: str = Form(...),
                          landing_tagline: str = Form(...),
                          primary_market: str = Form(...),
                          support_label: str = Form(...),
                          outreach_name: str = Form(...)):
    user = require_editor(request)
    save_product_profile({
        'product_name': product_name.strip(),
        'edition_name': edition_name.strip(),
        'company_name': company_name.strip(),
        'app_title': app_title.strip(),
        'hero_tagline': hero_tagline.strip(),
        'login_title': login_title.strip(),
        'login_subtitle': login_subtitle.strip(),
        'landing_tagline': landing_tagline.strip(),
        'primary_market': primary_market.strip(),
        'support_label': support_label.strip(),
        'outreach_name': outreach_name.strip(),
    })
    add_notification(f'{user.username} updated product settings', kind='settings', target_username=user.username)
    return redirect('/product-settings?success=saved')


@app.get('/export', response_class=HTMLResponse)
def export_home(request: Request):
    user = require_login(request)
    with engine.begin() as conn:
        products = conn.execute(text('SELECT * FROM export_products ORDER BY sort_order ASC, id DESC LIMIT 6')).fetchall()
        markets = conn.execute(text('SELECT * FROM export_markets ORDER BY sort_order ASC, id DESC LIMIT 6')).fetchall()
        shipment_count = conn.execute(text('SELECT COUNT(*) FROM shipments')).scalar() or 0
        active_followups = conn.execute(text("SELECT COUNT(*) FROM followups WHERE status!='Done'" )).scalar() or 0
        docs_count = conn.execute(text('SELECT COUNT(*) FROM shipment_documents')).scalar() or 0
    return templates.TemplateResponse('export_dashboard.html', {
        'request': request,
        'username': user.username,
        'user': user,
        'products': products,
        'markets': markets,
        'shipment_count': int(shipment_count),
        'active_followups': int(active_followups),
        'docs_count': int(docs_count),
        'title': 'Export Engine',
        'latest_notification_id': latest_notification_id_for_user(user.username),
    })


@app.get('/export/products', response_class=HTMLResponse)
def export_products_page(request: Request):
    user = require_login(request)
    with engine.begin() as conn:
        rows = conn.execute(text('SELECT * FROM export_products ORDER BY sort_order ASC, id DESC')).fetchall()
    return templates.TemplateResponse('export_products.html', {
        'request': request,
        'username': user.username,
        'user': user,
        'rows': rows,
        'title': 'Export Products',
        'latest_notification_id': latest_notification_id_for_user(user.username),
    })


@app.post('/export/products/add')
def export_products_add(request: Request,
                        product_name: str = Form(...),
                        category: str = Form(''),
                        origin_country: str = Form(''),
                        packaging: str = Form(''),
                        price_band: str = Form(''),
                        target_markets: str = Form(''),
                        market_fit: str = Form(''),
                        notes: str = Form('')):
    user = require_editor(request)
    now = time.time()
    with engine.begin() as conn:
        conn.execute(text("""
            INSERT INTO export_products
            (product_name, category, origin_country, packaging, price_band, target_markets, market_fit, notes, sort_order, created_at, updated_at)
            VALUES (:product_name, :category, :origin_country, :packaging, :price_band, :target_markets, :market_fit, :notes, 999, :t, :t)
        """), {
            'product_name': product_name.strip(), 'category': category.strip(), 'origin_country': origin_country.strip(),
            'packaging': packaging.strip(), 'price_band': price_band.strip(), 'target_markets': target_markets.strip(),
            'market_fit': market_fit.strip(), 'notes': notes.strip(), 't': now,
        })
    return redirect('/export/products?success=created')


@app.get('/export/markets', response_class=HTMLResponse)
def export_markets_page(request: Request):
    user = require_login(request)
    with engine.begin() as conn:
        rows = conn.execute(text('SELECT * FROM export_markets ORDER BY sort_order ASC, id DESC')).fetchall()
    return templates.TemplateResponse('export_markets.html', {
        'request': request,
        'username': user.username,
        'user': user,
        'rows': rows,
        'title': 'Export Markets',
        'latest_notification_id': latest_notification_id_for_user(user.username),
    })


@app.post('/export/markets/add')
def export_markets_add(request: Request,
                       country_name: str = Form(...),
                       region_name: str = Form(''),
                       demand_level: str = Form(''),
                       price_position: str = Form(''),
                       preferred_products: str = Form(''),
                       buyer_types: str = Form(''),
                       entry_points: str = Form(''),
                       market_notes: str = Form('')):
    user = require_editor(request)
    now = time.time()
    with engine.begin() as conn:
        conn.execute(text("""
            INSERT INTO export_markets
            (country_name, region_name, demand_level, price_position, preferred_products, buyer_types, entry_points, market_notes, sort_order, created_at, updated_at)
            VALUES (:country_name, :region_name, :demand_level, :price_position, :preferred_products, :buyer_types, :entry_points, :market_notes, 999, :t, :t)
        """), {
            'country_name': country_name.strip(), 'region_name': region_name.strip(), 'demand_level': demand_level.strip(),
            'price_position': price_position.strip(), 'preferred_products': preferred_products.strip(), 'buyer_types': buyer_types.strip(),
            'entry_points': entry_points.strip(), 'market_notes': market_notes.strip(), 't': now,
        })
    return redirect('/export/markets?success=created')


@app.get('/export/agent', response_class=HTMLResponse)
def export_agent_page(request: Request,
                      product_name: str = '',
                      target_country: str = '',
                      buyer_type: str = '',
                      goal: str = 'Open a new export conversation',
                      tone: str = 'professional'):
    user = require_login(request)
    result = None
    if product_name or target_country or buyer_type or goal:
        result = generate_export_agent_content(product_name, target_country, buyer_type, goal, tone)
    with engine.begin() as conn:
        products = conn.execute(text('SELECT id, product_name FROM export_products ORDER BY sort_order ASC, id DESC')).fetchall()
        markets = conn.execute(text('SELECT id, country_name FROM export_markets ORDER BY sort_order ASC, id DESC')).fetchall()
    return templates.TemplateResponse('export_agent.html', {
        'request': request,
        'username': user.username,
        'user': user,
        'products': products,
        'markets': markets,
        'result': result,
        'product_name': product_name,
        'target_country': target_country,
        'buyer_type': buyer_type,
        'goal': goal,
        'tone': tone,
        'title': 'AI Export Agent',
        'latest_notification_id': latest_notification_id_for_user(user.username),
    })


ensure_export_engine_tables()


# -----------------------------
# Shipment-centric customer deals & document packs
# -----------------------------
DOCUMENT_TEMPLATE_DIR = DATA_DIR / "document_templates"
DOCUMENT_TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
EXPORT_DOC_UPLOADS_DIR = UPLOADS_DIR / "export_docs"
EXPORT_DOC_UPLOADS_DIR.mkdir(parents=True, exist_ok=True)


def ensure_shipment_deal_tables():
    dialect = engine.dialect.name
    if dialect == 'postgresql':
        ddl = [
            """
            CREATE TABLE IF NOT EXISTS shipment_deals (
                id SERIAL PRIMARY KEY,
                shipment_id INTEGER NOT NULL,
                client_id INTEGER DEFAULT 0,
                client_name TEXT DEFAULT '',
                client_address TEXT DEFAULT '',
                client_country TEXT DEFAULT '',
                ice_no TEXT DEFAULT '',
                phone TEXT DEFAULT '',
                email TEXT DEFAULT '',
                bank_name TEXT DEFAULT '',
                bank_account_no TEXT DEFAULT '',
                bank_iban TEXT DEFAULT '',
                bank_swift TEXT DEFAULT '',
                invoice_no TEXT DEFAULT '',
                invoice_date TEXT DEFAULT '',
                booking_no TEXT DEFAULT '',
                product_name TEXT DEFAULT '',
                product_description TEXT DEFAULT '',
                unit TEXT DEFAULT 'TON',
                qty_ton DOUBLE PRECISION DEFAULT 0,
                cartons_count INTEGER DEFAULT 0,
                net_weight_ton DOUBLE PRECISION DEFAULT 0,
                gross_weight_ton DOUBLE PRECISION DEFAULT 0,
                unit_price DOUBLE PRECISION DEFAULT 0,
                freight_amount DOUBLE PRECISION DEFAULT 0,
                total_amount DOUBLE PRECISION DEFAULT 0,
                currency TEXT DEFAULT 'USD',
                payment_method TEXT DEFAULT 'CAD',
                maturity_date TEXT DEFAULT '',
                consignee_bank_name TEXT DEFAULT '',
                consignee_bank_account TEXT DEFAULT '',
                consignee_bank_iban TEXT DEFAULT '',
                consignee_bank_swift TEXT DEFAULT '',
                consignee_bank_postal_code TEXT DEFAULT '',
                consignee_bank_address TEXT DEFAULT '',
                notes TEXT DEFAULT '',
                status TEXT DEFAULT 'Draft',
                created_by TEXT DEFAULT '',
                created_at DOUBLE PRECISION DEFAULT 0,
                updated_at DOUBLE PRECISION DEFAULT 0
            );
            """,
            """
            CREATE TABLE IF NOT EXISTS deal_documents (
                id SERIAL PRIMARY KEY,
                deal_id INTEGER NOT NULL,
                document_type TEXT NOT NULL,
                file_kind TEXT DEFAULT '',
                filename TEXT NOT NULL,
                original_name TEXT DEFAULT '',
                created_at DOUBLE PRECISION DEFAULT 0,
                created_by TEXT DEFAULT ''
            );
            """
        ]
    else:
        ddl = [
            """
            CREATE TABLE IF NOT EXISTS shipment_deals (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                shipment_id INTEGER NOT NULL,
                client_id INTEGER DEFAULT 0,
                client_name TEXT DEFAULT '',
                client_address TEXT DEFAULT '',
                client_country TEXT DEFAULT '',
                ice_no TEXT DEFAULT '',
                phone TEXT DEFAULT '',
                email TEXT DEFAULT '',
                bank_name TEXT DEFAULT '',
                bank_account_no TEXT DEFAULT '',
                bank_iban TEXT DEFAULT '',
                bank_swift TEXT DEFAULT '',
                invoice_no TEXT DEFAULT '',
                invoice_date TEXT DEFAULT '',
                booking_no TEXT DEFAULT '',
                product_name TEXT DEFAULT '',
                product_description TEXT DEFAULT '',
                unit TEXT DEFAULT 'TON',
                qty_ton REAL DEFAULT 0,
                cartons_count INTEGER DEFAULT 0,
                net_weight_ton REAL DEFAULT 0,
                gross_weight_ton REAL DEFAULT 0,
                unit_price REAL DEFAULT 0,
                freight_amount REAL DEFAULT 0,
                total_amount REAL DEFAULT 0,
                currency TEXT DEFAULT 'USD',
                payment_method TEXT DEFAULT 'CAD',
                maturity_date TEXT DEFAULT '',
                consignee_bank_name TEXT DEFAULT '',
                consignee_bank_account TEXT DEFAULT '',
                consignee_bank_iban TEXT DEFAULT '',
                consignee_bank_swift TEXT DEFAULT '',
                consignee_bank_postal_code TEXT DEFAULT '',
                consignee_bank_address TEXT DEFAULT '',
                notes TEXT DEFAULT '',
                status TEXT DEFAULT 'Draft',
                created_by TEXT DEFAULT '',
                created_at REAL DEFAULT 0,
                updated_at REAL DEFAULT 0
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS deal_documents (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                deal_id INTEGER NOT NULL,
                document_type TEXT NOT NULL,
                file_kind TEXT DEFAULT '',
                filename TEXT NOT NULL,
                original_name TEXT DEFAULT '',
                created_at REAL DEFAULT 0,
                created_by TEXT DEFAULT ''
            )
            """
        ]
    with engine.begin() as conn:
        for stmt in ddl:
            conn.execute(text(stmt))


def _float(v):
    try:
        return float(v or 0)
    except Exception:
        return 0.0


def _int(v):
    try:
        return int(float(v or 0))
    except Exception:
        return 0


def _text(v):
    return (v or '').strip()


def amount_in_words_usd(value: float) -> str:
    units = ['ZERO','ONE','TWO','THREE','FOUR','FIVE','SIX','SEVEN','EIGHT','NINE','TEN','ELEVEN','TWELVE','THIRTEEN','FOURTEEN','FIFTEEN','SIXTEEN','SEVENTEEN','EIGHTEEN','NINETEEN']
    tens = ['', '', 'TWENTY', 'THIRTY', 'FORTY', 'FIFTY', 'SIXTY', 'SEVENTY', 'EIGHTY', 'NINETY']

    def words_under_1000(n: int) -> str:
        parts = []
        if n >= 100:
            parts.append(units[n // 100])
            parts.append('HUNDRED')
            n %= 100
        if n >= 20:
            parts.append(tens[n // 10])
            if n % 10:
                parts.append(units[n % 10])
        elif n > 0:
            parts.append(units[n])
        return ' '.join(parts) if parts else 'ZERO'

    def words(n: int) -> str:
        if n == 0:
            return 'ZERO'
        scales = [(1_000_000_000, 'BILLION'), (1_000_000, 'MILLION'), (1000, 'THOUSAND'), (1, '')]
        out = []
        for scale, label in scales:
            if n >= scale:
                part = n // scale
                n %= scale
                if part:
                    out.append(words_under_1000(part))
                    if label:
                        out.append(label)
        return ' '.join(out)

    whole = int(round(value or 0))
    return f"{words(whole)} USD ONLY"


def parse_form_date(value: str) -> str:
    value = _text(value)
    if not value:
        return ''
    try:
        return datetime.strptime(value, '%Y-%m-%d').strftime('%d-%b-%Y')
    except Exception:
        return value


def format_docx_date(value: str) -> str:
    value = _text(value)
    if not value:
        return ''
    for fmt in ('%d-%b-%Y', '%Y-%m-%d'):
        try:
            return datetime.strptime(value, fmt).strftime('%d %B %Y').upper()
        except Exception:
            pass
    return value.upper()


def maybe_convert_to_pdf(source_path: Path) -> Path | None:
    try:
        outdir = source_path.parent
        subprocess.run([
            'libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', str(outdir), str(source_path)
        ], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        pdf_path = outdir / f"{source_path.stem}.pdf"
        return pdf_path if pdf_path.exists() else None
    except Exception:
        return None


def build_deal_payload(deal, shipment) -> dict:
    qty_ton = _float(getattr(deal, 'qty_ton', 0))
    unit_price = _float(getattr(deal, 'unit_price', 0))
    freight = _float(getattr(deal, 'freight_amount', 0))
    fob = round(qty_ton * unit_price, 2)
    total = round(fob + freight, 2)
    invoice_date = _text(getattr(deal, 'invoice_date', ''))
    if not invoice_date:
        invoice_date = datetime.now().strftime('%d-%b-%Y')
    return {
        'shipment_no': _text(getattr(shipment, 'shipment_no', '')),
        'booking_no': _text(getattr(deal, 'booking_no', '')) or _text(getattr(shipment, 'shipment_no', '')),
        'invoice_no': _text(getattr(deal, 'invoice_no', '')),
        'invoice_date': invoice_date,
        'invoice_date_docx': format_docx_date(invoice_date),
        'client_name': _text(getattr(deal, 'client_name', '')),
        'client_address': _text(getattr(deal, 'client_address', '')),
        'client_country': _text(getattr(deal, 'client_country', '')),
        'ice_no': _text(getattr(deal, 'ice_no', '')),
        'phone': _text(getattr(deal, 'phone', '')),
        'email': _text(getattr(deal, 'email', '')),
        'product_name': _text(getattr(deal, 'product_name', '')),
        'product_description': _text(getattr(deal, 'product_description', '')) or _text(getattr(deal, 'product_name', '')),
        'qty_ton': qty_ton,
        'unit_price': unit_price,
        'fob_amount': fob,
        'freight_amount': freight,
        'total_amount': total,
        'total_amount_words': amount_in_words_usd(total),
        'cartons_count': _int(getattr(deal, 'cartons_count', 0)),
        'net_weight_ton': _float(getattr(deal, 'net_weight_ton', 0)),
        'gross_weight_ton': _float(getattr(deal, 'gross_weight_ton', 0)),
        'currency': _text(getattr(deal, 'currency', 'USD')) or 'USD',
        'payment_method': _text(getattr(deal, 'payment_method', 'CAD')) or 'CAD',
        'maturity_date': _text(getattr(deal, 'maturity_date', '')),
        'maturity_date_docx': format_docx_date(_text(getattr(deal, 'maturity_date', ''))),
        'origin_port': _text(getattr(shipment, 'origin_port', '')),
        'destination_port': _text(getattr(shipment, 'destination_port', '')),
        'vessel_name': _text(getattr(shipment, 'vessel_name', '')),
        'company_name': PRODUCT_PROFILE.get('company_name', 'AL TAHHAN GOLDEN DATES'),
        'consignee_bank_name': _text(getattr(deal, 'consignee_bank_name', '')),
        'consignee_bank_account': _text(getattr(deal, 'consignee_bank_account', '')),
        'consignee_bank_iban': _text(getattr(deal, 'consignee_bank_iban', '')),
        'consignee_bank_swift': _text(getattr(deal, 'consignee_bank_swift', '')),
        'consignee_bank_postal_code': _text(getattr(deal, 'consignee_bank_postal_code', '')),
        'consignee_bank_address': _text(getattr(deal, 'consignee_bank_address', '')),
        'shipper_bank_name': 'Abu Dhabi Commercial Bank - Egypt',
        'shipper_account_no': '1140724710010201',
        'shipper_iban': 'EG090027002001140724710010201',
        'shipper_swift': 'ADCBEGCXXXX',
        'shipper_bank_address': 'Tanta Branch',
    }


def generate_invoice_packing_workbook(payload: dict, target_path: Path) -> list[Path]:
    template_path = DOCUMENT_TEMPLATE_DIR / 'invoice_packing_template.xlsm'
    if not template_path.exists():
        raise RuntimeError('invoice_packing_template.xlsm is missing')
    wb = openpyxl.load_workbook(template_path, keep_vba=True)
    invoice = wb['Invoice']
    packing = wb['Packing List']

    # Fill the real customer invoice template using the exact business cells
    # already present in Altahhan's workbook.
    invoice['C11'] = payload['client_name']
    invoice['G11'] = payload['invoice_date']
    invoice['C12'] = payload['client_address']
    invoice['G12'] = payload['invoice_no']
    invoice['C13'] = payload['ice_no']
    invoice['C14'] = payload['phone']
    invoice['C15'] = payload['email']
    invoice['D20'] = round(payload['unit_price'] + (payload['freight_amount'] / payload['qty_ton'] if payload['qty_ton'] else 0), 2)
    invoice['B22'] = payload['product_description']
    invoice['D22'] = 'TON'
    invoice['E22'] = payload['qty_ton']
    invoice['F22'] = payload['unit_price']
    invoice['G22'] = payload['fob_amount']
    invoice['G23'] = payload['fob_amount']
    invoice['G24'] = payload['freight_amount']
    invoice['G25'] = payload['total_amount']
    invoice['A26'] = f"SAY : {payload['total_amount_words']}"
    invoice['A28'] = f"GROSS WEIGHT : {payload['gross_weight_ton']:,.3f} TON"
    invoice['A29'] = f"NET WEIGHT : {payload['net_weight_ton']:,.3f} TON"
    invoice['A30'] = f"CARTON : {payload['cartons_count']}"
    invoice['A33'] = f"PORT OF LOADING :{payload['origin_port'] or '-'}"
    invoice['A34'] = f"PORT OF DISCHARGE :{payload['destination_port'] or '-'}"
    shipping_method = payload['payment_method'] if payload['payment_method'] in {'FOB', 'CFR', 'CIF'} else 'CFR'
    invoice['A35'] = f"SHIPPING METHOD : {shipping_method}"

    # Packing List mirrors the invoice workbook layout used by the customer.
    packing['C9'] = payload['client_name']
    packing['F9'] = payload['invoice_date']
    packing['C10'] = payload['client_address']
    packing['F10'] = payload['invoice_no']
    packing['C11'] = payload['ice_no']
    packing['C12'] = payload['phone']
    packing['C13'] = payload['email']
    packing['C16'] = payload['product_description']
    packing['D16'] = payload['cartons_count']
    packing['E16'] = payload['gross_weight_ton']
    packing['F16'] = payload['net_weight_ton']
    packing['D17'] = payload['cartons_count']
    packing['E17'] = payload['gross_weight_ton']
    packing['F17'] = payload['net_weight_ton']
    packing['A18'] = f"Total Carton :{payload['cartons_count']} Carton - Net Weight {int(round(payload['net_weight_ton']*1000))} Kg - Gross Weight {int(round(payload['gross_weight_ton']*1000))} Kg"
    packing['A21'] = f"Booking No : {payload['booking_no']}"
    packing['A22'] = f"Exporter: {payload['company_name'].upper()}"
    packing['A23'] = "Exporter address : New Valley- Egypt"
    packing['A24'] = f"Producer: {payload['company_name']}"
    packing['A25'] = "ORIGIN OF GOODS EGYPT"

    wb.save(target_path)
    created = [target_path]
    pdf = maybe_convert_to_pdf(target_path)
    if pdf:
        created.append(pdf)
    return created


def replace_docx_text(document: Document, replacements: dict[str, str]):
    def replace_in_para(para):
        text = para.text
        new_text = text
        for old, new in replacements.items():
            new_text = new_text.replace(old, new)
        if new_text != text:
            para.text = new_text

    for para in document.paragraphs:
        replace_in_para(para)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para)


def generate_cad_docx(payload: dict, target_path: Path) -> list[Path]:
    template_path = DOCUMENT_TEMPLATE_DIR / 'cad_template.docx'
    doc = Document(template_path)

    # The real Altahhan CAD file is used as the template.  Instead of relying on
    # old sample phrases, update paragraphs by their business prefixes so the
    # generated file follows the same structure as the uploaded customer-ready
    # document.
    total_amount_int = int(round(payload['total_amount']))
    total_words = payload['total_amount_words']
    updates = {
        'SHIPPER:': f"SHIPPER:   {payload['company_name']} - New Valley Egypt.  Office 417 Harram st , Giza .",
        'Consignee:': f"Consignee: {payload['client_name']}",
        'Address:': f"Address: {payload['client_address']}",
        'Email:': f"Email: {payload['email']}",
        'BANK NAME: Abu Dhabi Commercial Bank - Egypt': f"BANK NAME: {payload['shipper_bank_name']}                               ACC NO: {payload['shipper_account_no']}",
        'IBAN ID:': f"IBAN ID:    {payload['shipper_iban']}",
        'SWIFT CODE: ADCBEGCXXXX': f"SWIFT CODE: {payload['shipper_swift']}                                             ADRESS:  {payload['shipper_bank_address']}",
        'Acc Name:': f"Acc Name: {payload['client_name']}",
        'BANK NAME: AttijariWafa bank': f"BANK NAME: {payload['consignee_bank_name'] or ''}",
        'ACC NO:': f"ACC NO: {payload['consignee_bank_account'] or payload['consignee_bank_iban'] or ''}",
        'IBAN:': f"IBAN: {payload['consignee_bank_iban'] or ''}",
        'Agency Name': f"Agency Name : {payload['consignee_bank_name'] or ''}   Postal Code : {payload['consignee_bank_postal_code'] or ''}",
        "ADRESS: CENTRE D'AFFAIRES": f"ADRESS: {payload['consignee_bank_address'] or ''}",
        'That Cover is Cash against Documents': f"That Cover is Cash against Documents With amount {total_amount_int} {payload['currency']} ({total_words})",
        'Inv number': f"Inv number {payload['invoice_no']} .",
    }
    swifts_done = 0
    for para in doc.paragraphs:
        line = para.text.strip()
        if not line:
            continue
        if line.startswith('SWIFT CODE:'):
            swifts_done += 1
            if swifts_done == 1:
                para.text = f"SWIFT CODE: {payload['shipper_swift']}                                             ADRESS:  {payload['shipper_bank_address']}"
            else:
                para.text = f"SWIFT CODE: {payload['consignee_bank_swift'] or ''}"
            continue
        for prefix, replacement in updates.items():
            if line.startswith(prefix):
                para.text = replacement
                break

    # Update table counts if present.
    if doc.tables:
        qty_map = {
            'BL Original': '3',
            'Invoice': '3',
            'Packing list': '3',
            'Origin certificate': '1',
            'HEALTH certificate': '1',
            'Phytosanitary certificate': '1',
            'Export Cr SAD': '1',
            'Total of Documents': '13 Document',
        }
        for row in doc.tables[0].rows:
            key = row.cells[0].text.strip()
            for label, qty in qty_map.items():
                if key.startswith(label):
                    row.cells[1].text = qty
                    break

    doc.save(target_path)
    created = [target_path]
    pdf = maybe_convert_to_pdf(target_path)
    if pdf:
        created.append(pdf)
    return created


def generate_bill_docx(payload: dict, target_path: Path) -> list[Path]:
    template_path = DOCUMENT_TEMPLATE_DIR / 'bill_template.docx'
    doc = Document(template_path)
    replacements = {
        '20 DECEMBER 2025                     Inv 1429': f"{payload['invoice_date_docx']}                     Inv {payload['invoice_no']}",
        'NAKHIL EL OUADI ELMAHABA AGHABI': payload['client_name'],
        'au276, BD IBN TACHEFINE ETG3.3 CASABLANCA, MOROCCO': payload['client_address'],
        '003557512000068': payload['ice_no'] or '',
        '+2120771181590': payload['phone'] or '',
        '106000 USD  ) ONE HUNDRED SIX THOUSAND USD ONLY)': f"{int(round(payload['total_amount']))} {payload['currency']} ) {payload['total_amount_words']})",
        '31 March 2026': payload['maturity_date_docx'] or payload['invoice_date_docx'],
        'Attijariwafa bank': payload['consignee_bank_name'] or 'Attijariwafa bank',
        '0000 105000005105': payload['consignee_bank_account'] or '',
        'MA64 007 780 0000105000005105 34': payload['consignee_bank_iban'] or '',
        'BCMAMAMC': payload['consignee_bank_swift'] or '',
        '20300': payload['consignee_bank_postal_code'] or '',
        'C.A Dakar SA Dakar - CASABLANCA. MOROCCO': payload['consignee_bank_address'] or '',
    }
    replace_docx_text(doc, replacements)
    doc.save(target_path)
    created = [target_path]
    pdf = maybe_convert_to_pdf(target_path)
    if pdf:
        created.append(pdf)
    return created


def register_generated_files(conn, deal_id: int, files: list[Path], created_by: str):
    now = time.time()
    for path in files:
        suffix = path.suffix.lower()
        if path.name.endswith('.xlsm'):
            doc_type = 'invoice_packing'
            file_kind = 'excel'
        elif 'cad_' in path.name:
            doc_type = 'cad_cover'
            file_kind = 'pdf' if suffix == '.pdf' else 'docx'
        elif 'bill_' in path.name:
            doc_type = 'bill_of_exchange'
            file_kind = 'pdf' if suffix == '.pdf' else 'docx'
        else:
            doc_type = 'invoice_packing'
            file_kind = 'pdf' if suffix == '.pdf' else suffix.lstrip('.')
        conn.execute(text("""
            INSERT INTO deal_documents (deal_id, document_type, file_kind, filename, original_name, created_at, created_by)
            VALUES (:deal_id, :document_type, :file_kind, :filename, :original_name, :created_at, :created_by)
        """), {
            'deal_id': deal_id,
            'document_type': doc_type,
            'file_kind': file_kind,
            'filename': path.name,
            'original_name': path.name,
            'created_at': now,
            'created_by': created_by,
        })


def generate_deal_document_pack(conn, deal_id: int, username: str) -> dict:
    deal = conn.execute(text('SELECT * FROM shipment_deals WHERE id=:id'), {'id': deal_id}).fetchone()
    if not deal:
        raise HTTPException(status_code=404, detail='Deal not found')
    shipment = conn.execute(text('SELECT * FROM shipments WHERE id=:id'), {'id': deal.shipment_id}).fetchone()
    if not shipment:
        raise HTTPException(status_code=404, detail='Shipment not found')
    payload = build_deal_payload(deal, shipment)
    ts_slug = f"deal_{deal_id}_{int(time.time())}"
    files = []
    files.extend(generate_invoice_packing_workbook(payload, EXPORT_DOC_UPLOADS_DIR / f'invoice_packing_{ts_slug}.xlsm'))
    if payload['payment_method'].upper() == 'CAD':
        files.extend(generate_cad_docx(payload, EXPORT_DOC_UPLOADS_DIR / f'cad_{ts_slug}.docx'))
    else:
        files.extend(generate_bill_docx(payload, EXPORT_DOC_UPLOADS_DIR / f'bill_{ts_slug}.docx'))
    zip_path = EXPORT_DOC_UPLOADS_DIR / f'document_pack_{ts_slug}.zip'
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        for path in files:
            if path.exists():
                zf.write(path, arcname=path.name)
    files.append(zip_path)
    register_generated_files(conn, deal_id, files, username)
    conn.execute(text("UPDATE shipment_deals SET status='Documents Generated', updated_at=:t WHERE id=:id"), {'t': time.time(), 'id': deal_id})
    log_activity(username, 'generate_document_pack', 'shipment_deal', deal_id, payload['invoice_no'] or str(deal_id), conn=conn)
    return {'payload': payload, 'files': files, 'shipment': shipment, 'deal': deal}


@app.post('/shipment/{shipment_id}/deal/add')
def add_customer_deal(request: Request,
                      shipment_id: int,
                      client_mode: str = Form('existing'),
                      client_id: int = Form(0),
                      client_name: str = Form(''),
                      client_address: str = Form(''),
                      client_country: str = Form(''),
                      ice_no: str = Form(''),
                      phone: str = Form(''),
                      email: str = Form(''),
                      bank_name: str = Form(''),
                      bank_account_no: str = Form(''),
                      bank_iban: str = Form(''),
                      bank_swift: str = Form(''),
                      invoice_no: str = Form(...),
                      invoice_date: str = Form(''),
                      booking_no: str = Form(''),
                      product_name: str = Form(''),
                      product_description: str = Form(''),
                      qty_ton: float = Form(0),
                      cartons_count: int = Form(0),
                      net_weight_ton: float = Form(0),
                      gross_weight_ton: float = Form(0),
                      unit_price: float = Form(0),
                      freight_amount: float = Form(0),
                      currency: str = Form('USD'),
                      payment_method: str = Form('CAD'),
                      maturity_date: str = Form(''),
                      consignee_bank_name: str = Form(''),
                      consignee_bank_account: str = Form(''),
                      consignee_bank_iban: str = Form(''),
                      consignee_bank_swift: str = Form(''),
                      consignee_bank_postal_code: str = Form(''),
                      consignee_bank_address: str = Form(''),
                      notes: str = Form('')):
    user = require_editor(request)
    now = time.time()
    with engine.begin() as conn:
        shipment = conn.execute(text('SELECT * FROM shipments WHERE id=:id'), {'id': shipment_id}).fetchone()
        if not shipment:
            return redirect('/shipments?error=shipment_not_found')
        payload = {
            'client_name': _text(client_name), 'client_address': _text(client_address), 'client_country': _text(client_country),
            'ice_no': _text(ice_no), 'phone': _text(phone), 'email': _text(email),
            'bank_name': _text(bank_name), 'bank_account_no': _text(bank_account_no), 'bank_iban': _text(bank_iban), 'bank_swift': _text(bank_swift),
        }
        if client_mode == 'existing' and _int(client_id) > 0:
            client = conn.execute(text('SELECT * FROM current_clients WHERE id=:id'), {'id': _int(client_id)}).fetchone()
            if client:
                payload.update({
                    'client_name': _text(getattr(client, 'company_en', '')) or _text(getattr(client, 'company_ar', '')),
                    'client_address': _text(getattr(client, 'address', '')),
                    'client_country': _text(getattr(client, 'country', '')),
                    'ice_no': _text(getattr(client, 'ice', '')),
                    'bank_name': _text(getattr(client, 'bank_name', '')),
                    'bank_iban': _text(getattr(client, 'iban_account', '')),
                    'bank_swift': _text(getattr(client, 'swift_code', '')),
                })
        deal_id = insert_and_get_id(conn, """
            INSERT INTO shipment_deals (
                shipment_id, client_id, client_name, client_address, client_country, ice_no, phone, email,
                bank_name, bank_account_no, bank_iban, bank_swift, invoice_no, invoice_date, booking_no,
                product_name, product_description, qty_ton, cartons_count, net_weight_ton, gross_weight_ton,
                unit_price, freight_amount, total_amount, currency, payment_method, maturity_date,
                consignee_bank_name, consignee_bank_account, consignee_bank_iban, consignee_bank_swift,
                consignee_bank_postal_code, consignee_bank_address, notes, status, created_by, created_at, updated_at
            ) VALUES (
                :shipment_id, :client_id, :client_name, :client_address, :client_country, :ice_no, :phone, :email,
                :bank_name, :bank_account_no, :bank_iban, :bank_swift, :invoice_no, :invoice_date, :booking_no,
                :product_name, :product_description, :qty_ton, :cartons_count, :net_weight_ton, :gross_weight_ton,
                :unit_price, :freight_amount, :total_amount, :currency, :payment_method, :maturity_date,
                :consignee_bank_name, :consignee_bank_account, :consignee_bank_iban, :consignee_bank_swift,
                :consignee_bank_postal_code, :consignee_bank_address, :notes, 'Draft', :created_by, :created_at, :updated_at
            )
        """, {
            'shipment_id': shipment_id,
            'client_id': _int(client_id),
            'client_name': payload['client_name'],
            'client_address': payload['client_address'],
            'client_country': payload['client_country'],
            'ice_no': payload['ice_no'],
            'phone': payload['phone'],
            'email': payload['email'],
            'bank_name': payload['bank_name'],
            'bank_account_no': payload['bank_account_no'],
            'bank_iban': payload['bank_iban'],
            'bank_swift': payload['bank_swift'],
            'invoice_no': _text(invoice_no),
            'invoice_date': parse_form_date(invoice_date) or datetime.now().strftime('%d-%b-%Y'),
            'booking_no': _text(booking_no) or _text(getattr(shipment, 'shipment_no', '')),
            'product_name': _text(product_name) or _text(getattr(shipment, 'product_name', '')),
            'product_description': _text(product_description) or _text(product_name) or _text(getattr(shipment, 'product_name', '')),
            'qty_ton': _float(qty_ton),
            'cartons_count': _int(cartons_count),
            'net_weight_ton': _float(net_weight_ton),
            'gross_weight_ton': _float(gross_weight_ton),
            'unit_price': _float(unit_price),
            'freight_amount': _float(freight_amount),
            'total_amount': round((_float(qty_ton) * _float(unit_price)) + _float(freight_amount), 2),
            'currency': _text(currency) or 'USD',
            'payment_method': _text(payment_method) or 'CAD',
            'maturity_date': parse_form_date(maturity_date),
            'consignee_bank_name': _text(consignee_bank_name) or payload['bank_name'],
            'consignee_bank_account': _text(consignee_bank_account) or payload['bank_account_no'],
            'consignee_bank_iban': _text(consignee_bank_iban) or payload['bank_iban'],
            'consignee_bank_swift': _text(consignee_bank_swift) or payload['bank_swift'],
            'consignee_bank_postal_code': _text(consignee_bank_postal_code),
            'consignee_bank_address': _text(consignee_bank_address),
            'notes': _text(notes),
            'created_by': user.username,
            'created_at': now,
            'updated_at': now,
        })
        log_activity(user.username, 'create_shipment_deal', 'shipment_deal', deal_id, _text(invoice_no), conn=conn)
    return redirect(f'/deal/{deal_id}?success=deal_created')


@app.get('/deal/{deal_id}', response_class=HTMLResponse)
def deal_detail(request: Request, deal_id: int):
    user = require_login(request)
    with engine.begin() as conn:
        deal = conn.execute(text('SELECT * FROM shipment_deals WHERE id=:id'), {'id': deal_id}).fetchone()
        if not deal:
            return redirect('/shipments?error=deal_not_found')
        shipment = conn.execute(text('SELECT * FROM shipments WHERE id=:id'), {'id': deal.shipment_id}).fetchone()
        docs = conn.execute(text('SELECT * FROM deal_documents WHERE deal_id=:id ORDER BY id DESC'), {'id': deal_id}).fetchall()
    payload = build_deal_payload(deal, shipment)
    return templates.TemplateResponse('deal_detail.html', {
        'request': request, 'username': user.username, 'user': user,
        'deal': deal, 'shipment': shipment, 'docs': docs, 'payload': payload,
        'title': 'Customer Deal', 'latest_notification_id': latest_notification_id_for_user(user.username)
    })


@app.post('/deal/{deal_id}/generate')
def generate_deal_documents(request: Request, deal_id: int):
    user = require_editor(request)
    with engine.begin() as conn:
        generate_deal_document_pack(conn, deal_id, user.username)
    return redirect(f'/deal/{deal_id}?success=documents_generated')


@app.get('/export/deals', response_class=HTMLResponse)
def deals_index(request: Request, q: str = ''):
    user = require_login(request)
    sql = """
        SELECT d.*, s.shipment_no, s.vessel_name, s.origin_port, s.destination_port
        FROM shipment_deals d
        LEFT JOIN shipments s ON s.id=d.shipment_id
        WHERE 1=1
    """
    params = {'q': f'%{_text(q)}%'}
    if _text(q):
        sql += " AND (d.invoice_no LIKE :q OR d.client_name LIKE :q OR d.product_name LIKE :q OR s.shipment_no LIKE :q)"
    sql += ' ORDER BY d.id DESC'
    with engine.begin() as conn:
        rows = conn.execute(text(sql), params).fetchall()
    return templates.TemplateResponse('deals_index.html', {
        'request': request, 'username': user.username, 'user': user, 'rows': rows, 'q': q,
        'title': 'Customer Deals', 'latest_notification_id': latest_notification_id_for_user(user.username)
    })


ensure_shipment_deal_tables()
